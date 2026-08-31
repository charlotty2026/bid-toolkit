# -*- coding: utf-8 -*-
"""
铁律校验器 - 10项硬约束检查（分层）
==================================

把铁律从"靠Agent自觉"变成"引擎强制校验"。
每次内容输出后必须过这10关，不通过的打回。

分层设计：
  🔴 fatal（废标级）- 不通过=废标风险，必须修复才能提交
    ① 过度承诺零容忍 - 不加招标文件没要求的承诺
    ② 增量审核 - 对照招标文件逐条核对
    ③ 不编数据 - 不自己编数字
    ④ 不编标题 - 标题以招标文件原文为准
    ⑨ 身份锚定 - 投标人角色/合同关系不能写反

  🟡 quality（质量级）- 不通过=扣分但不废标，建议修复
    ⑤ 格式规范 - 字体/缩进/编号连续性
    ⑥ 全角半角 - 中文标点必须全角
    ⑦ AI味扫描 - 禁用AI生成痕迹词汇
    ⑧ 排版格式检查 - 页眉页脚/标题层级/行距/首行缩进
    ⑩ 禁止手打编号 - 标题编号必须由引擎自动生成，禁止在标题文本中手打「第一章/一、/（一）/1、」
"""

import re
import json
from datetime import datetime
from typing import List, Dict, Optional

from bid_toolkit.orchestrator.bidder_profile import BidderProfile


class IronRuleChecker:
    """铁律校验器"""

    # 铁律分层定义
    TIER_FATAL = "fatal"      # 废标级
    TIER_QUALITY = "quality"  # 质量级

    # 每项铁律的层级
    RULE_TIERS = {
        1: TIER_FATAL,    # 过度承诺
        2: TIER_FATAL,    # 增量审核
        3: TIER_FATAL,    # 不编数据
        4: TIER_FATAL,    # 不编标题
        5: TIER_QUALITY,  # 格式规范
        6: TIER_QUALITY,  # 全角半角
        7: TIER_QUALITY,  # AI味扫描
        8: TIER_QUALITY,  # 排版格式检查
        9: TIER_FATAL,    # 身份锚定
        10: TIER_QUALITY, # 禁止手打编号
    }

    def __init__(self, tender_extract: Optional[dict] = None,
                 locked_titles: Optional[List[str]] = None,
                 bidder_profile: Optional[dict] = None):
        """
        Args:
            tender_extract: 招标文件提取结果（用于增量审核和过度承诺检测）
            locked_titles: 锁定的框架标题列表（用于标题篡改检测）
            bidder_profile: 投标人身份卡（用于身份锚定检测）
        """
        self.tender_extract = tender_extract or {}
        self.locked_titles = locked_titles or []
        self.bidder_profile = BidderProfile(bidder_profile) if bidder_profile else None

    def check_all(self, content: str, docx_path: Optional[str] = None) -> dict:
        """
        运行全部铁律检查（9项，分层）

        Args:
            content: 待检查的文本内容（Markdown或纯文本）
            docx_path: 如果有Word文档，检查排版格式

        Returns:
            {
                "all_passed": bool,          # 全部通过
                "fatal_passed": bool,         # 废标级全部通过（true=可提交）
                "total_rules": 10,
                "passed": int,
                "failed": int,
                "fatal_passed_count": int,
                "fatal_failed_count": int,
                "quality_passed_count": int,
                "quality_failed_count": int,
                "rules": [...],
                "timestamp": "..."
            }
        """
        rules = []

        # ① 过度承诺零容忍 [fatal]
        rules.append(self._check_overcommitment(content))

        # ② 增量审核 [fatal]
        rules.append(self._check_incremental(content))

        # ③ 不编数据 [fatal]
        rules.append(self._check_fabricated_data(content))

        # ④ 不编标题 [fatal]
        rules.append(self._check_title_tampering(content))

        # ⑤ 格式规范 [quality]
        rules.append(self._check_format_standard(content))

        # ⑥ 全角半角 [quality]
        rules.append(self._check_punctuation(content))

        # ⑦ AI味扫描 [quality]
        rules.append(self._check_ai_flavor(content))

        # ⑧ 排版格式检查 [quality]（需要docx文件）
        if docx_path:
            rules.append(self._check_layout(docx_path))
        else:
            rules.append({
                "id": 8,
                "name": "排版格式检查",
                "tier": self.TIER_QUALITY,
                "passed": True,
                "skipped": True,
                "message": "未提供docx文件，跳过排版格式检查",
                "issues": []
            })

        # ⑨ 身份锚定 [fatal]
        rules.append(self._check_identity_anchor(content))

        # ⑩ 禁止手打编号 [quality]
        rules.append(self._check_hand_typed_numbering(content))

        # 统一补上tier字段（跳过分支已设的不会重复）
        for rule in rules:
            if "tier" not in rule:
                rule["tier"] = self.RULE_TIERS.get(rule["id"], self.TIER_QUALITY)

        # 统计
        passed = sum(1 for r in rules if r.get("passed", False))
        failed = len(rules) - passed

        fatal_rules = [r for r in rules if r.get("tier") == self.TIER_FATAL]
        quality_rules = [r for r in rules if r.get("tier") == self.TIER_QUALITY]

        fatal_passed = sum(1 for r in fatal_rules if r.get("passed", False))
        fatal_failed = len(fatal_rules) - fatal_passed
        quality_passed = sum(1 for r in quality_rules if r.get("passed", False))
        quality_failed = len(quality_rules) - quality_passed

        return {
            "all_passed": failed == 0,
            "fatal_passed": fatal_failed == 0,  # 废标级全过=可提交
            "total_rules": len(rules),
            "passed": passed,
            "failed": failed,
            "fatal_passed_count": fatal_passed,
            "fatal_failed_count": fatal_failed,
            "quality_passed_count": quality_passed,
            "quality_failed_count": quality_failed,
            "rules": rules,
            "timestamp": datetime.now().isoformat()
        }

    # ===== ① 过度承诺零容忍 =====

    # 招标文件没要求但AI喜欢加的承诺
    OVERCOMMITMENT_KEYWORDS = [
        "保险金额", "意外险", "雇主责任险", "安全生产责任险",
        "薪酬保障金", "培训基金", "帮扶基金", "奖励基金",
        "违约罚款", "违约金", "赔偿金",
        "猎头", "猎才计划",
        "晋升通道", "晋升比例", "晋升数据",
        "节省成本", "节约费用", "降本增效",
        "零投诉", "零事故", "零离职",
        "100%满意", "100%通过", "100%覆盖",
        "全程跟踪", "全程监控", "全天候响应",
    ]

    # 招标文件可能要求的合理承诺（白名单）
    LEGITIMATE_KEYWORDS = [
        "季度考核", "服务质量", "24小时值守", "全年无休",
        "承诺", "保证", "确保",  # 这些词本身没问题，看上下文
    ]

    def _check_overcommitment(self, content: str) -> dict:
        """① 过度承诺零容忍"""
        issues = []

        # 提取招标文件中实际要求的承诺
        tender_commitments = set()
        for req in self.tender_extract.get("requirements", []):
            tender_commitments.add(req.get("text", ""))

        for keyword in self.OVERCOMMITMENT_KEYWORDS:
            if keyword in content:
                # 检查是否在招标文件中出现过
                in_tender = any(keyword in tc for tc in tender_commitments)
                if not in_tender:
                    # 找到上下文
                    idx = content.find(keyword)
                    context = content[max(0, idx - 30):idx + len(keyword) + 30]
                    issues.append({
                        "type": "过度承诺",
                        "keyword": keyword,
                        "context": context,
                        "severity": "high",
                        "suggestion": f"删除「{keyword}」--招标文件未要求此承诺"
                    })

        return {
            "id": 1,
            "name": "过度承诺零容忍",
            "passed": len(issues) == 0,
            "issues": issues,
            "message": f"发现{len(issues)}处过度承诺" if issues else "无过度承诺"
        }

    # ===== ② 增量审核 =====

    def _check_incremental(self, content: str) -> dict:
        """② 增量审核 - 对照招标文件逐条核对"""
        issues = []

        requirements = self.tender_extract.get("requirements", [])
        if not requirements:
            return {
                "id": 2,
                "name": "增量审核",
                "passed": True,
                "skipped": True,
                "issues": [],
                "message": "无招标文件提取结果，跳过增量审核"
            }

        # 检查每个招标文件要求是否在投标文件中有响应
        for req in requirements:
            req_text = req.get("text", "")
            req_keywords = req.get("keywords", [])
            if not req_keywords:
                # 如果没有提取关键词，用原文前20字作为关键词
                req_keywords = [req_text[:20]] if req_text else []

            found = False
            for kw in req_keywords:
                if kw and kw in content:
                    found = True
                    break

            if not found and req.get("required", True):
                issues.append({
                    "type": "需求未响应",
                    "requirement": req_text[:50],
                    "keywords": req_keywords,
                    "severity": "medium",
                    "suggestion": f"招标文件要求「{req_text[:30]}...」未在投标文件中响应"
                })

        return {
            "id": 2,
            "name": "增量审核",
            "passed": len(issues) == 0,
            "issues": issues,
            "message": f"发现{len(issues)}处需求未响应" if issues else "所有需求已响应"
        }

    # ===== ③ 不编数据 =====

    # 可疑数据模式
    FABRICATED_DATA_PATTERNS = [
        (r'(\d+(?:\.\d+)?)\s*%(?:满意度|通过率|合格率|覆盖率)', "百分比数据"),
        (r'(\d+(?:\.\d+)?)\s*(?:万元|元)(?:节约|节省|降本|减少)', "节约金额"),
        (r'(?:离职率|流失率|离职人数)\s*[：:]\s*\d+', "离职率数据"),
        (r'(?:晋升率|晋升比例|晋升人数)\s*[：:]\s*\d+', "晋升数据"),
        (r'(?:培训次数|培训场次|培训人数)\s*[：:]\s*\d+', "培训数据"),
        (r'(?:项目数|案例数|业绩数)\s*[：:]\s*\d+', "项目数据"),
    ]

    def _check_fabricated_data(self, content: str) -> dict:
        """③ 不编数据"""
        issues = []

        for pattern, data_type in self.FABRICATED_DATA_PATTERNS:
            for m in re.finditer(pattern, content):
                context = content[max(0, m.start() - 20):m.end() + 20]
                issues.append({
                    "type": "疑似编造数据",
                    "data_type": data_type,
                    "matched": m.group(),
                    "context": context,
                    "severity": "high",
                    "suggestion": f"核实「{m.group()}」是否为真实数据，如非招标文件要求或真实业绩，删除"
                })

        return {
            "id": 3,
            "name": "不编数据",
            "passed": len(issues) == 0,
            "issues": issues,
            "message": f"发现{len(issues)}处疑似编造数据" if issues else "无编造数据"
        }

    # ===== ④ 不编标题 =====

    def _check_title_tampering(self, content: str) -> dict:
        """④ 不编标题 - 标题以招标文件原文为准"""
        issues = []

        if not self.locked_titles:
            return {
                "id": 4,
                "name": "不编标题",
                "passed": True,
                "skipped": True,
                "issues": [],
                "message": "无锁定框架，跳过标题篡改检测"
            }

        # 提取内容中的标题
        from .framework import extract_titles_from_markdown
        current_titles = extract_titles_from_markdown(content)
        current_title_texts = [t["title"] for t in current_titles]

        # 检查是否有不在锁定框架中的标题
        for ct in current_title_texts:
            found = False
            for lt in self.locked_titles:
                if ct == lt or lt in ct or ct in lt:
                    found = True
                    break
            if not found:
                issues.append({
                    "type": "未经授权的新增标题",
                    "title": ct,
                    "severity": "high",
                    "suggestion": f"标题「{ct}」不在锁定的框架中，需人工确认是否保留"
                })

        return {
            "id": 4,
            "name": "不编标题",
            "passed": len(issues) == 0,
            "issues": issues,
            "message": f"发现{len(issues)}个未经授权的新增标题" if issues else "标题与框架一致"
        }

    # ===== ⑤ 格式规范 =====

    def _check_format_standard(self, content: str) -> dict:
        """⑤ 格式规范 - 字体/缩进/编号连续性"""
        issues = []

        lines = content.split("\n")
        prev_num = 0
        prev_level = 0

        for i, line in enumerate(lines):
            line_stripped = line.strip()
            if not line_stripped:
                continue

            # 检查编号连续性 (1. 2. 3. ...)
            m = re.match(r'^(\d+)\.\s', line_stripped)
            if m:
                num = int(m.group(1))
                if num > 1 and num != prev_num + 1 and prev_num > 0:
                    issues.append({
                        "type": "编号不连续",
                        "line": i + 1,
                        "expected": prev_num + 1,
                        "actual": num,
                        "context": line_stripped[:50],
                        "severity": "low"
                    })
                prev_num = num

            # 检查Markdown残留标记
            for residue in ["**", "##", "###", "```", "---"]:
                if residue in line_stripped and not line_stripped.startswith("#"):
                    # 允许Markdown标题
                    if residue == "##" and line_stripped.startswith("#"):
                        continue
                    if residue == "**":
                        # 允许加粗，但如果整行都是**包裹则可能是残留
                        if line_stripped.count("**") >= 2 and line_stripped.startswith("**"):
                            issues.append({
                                "type": "Markdown残留",
                                "line": i + 1,
                                "residue": residue,
                                "context": line_stripped[:50],
                                "severity": "low"
                            })

        return {
            "id": 5,
            "name": "格式规范",
            "passed": len(issues) == 0,
            "issues": issues,
            "message": f"发现{len(issues)}处格式问题" if issues else "格式规范"
        }

    # ===== ⑥ 全角半角 =====

    HALF_TO_FULL = {
        ',': '，', '.': '。', '!': '！', '?': '？', ':': '：', ';': '；',
        '(': '（', ')': '）',
    }

    def _check_punctuation(self, content: str) -> dict:
        """⑥ 全角半角检查"""
        issues = []

        for i, line in enumerate(content.split("\n")):
            for half, full in self.HALF_TO_FULL.items():
                # 检查中文之间的半角标点
                for m in re.finditer(rf'[\u4e00-\u9fff]{re.escape(half)}[\u4e00-\u9fff]', line):
                    issues.append({
                        "type": "半角标点混入中文",
                        "line": i + 1,
                        "char": half,
                        "should_be": full,
                        "context": line[max(0, m.start() - 5):m.end() + 5],
                        "severity": "medium"
                    })

            # 检查全角数字/字母
            for m in re.finditer(r'[\uff10-\uff19\uff21-\uff3a\uff41-\uff5a]+', line):
                issues.append({
                    "type": "全角数字/字母",
                    "line": i + 1,
                    "char": m.group(),
                    "context": line[max(0, m.start() - 5):m.end() + 5],
                    "severity": "low"
                })

        return {
            "id": 6,
            "name": "全角半角检查",
            "passed": len(issues) == 0,
            "issues": issues,
            "message": f"发现{len(issues)}处标点问题" if issues else "标点规范"
        }

    # ===== ⑦ AI味扫描 =====

    AI_FLAVOR_WORDS = [
        "作为AI", "作为人工智能", "我是一个AI", "我无法", "我不能",
        "温馨提示", "请注意", "值得注意", "需要指出",
        "首先", "其次", "再次", "最后",
        "综上所述", "总而言之", "由此可见",
        "不可或缺", "至关重要", "重中之重",
        "赋能", "助力", "打造", "构建",
        "全面覆盖", "全方位", "多维度", "立体化",
        "致力于", "秉承", "秉持", "坚守",
        "行业领先", "业界标杆", "一流水平",
        "我们深知", "我们始终", "我们将持续",
        "基于此", "在此基础上", "与此同时",
    ]

    def _check_ai_flavor(self, content: str) -> dict:
        """⑦ AI味扫描"""
        issues = []

        for word in self.AI_FLAVOR_WORDS:
            if word in content:
                idx = content.find(word)
                context = content[max(0, idx - 15):idx + len(word) + 15]
                issues.append({
                    "type": "AI味词汇",
                    "word": word,
                    "context": context,
                    "severity": "high",
                    "suggestion": f"删除或替换「{word}」"
                })

        return {
            "id": 7,
            "name": "AI味扫描",
            "passed": len(issues) == 0,
            "issues": issues,
            "message": f"发现{len(issues)}处AI味词汇" if issues else "无AI味"
        }

    # ===== ⑧ 排版格式检查 =====

    def _check_layout(self, docx_path: str) -> dict:
        """⑧ 排版格式检查 - 页眉页脚/标题层级/行距/首行缩进"""
        issues = []

        try:
            from docx import Document
            from docx.shared import Pt
            from docx.oxml.ns import qn

            doc = Document(docx_path)

            # 检查标题层级
            prev_level = 0
            for para in doc.paragraphs:
                style_name = (para.style.name if para.style else "") or ""
                text = para.text.strip()
                if not text:
                    continue

                if style_name.startswith("Heading"):
                    try:
                        level = int(style_name.replace("Heading ", ""))
                    except ValueError:
                        level = 1

                    # 检查层级跳跃（如从H1直接跳到H3）
                    if prev_level > 0 and level > prev_level + 1:
                        issues.append({
                            "type": "标题层级跳跃",
                            "text": text[:30],
                            "from_level": prev_level,
                            "to_level": level,
                            "severity": "medium"
                        })
                    prev_level = level

                # 检查正文首行缩进
                elif style_name == "Normal" and text:
                    pf = para.paragraph_format
                    if pf.first_line_indent is None or pf.first_line_indent == 0:
                        # 检查是否是表格内文本（表格内不需要首行缩进）
                        parent = para._element.getparent()
                        if parent is not None and parent.tag.endswith('}tc'):
                            continue  # 表格内，跳过
                        issues.append({
                            "type": "缺少首行缩进",
                            "text": text[:30],
                            "severity": "low"
                        })

                    # 检查行距
                    if pf.line_spacing is not None and pf.line_spacing < 1.0:
                        issues.append({
                            "type": "行距过小",
                            "text": text[:30],
                            "line_spacing": pf.line_spacing,
                            "severity": "low"
                        })

            # 检查页眉页脚
            for section in doc.sections:
                header = section.header
                footer = section.footer
                if header and not header.paragraphs[0].text.strip():
                    issues.append({
                        "type": "页眉为空",
                        "severity": "medium"
                    })
                if footer and not footer.paragraphs[0].text.strip():
                    issues.append({
                        "type": "页脚为空",
                        "severity": "medium"
                    })

        except ImportError:
            return {
                "id": 8,
                "name": "排版格式检查",
                "passed": True,
                "skipped": True,
                "issues": [],
                "message": "未安装python-docx，跳过排版格式检查"
            }
        except Exception as e:
            return {
                "id": 8,
                "name": "排版格式检查",
                "passed": False,
                "issues": [],
                "message": f"检查失败: {e}"
            }

        return {
            "id": 8,
            "name": "排版格式检查",
            "passed": len(issues) == 0,
            "issues": issues,
            "message": f"发现{len(issues)}处排版问题" if issues else "排版规范"
        }

    # ===== ⑨ 身份锚定 =====

    def _check_identity_anchor(self, content: str) -> dict:
        """⑨ 身份锚定 - 投标人角色/合同关系不能写反 [fatal]"""
        if not self.bidder_profile:
            return {
                "id": 9,
                "name": "身份锚定",
                "passed": True,
                "skipped": True,
                "issues": [],
                "message": "未设置投标人身份卡，跳过身份锚定检查（建议执行 profile --set-file <json>）"
            }

        if not self.bidder_profile.is_complete():
            return {
                "id": 9,
                "name": "身份锚定",
                "passed": False,
                "issues": [],
                "message": "⚠️ 投标人身份卡未填写完整，无法进行身份锚定检查"
            }

        issues = self.bidder_profile.check_content(content)

        return {
            "id": 9,
            "name": "身份锚定",
            "passed": len(issues) == 0,
            "issues": issues,
            "message": f"发现{len(issues)}处身份关系问题" if issues else "身份关系正确"
        }

    # ===== ⑩ 禁止手打编号 =====

    # 手打编号特征：标题/行首出现这些前缀 = 应交给引擎自动编号（numPr）
    HAND_TYPED_NUMBERING_PATTERNS = [
        (r'^#*\s*第[一二三四五六七八九十百零]+章', "第一章/第X章"),
        (r'^#*\s*[一二三四五六七八九十]+、', "一、"),
        (r'^#*\s*[（(][一二三四五六七八九十]+[）)]', "（一）"),
        (r'^#*\s*\d+[、.]', "1、/1."),
    ]

    def _check_hand_typed_numbering(self, content: str) -> dict:
        """⑩ 禁止手打编号 - 标题编号必须由引擎自动生成 [quality]

        死规矩：编号一律由 render/排版引擎走 numbering.xml 自动生成，
        禁止在标题文本里手打「第一章 / 一、 / （一） / 1、」等前缀，
        否则会出现「第一章 第一章 运营模式」双重编号。
        """
        issues = []
        for i, line in enumerate(content.split("\n")):
            stripped = line.strip()
            if not stripped:
                continue
            for pat, label in self.HAND_TYPED_NUMBERING_PATTERNS:
                if re.match(pat, stripped):
                    issues.append({
                        "type": "手打编号",
                        "label": label,
                        "line": i + 1,
                        "text": stripped[:40],
                        "severity": "high",
                        "suggestion": "删除标题文本中的手打编号「%s」，编号由引擎自动生成（numPr）" % label
                    })
                    break
        return {
            "id": 10,
            "name": "禁止手打编号",
            "passed": len(issues) == 0,
            "issues": issues,
            "message": f"发现{len(issues)}处手打编号" if issues else "标题编号均由引擎自动生成"
        }

    # ===== 报告生成 =====

    @staticmethod
    def format_report(report: dict) -> str:
        """格式化铁律校验报告为人类可读文本（分层显示）"""
        fatal_passed = report.get("fatal_passed", False)

        lines = [
            "=" * 60,
            "铁律校验报告（10项·分层）",
            f"时间: {report.get('timestamp', '')}",
            "=" * 60,
            "",
        ]

        # 总览
        if report["all_passed"]:
            lines.append("🎉 全部10项铁律通过，可以提交！")
        elif fatal_passed:
            lines.append("✅ 废标级全部通过（可提交），质量级有问题建议修复")
        else:
            lines.append("⛔ 废标级铁律未通过，禁止提交！必须修复后重新校验")

        lines.append(
            f"   废标级: {report.get('fatal_passed_count', 0)}通过 "
            f"/ {report.get('fatal_failed_count', 0)}未通过"
        )
        lines.append(
            f"   质量级: {report.get('quality_passed_count', 0)}通过 "
            f"/ {report.get('quality_failed_count', 0)}未通过"
        )
        lines.append("")

        # 废标级
        fatal_rules = [r for r in report.get("rules", []) if r.get("tier") == "fatal"]
        if fatal_rules:
            lines.append("━" * 40)
            lines.append("🔴 废标级（不通过=废标风险）")
            lines.append("━" * 40)
            for rule in fatal_rules:
                status = "✅" if rule.get("passed") else "❌"
                if rule.get("skipped"):
                    status = "⏭️"
                lines.append(f"{status} 铁律{rule['id']}: {rule['name']}")
                lines.append(f"   {rule.get('message', '')}")
                for issue in rule.get("issues", []):
                    severity = issue.get("severity", "medium")
                    icon = {"fatal": "💀", "high": "🔴", "medium": "🟡", "low": "🔵"}.get(severity, "🟡")
                    suggestion = issue.get("suggestion", issue.get("type", ""))
                    lines.append(f"   {icon} {suggestion}")
                lines.append("")

        # 质量级
        quality_rules = [r for r in report.get("rules", []) if r.get("tier") == "quality"]
        if quality_rules:
            lines.append("━" * 40)
            lines.append("🟡 质量级（不通过=扣分建议修复）")
            lines.append("━" * 40)
            for rule in quality_rules:
                status = "✅" if rule.get("passed") else "❌"
                if rule.get("skipped"):
                    status = "⏭️"
                lines.append(f"{status} 铁律{rule['id']}: {rule['name']}")
                lines.append(f"   {rule.get('message', '')}")
                for issue in rule.get("issues", []):
                    severity = issue.get("severity", "medium")
                    icon = {"fatal": "💀", "high": "🔴", "medium": "🟡", "low": "🔵"}.get(severity, "🟡")
                    suggestion = issue.get("suggestion", issue.get("type", ""))
                    lines.append(f"   {icon} {suggestion}")
                lines.append("")

        # 结论
        lines.append("=" * 60)
        if report["all_passed"]:
            lines.append("🎉 全部铁律通过，可以提交！")
        elif fatal_passed:
            lines.append("✅ 废标级通过，质量级建议修复后提交")
        else:
            lines.append("⛔ 废标级铁律未通过，禁止提交！")

        return "\n".join(lines)
