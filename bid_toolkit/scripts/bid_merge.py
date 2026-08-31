#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
标书拼装器 v1.0
================
多章节合并 + 格式统一 + 占位符检查 + 终审自检

用法:
  python bid_merge.py -d chapters/ -o 投标书_完整版.docx
  python bid_merge.py -d chapters/ -c config.json -o 投标书.docx
"""

import os, sys, re, json, argparse
from pathlib import Path

sys.stdout.reconfigure(encoding='utf-8')

try:
    from docx import Document
    from docx.shared import Pt, Twips, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("❌ 需要 python-docx: pip install python-docx")
    sys.exit(1)

# ============================================================
#  字体工具
# ============================================================

def set_run_font(run, font_name='宋体', size=12, bold=False):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    for attr in ['w:asciiTheme', 'w:hAnsiTheme', 'w:eastAsiaTheme']:
        try:
            del rFonts.attrib[qn(attr)]
        except KeyError:
            pass

# ============================================================
#  拼装
# ============================================================

def merge_chapters(chapter_dir, output_path, config=None):
    """合并所有章节docx为一个文档"""
    chapter_dir = Path(chapter_dir)
    
    # 按文件名排序（01_xxx, 02_xxx...）
    docx_files = sorted(chapter_dir.glob('*.docx'), 
                        key=lambda f: f.name)
    
    if not docx_files:
        print(f'❌ 目录中没有docx文件: {chapter_dir}')
        return None
    
    print(f'📚 发现 {len(docx_files)} 个章节:')
    for f in docx_files:
        print(f'  {f.name}')
    
    # 创建合并文档
    merged = Document()
    
    # 设置页面（标准A4）
    for section in merged.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)
    
    # 初始化Normal样式
    try:
        style = merged.styles['Normal']
        style.font.name = '宋体'
        style.font.size = Pt(12)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    except KeyError:
        pass
    
    # 逐章合并
    total_paras = 0
    total_tables = 0
    
    for docx_file in docx_files:
        print(f'  合并: {docx_file.name}...', end=' ')
        doc = Document(str(docx_file))
        
        # 复制段落
        for para in doc.paragraphs:
            # 跳过空段落
            if not para.text.strip() and not para.runs:
                continue
            
            # 根据样式添加
            if para.style.name.startswith('Heading'):
                level = int(para.style.name.split()[-1])
                new_para = merged.add_heading(para.text, level=level)
            elif para.style.name == 'List Bullet':
                new_para = merged.add_paragraph(style='List Bullet')
                new_para.clear()
                run = new_para.add_run(para.text)
                set_run_font(run, '宋体', 12, False)
            elif 'Quote' in para.style.name:
                new_para = merged.add_paragraph()
                run = new_para.add_run(para.text)
                set_run_font(run, '宋体', 12, False)
                new_para.paragraph_format.left_indent = Cm(1)
            else:
                new_para = merged.add_paragraph(para.text, style='Normal')
            
            # 复制格式（字体/缩进/行距）
            for run in new_para.runs:
                set_run_font(run, '宋体', 12, False)
            new_para.paragraph_format.first_line_indent = Twips(480)
            new_para.paragraph_format.line_spacing = 1.5
            new_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            total_paras += 1
        
        # 复制表格
        for table in doc.tables:
            # 复制表格数据
            rows = len(table.rows)
            cols = len(table.columns)
            if rows == 0 or cols == 0:
                continue
            
            new_table = merged.add_table(rows=rows, cols=cols)
            new_table.style = 'Table Grid'
            
            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    new_table.rows[i].cells[j].text = cell.text
                    # 设置字体
                    for para in new_table.rows[i].cells[j].paragraphs:
                        for run in para.runs:
                            set_run_font(run, '宋体', 10.5, False)
            
            total_tables += 1
        
        # 章节间加分页符（最后一个章节不加）
        if docx_file != docx_files[-1]:
            merged.add_page_break()
        
        print(f'{len(doc.paragraphs)}段落 + {len(doc.tables)}表格')
    
    # 保存
    merged.save(output_path)
    print(f'\n✅ 合并完成: {output_path}')
    print(f'  总计: {total_paras}段落, {total_tables}表格')
    
    return output_path

# ============================================================
#  占位符检查
# ============================================================

def check_placeholders(docx_path):
    """检查文档中的未填写占位符"""
    doc = Document(docx_path)
    issues = []
    
    patterns = [
        (r'\[XX\]', '未填占位符'),
        (r'\[.*?XX.*?\]', '未填占位符'),
        (r'（投标人名称）', '投标人名称未填'),
        (r'（项目名称）', '项目名称未填'),
        (r'（招标编号）', '招标编号未填'),
        (r'\d{4}年\d{1,2}月\s+日', '日期未更新'),
        (r'XXX', '占位符XXX未替换'),
        (r'_____', '空白下划线未填写'),
    ]
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        for pattern, desc in patterns:
            for match in re.finditer(pattern, text):
                issues.append({
                    'paragraph': i+1,
                    'type': desc,
                    'match': match.group(),
                    'context': text[max(0,match.start()-20):match.end()+20]
                })
    
    if issues:
        print(f'\n⚠️  发现 {len(issues)} 处占位符/未填写:')
        for issue in issues:
            print(f'  段落{issue["paragraph"]}: [{issue["type"]}] "{issue["match"]}"')
    else:
        print('✅ 无未填写占位符')
    
    return issues

# ============================================================
#  CLI
# ============================================================

def main():
    parser = argparse.ArgumentParser(description='标书拼装器 v1.0')
    parser.add_argument('-d', '--dir', required=True, help='章节目录(含多个docx)')
    parser.add_argument('-o', '--output', default='投标书_完整版.docx', help='输出文件')
    parser.add_argument('-c', '--config', default=None, help='排版配置JSON')
    parser.add_argument('--check', action='store_true', help='拼装后检查占位符')
    args = parser.parse_args()
    
    config = None
    if args.config and os.path.exists(args.config):
        with open(args.config, 'r', encoding='utf-8') as f:
            config = json.load(f)
    
    result = merge_chapters(args.dir, args.output, config)
    
    if result and args.check:
        print('\n📋 占位符检查:')
        check_placeholders(result)
    
    sys.exit(0)

if __name__ == '__main__':
    main()
