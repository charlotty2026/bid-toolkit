#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
工程标两列表格生成器 v1.0
输入：DOC/DOCX/PDF文件 → 提取文字 → 自动识别"标签：内容"对
输出：Word文档，左列标签+右列内容的两列表格，样式全可调

用法：
  python bid_table_generator.py 输入文件.docx -o 输出.docx
  python bid_table_generator.py 输入文件.pdf -o 输出.docx
  python bid_table_generator.py 输入文件.docx -o 输出.docx --config my.yaml
  python bid_table_generator.py 输入文件.docx -o 输出.docx --left-width 3 --right-width 5
  python bid_table_generator.py 输入文件.docx -o 输出.docx --left-color D9E2F3 --right-color FFFFFF
  python bid_table_generator.py 输入文件.docx -o 输出.docx --font-size 12 --font-name 宋体
  python bid_table_generator.py 输入文件.docx --no-merge    # 不分组合并，每行独立
  python bid_table_generator.py 输入文件.docx --list         # 仅列出识别的标签-内容对
"""

import os, sys, re, json, argparse, yaml
from pathlib import Path
from typing import List, Tuple, Optional, Dict

# 文档提取
try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, Inches, Cm, RGBColor, Twips, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import fitz  # PyMuPDF
    HAS_PDF = True
except ImportError:
    HAS_PDF = False


# ===== 默认配置 =====
DEFAULT_CONFIG = {
    'table': {
        'left_width_cm': 4.0,      # 左列宽度（标签列）
        'right_width_cm': 10.0,    # 右列宽度（内容列）
        'border_color': '000000',  # 边框颜色（十六进制，无#）
        'border_size': 1,          # 边框粗细（pt）
    },
    'left_cell': {
        'font_name': '黑体',
        'font_size': 11,           # pt
        'bold': True,
        'color': '000000',         # 字体颜色
        'bg_color': 'D9E2F3',      # 背景色（浅蓝灰）
        'alignment': 'center',     # left/center/right
        'vertical_alignment': 'center',
    },
    'right_cell': {
        'font_name': '宋体',
        'font_size': 11,
        'bold': False,
        'color': '000000',
        'bg_color': 'FFFFFF',      # 白色
        'alignment': 'left',
        'vertical_alignment': 'center',
    },
    'header': {
        'enabled': False,          # 是否显示表头行
        'text': '主要技术参数',
        'font_name': '黑体',
        'font_size': 14,
        'bold': True,
        'color': 'FFFFFF',
        'bg_color': '4472C4',      # 深蓝色
    },
    'section': {
        'merge_same_left': True,   # 左列相同内容是否合并（如"一、施工方案"合并多行）
        'section_font_name': '黑体',
        'section_font_size': 12,
        'section_bold': True,
        'section_bg_color': 'B4C6E7',  # 章节标题行背景色
    },
    'extraction': {
        'pattern': r'^([^：:]+)[：:]\s*(.*)',  # 识别"标签：内容"的正则
        'min_label_length': 1,     # 标签最小长度
        'max_label_length': 30,    # 标签最大长度
        'skip_lines_startswith': ['#', '//', '/*', '*', '注', '说明'],  # 跳过的行前缀
        'skip_empty_labels': True, # 跳过空标签
    },
    'page': {
        'margin_top_cm': 2.54,
        'margin_bottom_cm': 2.54,
        'margin_left_cm': 3.17,
        'margin_right_cm': 3.17,
        'orientation': 'portrait',  # portrait/landscape
    },
}


# ===== 颜色工具 =====
def hex_to_rgb(hex_color: str) -> RGBColor:
    """十六进制颜色转RGBColor对象"""
    hex_color = hex_color.lstrip('#')
    if len(hex_color) == 3:
        hex_color = ''.join(c*2 for c in hex_color)
    try:
        return RGBColor(int(hex_color[0:2], 16),
                        int(hex_color[2:4], 16),
                        int(hex_color[4:6], 16))
    except (ValueError, IndexError):
        return RGBColor(0, 0, 0)


def set_cell_shading(cell, color_hex: str):
    """设置单元格背景色"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_cell_border(cell, color_hex: str = '000000', size: int = 4):
    """设置单元格边框"""
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{size}" w:color="{color_hex}"/>'
        f'  <w:left w:val="single" w:sz="{size}" w:color="{color_hex}"/>'
        f'  <w:bottom w:val="single" w:sz="{size}" w:color="{color_hex}"/>'
        f'  <w:right w:val="single" w:sz="{size}" w:color="{color_hex}"/>'
        f'</w:tcBorders>'
    )
    cell._tc.get_or_add_tcPr().append(borders)


def set_cell_vertical_alignment(cell, alignment: str = 'center'):
    """设置单元格垂直对齐"""
    val_map = {'top': 'top', 'center': 'center', 'bottom': 'bottom'}
    vAlign = parse_xml(
        f'<w:vAlign {nsdecls("w")} w:val="{val_map.get(alignment, "center")}"/>'
    )
    cell._tc.get_or_add_tcPr().append(vAlign)


def set_cell_width(cell, width_cm: float):
    """设置单元格宽度"""
    tc_w = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{int(width_cm * 567)}" w:type="dxa"/>')
    cell._tc.get_or_add_tcPr().append(tc_w)


# ===== 文档提取 =====
def extract_text_from_docx(filepath: str) -> str:
    """从DOCX文件提取所有文本"""
    doc = DocxDocument(filepath)
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)
    # 也提取表格中的文字
    for table in doc.tables:
        for row in table.rows:
            row_texts = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_texts.append(cell_text)
            if row_texts:
                paragraphs.append(' | '.join(row_texts))
    return '\n'.join(paragraphs)


def extract_text_from_pdf(filepath: str) -> str:
    """从PDF文件提取所有文本"""
    doc = fitz.open(filepath)
    paragraphs = []
    for page in doc:
        text = page.get_text().strip()
        if text:
            paragraphs.append(text)
    doc.close()
    return '\n'.join(paragraphs)


def extract_text(filepath: str) -> str:
    """自动识别文件类型并提取文本"""
    ext = Path(filepath).suffix.lower()
    if ext in ('.doc', '.docx'):
        if not HAS_DOCX:
            raise ImportError("需要安装python-docx: pip install python-docx")
        return extract_text_from_docx(filepath)
    elif ext == '.pdf':
        if not HAS_PDF:
            raise ImportError("需要安装PyMuPDF: pip install pymupdf")
        return extract_text_from_pdf(filepath)
    elif ext in ('.txt', '.md'):
        with open(filepath, 'r', encoding='utf-8') as f:
            return f.read()
    else:
        raise ValueError(f"不支持的文件格式: {ext}")


# ===== 标签-内容解析 =====
def parse_label_content_pairs(text: str, config: dict) -> List[Tuple[str, str]]:
    """
    从文本中识别"标签：内容"对
    
    支持格式：
      - 项目名称：XX工程
      - 项目名称: XX工程
      - 一、施工方案 → 施工方案
      - 带编号的：1. 项目名称：XX工程
    """
    pattern = config['extraction']['pattern']
    min_len = config['extraction']['min_label_length']
    max_len = config['extraction']['max_label_length']
    skip_prefixes = tuple(config['extraction']['skip_lines_startswith'])
    skip_empty = config['extraction']['skip_empty_labels']
    
    pairs = []
    lines = text.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        # 跳过注释/说明行
        if line.startswith(skip_prefixes):
            continue
        
        # 尝试匹配"标签：内容"
        m = re.match(pattern, line)
        if m:
            label = m.group(1).strip()
            content = m.group(2).strip()
            # 清理标签中的编号前缀（如"一、""1.""（1）"）
            label = re.sub(r'^[一二三四五六七八九十]+[、.．]', '', label)
            label = re.sub(r'^[\d]+[、.．)）]', '', label)
            label = re.sub(r'^（[\d一二三四五六七八九十]+）', '', label)
            label = label.strip()
            
            if label and len(label) >= min_len and len(label) <= max_len:
                if skip_empty and not content:
                    continue
                pairs.append((label, content))
        else:
            # 不是"标签：内容"格式，可能是章节标题或普通段落
            # 尝试作为纯文本段落处理
            if len(line) > 3 and not line.startswith(skip_prefixes):
                # 检查是否像章节标题（"一、XX""1.1 XX"）
                if re.match(r'^[一二三四五六七八九十]+[、.．]', line) or \
                   re.match(r'^[\d]+\.[\d]+\s', line) or \
                   re.match(r'^第[一二三四五六七八九十]+章', line) or \
                   re.match(r'^第[一二三四五六七八九十]+节', line) or \
                   re.match(r'^[\d]+、', line):
                    pairs.append(('__section__', line))
    
    return pairs


def group_by_section(pairs: List[Tuple[str, str]]) -> List[List[Tuple[str, str]]]:
    """
    按章节分组，章节标题后的内容归入该章节
    
    返回：[[('__section__', '一、施工方案'), ('工期', '120天'), ...], ...]
    """
    groups = []
    current_group = []
    
    for label, content in pairs:
        if label == '__section__':
            if current_group:
                groups.append(current_group)
            current_group = [(label, content)]
        else:
            current_group.append((label, content))
    
    if current_group:
        groups.append(current_group)
    
    return groups


def merge_same_left(pairs: List[Tuple[str, str]]) -> List[Tuple[str, str, int]]:
    """
    合并左列相同标签的行，返回 (label, content, rowspan)
    rowspan=1表示不合并，>1表示合并
    """
    merged = []
    if not pairs:
        return merged
    
    # 去掉章节标题行，单独处理
    content_pairs = [(l, c) for l, c in pairs if l != '__section__']
    
    i = 0
    while i < len(content_pairs):
        current_label = content_pairs[i][0]
        # 看后面有多少行标签相同
        same_count = 1
        contents = [content_pairs[i][1]]
        j = i + 1
        while j < len(content_pairs) and content_pairs[j][0] == current_label:
            same_count += 1
            contents.append(content_pairs[j][1])
            j += 1
        
        if same_count > 1:
            merged.append((current_label, '\n'.join(contents), same_count))
        else:
            merged.append((current_label, contents[0], 1))
        i = j
    
    return merged


# ===== Word表格生成 =====
def create_two_column_table(doc, pairs: List[Tuple[str, str]], config: dict,
                            section_title: str = None):
    """
    创建一个两列表格
    
    Args:
        doc: Document对象
        pairs: (标签, 内容) 列表，可能包含章节标题
        config: 完整配置
        section_title: 可选的章节标题
    """
    table_cfg = config['table']
    left_cfg = config['left_cell']
    right_cfg = config['right_cell']
    header_cfg = config['header']
    section_cfg = config['section']
    
    # 计算行数
    content_pairs = [(l, c) for l, c in pairs if l != '__section__']
    total_rows = len(content_pairs)
    if header_cfg['enabled']:
        total_rows += 1
    
    if total_rows == 0:
        return  # 没有数据，跳过
    
    # 创建表格（2列）
    table = doc.add_table(rows=total_rows, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    # 设置列宽
    left_width = Cm(table_cfg['left_width_cm'])
    right_width = Cm(table_cfg['right_width_cm'])
    for row in table.rows:
        set_cell_width(row.cells[0], table_cfg['left_width_cm'])
        set_cell_width(row.cells[1], table_cfg['right_width_cm'])
    
    # 配置边框
    border_color = table_cfg['border_color']
    border_size = table_cfg['border_size'] * 4  # 转换到Word的1/8 pt单位
    
    # 填充表头行
    row_idx = 0
    if header_cfg['enabled']:
        header_row = table.rows[0]
        # 合并两列作为表头
        header_row.cells[0].merge(header_row.cells[1])
        cell = header_row.cells[0]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header_cfg['text'])
        run.font.name = header_cfg['font_name']
        run.font.size = Pt(header_cfg['font_size'])
        run.font.bold = header_cfg['bold']
        run.font.color.rgb = hex_to_rgb(header_cfg['color'])
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, header_cfg['bg_color'])
        set_cell_border(cell, border_color, border_size)
        set_cell_vertical_alignment(cell, 'center')
        row_idx = 1
    
    # 填充内容行
    for i, (label, content) in enumerate(content_pairs):
        if row_idx >= len(table.rows):
            break
        row = table.rows[row_idx]
        
        # 左列 - 标签
        left_cell = row.cells[0]
        left_cell.text = ''
        p = left_cell.paragraphs[0]
        run = p.add_run(label)
        run.font.name = left_cfg['font_name']
        run.font.size = Pt(left_cfg['font_size'])
        run.font.bold = left_cfg['bold']
        run.font.color.rgb = hex_to_rgb(left_cfg['color'])
        p.alignment = getattr(WD_ALIGN_PARAGRAPH, left_cfg['alignment'].upper())
        set_cell_shading(left_cell, left_cfg['bg_color'])
        set_cell_border(left_cell, border_color, border_size)
        set_cell_vertical_alignment(left_cell, left_cfg['vertical_alignment'])
        
        # 右列 - 内容
        right_cell = row.cells[1]
        right_cell.text = ''
        p = right_cell.paragraphs[0]
        run = p.add_run(content)
        run.font.name = right_cfg['font_name']
        run.font.size = Pt(right_cfg['font_size'])
        run.font.bold = right_cfg['bold']
        run.font.color.rgb = hex_to_rgb(right_cfg['color'])
        p.alignment = getattr(WD_ALIGN_PARAGRAPH, right_cfg['alignment'].upper())
        set_cell_shading(right_cell, right_cfg['bg_color'])
        set_cell_border(right_cell, border_color, border_size)
        set_cell_vertical_alignment(right_cell, right_cfg['vertical_alignment'])
        
        # 如果内容有多行，自动添加段落
        if '\n' in content:
            lines = content.split('\n')
            # 第一个段落已经有了
            for line_idx, line in enumerate(lines[1:], 1):
                if line_idx >= len(p.paragraphs):
                    new_p = right_cell.add_paragraph()
                else:
                    new_p = p.paragraphs[line_idx]
                new_run = new_p.add_run(line)
                new_run.font.name = right_cfg['font_name']
                new_run.font.size = Pt(right_cfg['font_size'])
                new_run.font.bold = right_cfg['bold']
                new_run.font.color.rgb = hex_to_rgb(right_cfg['color'])
                new_p.alignment = getattr(WD_ALIGN_PARAGRAPH, right_cfg['alignment'].upper())
        
        row_idx += 1


def generate_document(pairs: List[Tuple[str, str]], config: dict,
                      output_path: str, no_merge: bool = False):
    """生成完整Word文档"""
    doc = DocxDocument()
    
    # 页面设置
    section = doc.sections[0]
    page_cfg = config['page']
    section.top_margin = Cm(page_cfg['margin_top_cm'])
    section.bottom_margin = Cm(page_cfg['margin_bottom_cm'])
    section.left_margin = Cm(page_cfg['margin_left_cm'])
    section.right_margin = Cm(page_cfg['margin_right_cm'])
    if page_cfg['orientation'] == 'landscape':
        section.orientation = 1  # WD_ORIENT.LANDSCAPE
    
    # 按章节分组
    groups = group_by_section(pairs)
    
    if not groups:
        # 没有章节标题，所有内容作为一个表格
        if no_merge:
            create_two_column_table(doc, pairs, config)
        else:
            merged = merge_same_left(pairs)
            # 转为 (label, content) 格式
            flat_pairs = [(l, c) for l, c, _ in merged]
            create_two_column_table(doc, flat_pairs, config)
    else:
        for group in groups:
            # 检查是否有章节标题
            section_title = None
            section_pairs = []
            for label, content in group:
                if label == '__section__':
                    section_title = content
                else:
                    section_pairs.append((label, content))
            
            if section_title:
                # 添加章节标题段落
                p = doc.add_paragraph()
                run = p.add_run(section_title)
                run.font.name = config['section']['section_font_name']
                run.font.size = Pt(config['section']['section_font_size'])
                run.font.bold = config['section']['section_bold']
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(6)
            
            if section_pairs:
                if no_merge:
                    create_two_column_table(doc, section_pairs, config, section_title)
                else:
                    merged = merge_same_left(section_pairs)
                    flat_pairs = [(l, c) for l, c, _ in merged]
                    create_two_column_table(doc, flat_pairs, config, section_title)
                
                # 表格后空一行
                doc.add_paragraph()
    
    doc.save(output_path)
    return output_path


# ===== 配置管理 =====
def load_config(config_path: str = None) -> dict:
    """加载配置，未指定则使用默认配置"""
    config = DEFAULT_CONFIG.copy()
    if config_path:
        with open(config_path, 'r', encoding='utf-8') as f:
            user_config = yaml.safe_load(f)
            if user_config:
                # 深层合并
                _deep_merge(config, user_config)
    return config


def _deep_merge(base: dict, override: dict):
    """深度合并字典"""
    for key, value in override.items():
        if key in base and isinstance(base[key], dict) and isinstance(value, dict):
            _deep_merge(base[key], value)
        else:
            base[key] = value


def generate_default_config(output_path: str):
    """生成默认配置文件"""
    with open(output_path, 'w', encoding='utf-8') as f:
        yaml.dump(DEFAULT_CONFIG, f, allow_unicode=True, default_flow_style=False)
    print(f'✅ 默认配置文件已生成: {output_path}')


# ===== 主入口 =====
def main():
    parser = argparse.ArgumentParser(
        description='工程标两列表格生成器 - 从DOC/PDF提取文字，自动生成两列表格Word文档',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例：
  python bid_table_generator.py 技术方案.docx -o 输出.docx
  python bid_table_generator.py 技术方案.pdf -o 输出.docx
  python bid_table_generator.py 技术方案.docx -o 输出.docx --left-color D9E2F3
  python bid_table_generator.py 技术方案.docx -o 输出.docx --font-size 12 --font-name 宋体
  python bid_table_generator.py 技术方案.docx --list             # 仅列出识别结果
  python bid_table_generator.py --gen-config my_config.yaml       # 生成配置模板
        """
    )
    parser.add_argument('input', nargs='?', help='输入文件（DOC/DOCX/PDF）')
    parser.add_argument('-o', '--output', default=None, help='输出Word文件路径')
    parser.add_argument('--config', default=None, help='YAML配置文件路径')
    parser.add_argument('--gen-config', default=None, help='生成默认配置文件到指定路径')
    
    # 快速样式覆盖
    parser.add_argument('--left-width', type=float, default=None, help='左列宽度（cm）')
    parser.add_argument('--right-width', type=float, default=None, help='右列宽度（cm）')
    parser.add_argument('--left-color', default=None, help='左列背景色（十六进制，如D9E2F3）')
    parser.add_argument('--right-color', default=None, help='右列背景色（十六进制，如FFFFFF）')
    parser.add_argument('--border-color', default=None, help='边框颜色（十六进制，如000000）')
    parser.add_argument('--font-size', type=int, default=None, help='正文字号（pt）')
    parser.add_argument('--font-name', default=None, help='正文字体名称')
    parser.add_argument('--bold-left', type=bool, default=None, help='左列是否加粗')
    
    parser.add_argument('--no-merge', action='store_true', help='不分组合并相同标签')
    parser.add_argument('--list', action='store_true', help='仅列出识别的标签-内容对，不生成表格')
    parser.add_argument('--section', action='store_true', default=True,
                        help='按章节分组（默认开启）')
    
    args = parser.parse_args()
    
    # 生成配置模板
    if args.gen_config:
        generate_default_config(args.gen_config)
        return
    
    if not args.input:
        parser.print_help()
        return
    
    # 检查输入文件
    input_path = Path(args.input)
    if not input_path.exists():
        print(f'❌ 文件不存在: {args.input}')
        sys.exit(1)
    
    # 加载配置
    config = load_config(args.config)
    
    # 快速样式覆盖
    if args.left_width is not None:
        config['table']['left_width_cm'] = args.left_width
    if args.right_width is not None:
        config['table']['right_width_cm'] = args.right_width
    if args.left_color is not None:
        config['left_cell']['bg_color'] = args.left_color.lstrip('#')
    if args.right_color is not None:
        config['right_cell']['bg_color'] = args.right_color.lstrip('#')
    if args.border_color is not None:
        config['table']['border_color'] = args.border_color.lstrip('#')
    if args.font_size is not None:
        config['left_cell']['font_size'] = args.font_size
        config['right_cell']['font_size'] = args.font_size
    if args.font_name is not None:
        config['left_cell']['font_name'] = args.font_name
        config['right_cell']['font_name'] = args.font_name
    if args.bold_left is not None:
        config['left_cell']['bold'] = args.bold_left
    
    # 步骤1：提取文字
    print(f'📖 读取文件: {input_path}')
    try:
        text = extract_text(str(input_path))
    except Exception as e:
        print(f'❌ 提取文字失败: {e}')
        sys.exit(1)
    
    print(f'   ✅ 提取文字 {len(text)} 字符')
    
    # 步骤2：解析标签-内容对
    pairs = parse_label_content_pairs(text, config)
    print(f'   📋 识别到 {len(pairs)} 个标签-内容对')
    
    # 仅列出模式
    if args.list:
        print('\n' + '='*60)
        print('识别的标签-内容对：')
        print('='*60)
        section_count = 0
        for label, content in pairs:
            if label == '__section__':
                print(f'\n  📂 {content}')
                section_count += 1
            else:
                # 截断过长的内容
                display_content = content[:60] + '...' if len(content) > 60 else content
                print(f'  📌 {label} → {display_content}')
        print(f'\n共 {section_count} 个章节，{len(pairs) - section_count} 个标签-内容对')
        return
    
    # 步骤3：生成Word文档
    output_path = args.output or str(input_path.with_name(f'{input_path.stem}_表格化.docx'))
    print(f'📝 生成表格文档: {output_path}')
    
    try:
        result = generate_document(pairs, config, output_path, args.no_merge)
        print(f'✅ 完成！输出文件: {result}')
        print(f'   📊 共 {len(pairs)} 个标签-内容对')
    except Exception as e:
        print(f'❌ 生成失败: {e}')
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == '__main__':
    main()