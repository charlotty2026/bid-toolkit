#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
投标文件资质响应检查器 v1.0
==========================
检查投标文件中的资质响应是否完整：
  1. 从招标文件解析结果（parse_bid.py 输出的JSON）中提取资质要求
  2. 从投标文件中扫描资质响应情况
  3. 对比：招标要求的每项资质是否在投标文件中有对应响应
  4. 检查资质证书有效期是否在项目期内有效
  5. 输出检查报告

日期: 2026-07-27

用法：
    python scripts/qualification_check.py check 投标文件.md --rfp 招标文件.json
    python scripts/qualification_check.py check 投标文件.md --rfp 招标文件.json --json
"""

import argparse
import json
import re
import sys
from pathlib import Path
from datetime import date, datetime

# ============================================================
# 配置
# ============================================================

BASE_DIR = Path(__file__).resolve().parent.parent
QUALIFICATIONS_FILE = BASE_DIR / "company_profile" / "qualifications.md"
CONFIG_FILE = BASE_DIR / "user_config.yaml"

# 资质证书名称 → 搜索别名映射（用于在投标文件中匹配资质响应）
QUAL_ALIAS: dict[str, list[str]] = {
    "营业执照": ["营业执照", "统一社会信用代码", "工商注册"],
    "ISO9001": ["ISO9001", "ISO 9001", "质量管理体系认证", "质量管理体系"],
    "ISO14001": ["ISO14001", "ISO 14001", "环境管理体系认证", "环境管理体系"],
    "ISO45001": ["ISO45001", "ISO 45001", "职业健康安全管理体系", "OHSAS18001"],
    "安全生产许可证": ["安全生产许可证", "安全许可证"],
    "劳务派遣经营许可证": ["劳务派遣经营许可证", "劳务派遣证", "劳务派遣"],
    "人力资源服务许可证": ["人力资源服务许可证", "人力资源证"],
    "建筑业企业资质": ["建筑业企业资质", "施工资质", "建筑资质"],
    "建造师": ["建造师", "注册建造师"],
    "安全生产考核合格证": ["安全生产考核合格证", "安全员证", "A证", "B证", "C证"],
    "食品经营许可证": ["食品经营许可证", "食品流通许可证"],
    "医疗器械经营许可证": ["医疗器械经营许可证", "医疗器械经营"],
    "软件企业认定": ["软件企业认定", "软件企业证书"],
    "高新技术企业": ["高新技术企业", "高新证书"],
    "CMMI": ["CMMI", "能力成熟度模型"],
    "ITSS": ["ITSS", "信息技术服务标准"],
    "3C认证": ["3C认证", "CCC认证", "强制性产品认证"],
    "环保产品认证": ["环保产品认证", "环境标志认证", "十环认证"],
}

# 项目周期内的"即将到期"阈值（天）
EXPIRY_WARNING_DAYS = 90


# ============================================================
# 文件读取（支持MD/TXT和DOCX）
# ============================================================

def read_file(file_path: str) -> list[str]:
    """读取输入文件，返回行列表。支持 .md/.txt/.docx 格式。"""
    path = Path(file_path)
    if not path.exists():
        print(f"错误：文件不存在：{file_path}", file=sys.stderr)
        sys.exit(1)

    suffix = path.suffix.lower()
    if suffix in (".md", ".txt"):
        with open(path, "r", encoding="utf-8") as f:
            return f.read().split("\n")
    elif suffix == ".docx":
        return _read_docx(path)
    else:
        print(f"错误：不支持的文件格式：{suffix}", file=sys.stderr)
        print("支持：.md / .txt / .docx", file=sys.stderr)
        sys.exit(1)


def _read_docx(path: Path) -> list[str]:
    """读取DOCX文件，返回行列表（含表格内容）。"""
    try:
        from docx import Document
    except ImportError:
        print("错误：需要 python-docx：pip install python-docx", file=sys.stderr)
        sys.exit(1)

    doc = Document(str(path))
    lines: list[str] = []
    for para in doc.paragraphs:
        lines.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text for cell in row.cells)
            lines.append(row_text)
    return lines


# ============================================================
# 招标文件JSON解析
# ============================================================

def load_rfp(rfp_path: str) -> dict:
    """加载招标文件解析结果JSON。"""
    path = Path(rfp_path)
    if not path.exists():
        print(f"错误：招标文件不存在：{rfp_path}", file=sys.stderr)
        sys.exit(1)

    with open(path, "r", encoding="utf-8") as f:
        data = json.load(f)
    return data


def extract_qual_requirements(rfp_data: dict) -> list[dict]:
    """从招标文件JSON中提取资质要求列表。

    parse_bid.py 输出的JSON中，"资质要求" 是一个字符串列表。
    本函数将其标准化为结构化的资质要求条目。
    """
    raw_quals = rfp_data.get("资质要求", [])
    if not raw_quals:
        # 也尝试其他可能的key
        raw_quals = rfp_data.get("qualifications", [])

    requirements: list[dict] = []
    for item in raw_quals:
        if isinstance(item, str):
            # 从原始字符串中提取资质名称关键词
            qual_name = _extract_qual_name(item)
            requirements.append({
                "requirement": item,
                "qual_name": qual_name,
                "qual_type": _classify_qual(qual_name),
            })
        elif isinstance(item, dict):
            qual_name = item.get("name", item.get("名称", ""))
            requirements.append({
                "requirement": item.get("requirement", item.get("要求", qual_name)),
                "qual_name": qual_name,
                "qual_type": _classify_qual(qual_name),
            })
    return requirements


def _extract_qual_name(text: str) -> str:
    """从资质要求文本中提取资质名称关键词。"""
    # 优先匹配已知的资质证书名称
    for key, aliases in QUAL_ALIAS.items():
        for alias in aliases:
            if alias in text:
                return key
    # 提取引号内的名称
    quote_match = re.search(r"[「""'《【](.+?)[」""'》】]", text)
    if quote_match:
        return quote_match.group(1)
    # 截取前20个字符作为名称
    return text[:20].strip()


def _classify_qual(qual_name: str) -> str:
    """对资质名称进行分类。"""
    if any(kw in qual_name for kw in ["营业执照", "工商", "信用代码"]):
        return "工商注册"
    if any(kw in qual_name for kw in ["ISO", "管理体系", "认证"]):
        return "体系认证"
    if any(kw in qual_name for kw in ["许可证", "许可"]):
        return "经营许可"
    if any(kw in qual_name for kw in ["资质", "建造师", "资格"]):
        return "行业资质"
    return "其他"


# ============================================================
# 投标文件资质响应扫描
# ============================================================

def scan_bid_qualifications(lines: list[str], requirements: list[dict]) -> list[dict]:
    """扫描投标文件，检查每项资质要求是否有响应。

    返回检查结果列表，每项包含：
      - requirement: 资质要求文本
      - qual_name: 资质名称
      - qual_type: 资质类型
      - responded: 是否已响应
      - line: 响应所在行号（未找到为0）
      - expiry_date: 证书有效期（如果能提取到）
      - expiry_status: 有效期状态（valid/expired/expiring_soon/unknown）
    """
    results: list[dict] = []

    for req in requirements:
        qual_name = req["qual_name"]
        aliases = QUAL_ALIAS.get(qual_name, [qual_name])

        # 在投标文件中搜索资质响应
        found_line = 0
        found_context = ""
        for i, line in enumerate(lines, 1):
            for alias in aliases:
                if alias in line:
                    found_line = i
                    found_context = line.strip()[:120]
                    break
            if found_line:
                break

        # 尝试提取有效期
        expiry_date: str | None = None
        expiry_status: str = "unknown"
        if found_line:
            # 在找到行附近搜索有效期
            search_start = max(0, found_line - 1)
            search_end = min(len(lines), found_line + 3)
            for j in range(search_start, search_end):
                expiry_date = _extract_expiry_date(lines[j])
                if expiry_date:
                    expiry_status = _check_expiry_status(expiry_date)
                    break

        results.append({
            "requirement": req["requirement"],
            "qual_name": qual_name,
            "qual_type": req["qual_type"],
            "responded": found_line > 0,
            "line": found_line,
            "context": found_context,
            "expiry_date": expiry_date,
            "expiry_status": expiry_status,
        })

    return results


def _extract_expiry_date(text: str) -> str | None:
    """从文本中提取有效期日期。"""
    patterns = [
        r"有效期[至到]?[：:]?\s*(\d{4}[-/年]\d{1,2}[-/月]\d{1,2}日?)",
        r"有效期[至到]?[：:]?\s*(\d{4}年\d{1,2}月\d{1,2}日)",
        r"至\s*(\d{4}[-/年]\d{1,2}[-/月]\d{1,2}日?)",
        r"有效期至\s*(\d{4}[-/]\d{1,2}[-/]\d{1,2})",
    ]
    for pat in patterns:
        m = re.search(pat, text)
        if m:
            raw = m.group(1)
            # 标准化为 YYYY-MM-DD
            normalized = _normalize_date(raw)
            if normalized:
                return normalized
    return None


def _normalize_date(raw: str) -> str | None:
    """将各种日期格式标准化为 YYYY-MM-DD。"""
    # 2024年12月31日 → 2024-12-31
    m = re.match(r"(\d{4})年(\d{1,2})月(\d{1,2})日?", raw)
    if m:
        return f"{int(m.group(1)):04d}-{int(m.group(2)):02d}-{int(m.group(3)):02d}"
    # 2024-12-31 or 2024/12/31
    m = re.match(r"(\d{4})[-/](\d{1,2})[-/](\d{1,2})", raw)
    if m:
        return f"{int(m.group(1)):04d}-{int(m.group(2)):02d}-{int(m.group(3)):02d}"
    return None


def _check_expiry_status(expiry_date_str: str) -> str:
    """检查有效期状态。"""
    try:
        parts = expiry_date_str.split("-")
        expiry_date = date(int(parts[0]), int(parts[1]), int(parts[2]))
    except (ValueError, IndexError):
        return "unknown"

    today = date.today()
    days_until_expiry = (expiry_date - today).days

    if expiry_date < today:
        return "expired"
    if days_until_expiry <= EXPIRY_WARNING_DAYS:
        return "expiring_soon"
    return "valid"


# ============================================================
# 企业资质证书加载（从 qualifications.md）
# ============================================================

def load_company_certificates() -> list[dict]:
    """从 qualifications.md 加载企业持有的资质证书信息。

    文件格式：Markdown表格，列为：证书名称 | 证书编号 | 发证机构 | 有效期至
    """
    if not QUALIFICATIONS_FILE.exists():
        return []

    certs: list[dict] = []
    with open(QUALIFICATIONS_FILE, "r", encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if not line or line.startswith("#") or line.startswith("<!--") or line.startswith("-->"):
                continue
            if line.startswith("|") and not line.startswith("|---") and not line.startswith("| ---"):
                parts = [p.strip() for p in line.split("|")]
                # 去掉首尾空元素
                parts = [p for p in parts if p]
                if len(parts) >= 4 and parts[0] != "证书名称":
                    cert_name = parts[0].replace("***", "").strip()
                    cert_no = parts[1].replace("***", "").strip()
                    issuer = parts[2].replace("***", "").strip()
                    expiry = parts[3].replace("***", "").strip()
                    if cert_name and cert_name != "证书名称":
                        certs.append({
                            "name": cert_name,
                            "number": cert_no,
                            "issuer": issuer,
                            "expiry": expiry if expiry else "长期",
                        })
    return certs


# ============================================================
# 配置加载
# ============================================================

def load_config() -> dict:
    """加载配置文件（user_config.yaml 或 config.yaml）。"""
    # 优先读取 user_config.yaml
    for config_path in [CONFIG_FILE, BASE_DIR / "config.yaml"]:
        if config_path.exists():
            try:
                import yaml
                with open(config_path, "r", encoding="utf-8") as f:
                    return yaml.safe_load(f) or {}
            except ImportError:
                # yaml未安装，手动解析简单key
                return _parse_simple_yaml(config_path)
            except Exception:
                pass
    return {}


def _parse_simple_yaml(path: Path) -> dict:
    """简易YAML解析（不依赖pyyaml，仅支持简单嵌套）。"""
    config: dict = {}
    current_section: dict | None = None
    with open(path, "r", encoding="utf-8") as f:
        for line in f:
            stripped = line.rstrip()
            if not stripped or stripped.startswith("#"):
                continue
            if not stripped.startswith(" ") and stripped.endswith(":"):
                key = stripped[:-1]
                current_section = {}
                config[key] = current_section
            elif current_section is not None and ":" in stripped:
                parts = stripped.strip().split(":", 1)
                key = parts[0].strip()
                val = parts[1].strip() if len(parts) > 1 else ""
                current_section[key] = val
    return config


# ============================================================
# 报告生成
# ============================================================

def print_text_report(
    bid_file: str,
    rfp_file: str,
    results: list[dict],
    company_certs: list[dict],
) -> None:
    """打印文本格式报告。"""
    total = len(results)
    responded = sum(1 for r in results if r["responded"])
    missing = sum(1 for r in results if not r["responded"])
    expiring = sum(1 for r in results if r["expiry_status"] == "expiring_soon")
    expired = sum(1 for r in results if r["expiry_status"] == "expired")

    print("资质响应检查报告")
    print("=" * 60)
    print(f"检查文件: {bid_file}")
    print(f"招标文件: {rfp_file}")
    print()

    # 企业资质证书概览
    if company_certs:
        print("企业持有证书:")
        for cert in company_certs:
            expiry_display = cert["expiry"]
            expiry_status = _check_expiry_status(cert["expiry"]) if cert["expiry"] != "长期" else "valid"
            status_icon = "✓" if expiry_status == "valid" else "⚠" if expiry_status == "expiring_soon" else "✗"
            print(f"  {status_icon} {cert['name']} — 有效期: {expiry_display}")
        print()

    print("资质要求清单:")
    for r in results:
        if not r["responded"]:
            print(f"[✗] {r['qual_name']} — 未找到响应")
        elif r["expiry_status"] == "expired":
            print(f"[✗] {r['qual_name']} — 已响应（第{r['line']}页/行），但有效期{r['expiry_date']}已过期")
        elif r["expiry_status"] == "expiring_soon":
            print(f"[⚠] {r['qual_name']} — 已响应（第{r['line']}页/行），但有效期{r['expiry_date']}即将到期")
        elif r["expiry_date"]:
            print(f"[✓] {r['qual_name']} — 已响应（第{r['line']}页/行），有效期: {r['expiry_date']} ✓")
        else:
            print(f"[✓] {r['qual_name']} — 已响应（第{r['line']}页/行）")

    print()
    summary_parts = [f"{total}项资质要求", f"{responded}项已响应", f"{missing}项缺失"]
    if expired > 0:
        summary_parts.append(f"{expired}项已过期")
    if expiring > 0:
        summary_parts.append(f"{expiring}项即将到期")
    print(f"汇总: {', '.join(summary_parts)}")


def print_json_report(
    bid_file: str,
    rfp_file: str,
    results: list[dict],
    company_certs: list[dict],
) -> None:
    """打印JSON格式报告。"""
    total = len(results)
    responded = sum(1 for r in results if r["responded"])
    missing = sum(1 for r in results if not r["responded"])
    expiring = sum(1 for r in results if r["expiry_status"] == "expiring_soon")
    expired = sum(1 for r in results if r["expiry_status"] == "expired")

    report = {
        "检查文件": bid_file,
        "招标文件": rfp_file,
        "资质要求总数": total,
        "已响应": responded,
        "缺失": missing,
        "已过期": expired,
        "即将到期": expiring,
        "企业证书": company_certs,
        "资质检查明细": results,
    }
    print(json.dumps(report, ensure_ascii=False, indent=2))


# ============================================================
# CLI入口
# ============================================================

def main() -> None:
    parser = argparse.ArgumentParser(
        description="投标文件资质响应检查器 v1.0",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
    python scripts/qualification_check.py check 投标文件.md --rfp 招标文件.json
    python scripts/qualification_check.py check 投标文件.md --rfp 招标文件.json --json
        """,
    )
    subparsers = parser.add_subparsers(dest="command", help="子命令")

    check_parser = subparsers.add_parser("check", help="检查资质响应")
    check_parser.add_argument("file", help="投标文件路径（.md/.txt/.docx）")
    check_parser.add_argument("--rfp", required=True, help="招标文件解析结果JSON（parse_bid.py输出）")
    check_parser.add_argument("--json", action="store_true", help="输出JSON格式报告")

    args = parser.parse_args()

    if args.command == "check":
        # 1. 读取投标文件
        lines = read_file(args.file)

        # 2. 加载招标文件解析结果
        rfp_data = load_rfp(args.rfp)
        requirements = extract_qual_requirements(rfp_data)

        # 3. 如果未从招标文件提取到资质要求，使用配置文件中的默认要求
        if not requirements:
            config = load_config()
            qual_config = config.get("qualification_check", {})
            default_quals = qual_config.get("required_quals", [])
            for q in default_quals:
                requirements.append({
                    "requirement": q,
                    "qual_name": _extract_qual_name(q),
                    "qual_type": _classify_qual(_extract_qual_name(q)),
                })

        # 4. 扫描投标文件中的资质响应
        results = scan_bid_qualifications(lines, requirements)

        # 5. 加载企业资质证书信息
        company_certs = load_company_certificates()

        # 6. 输出报告
        if args.json:
            print_json_report(args.file, args.rfp, results, company_certs)
        else:
            print_text_report(args.file, args.rfp, results, company_certs)
    else:
        parser.print_help()


if __name__ == "__main__":
    main()
