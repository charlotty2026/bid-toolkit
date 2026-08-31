#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
render_docx.py — 擎标排版引擎

输入 content.json（结构化内容），输出格式化的 .docx 文件。
所有排版参数由 format_config.yaml 控制，改配置不改代码。

用法:
    python -m bid_toolkit.render_docx content.json output.docx
    python -m bid_toolkit.render_docx content.json output.docx --config custom.yaml

块类型:
    h1-h5      标题（真实 Heading 1-5 样式 + 自动编号，默认 第一章/一、/（一）/1、/1.1）
    p          正文段落
    list       列表（最多三级）
    table      表格（分章自动编号）
    figure     图片（分章自动编号）

架构参考: tender-bid-writer 的 build_docx.py（793行排版引擎）
"""

import argparse
import json
import os
import re
import sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.shared import Cm, Emu, Pt, RGBColor

# ---------------------------------------------------------------------------
# 全局配置
# ---------------------------------------------------------------------------

_CONFIG = {}  # 加载后全局持有


def _load_config(path):
    """加载 YAML 格式排版配置，提供默认值兜底。"""
    try:
        import yaml
        with open(path, 'r', encoding='utf-8') as f:
            raw = yaml.safe_load(f)
    except ImportError:
        # 用 json 兜底（兼容无 pyyaml 环境）
        with open(path, 'r', encoding='utf-8') as f:
            raw = json.load(f)
    except FileNotFoundError:
        raw = {}

    defaults = {
        "页面": {"纸张宽cm": 21.0, "纸张高cm": 29.7,
                 "上边距cm": 2.5, "下边距cm": 2.0,
                 "左边距cm": 2.0, "右边距cm": 2.0,
                 "页眉距cm": 1.5, "页脚距cm": 1.75},
        "字体": {"中文字体": "宋体", "西文字体": "Times New Roman", "字体颜色": "000000"},
        "正文": {"字号pt": 12.0, "行距倍数": 1.5, "首行缩进字符": 2, "对齐": "两端对齐"},
        "标题": {"字号pt": 16.0, "加粗": True, "对齐": "左对齐",
                 "段前行": 1.0, "段后行": 1.0, "行距倍数": 1.5,
                 "各级字号pt": [22.0, 18.0, 16.0, 15.0, 14.0],
                 "各级对齐": ["居中", "左对齐", "左对齐", "左对齐", "左对齐"],
                 "一级标题另起页": True},
        "标题编号": {"级别数": 5,
                 "格式": [
                     {"numFmt": "chineseCounting", "lvlText": "第%1章"},
                     {"numFmt": "chineseCounting", "lvlText": "%2、"},
                     {"numFmt": "chineseCounting", "lvlText": "（%3）"},
                     {"numFmt": "decimal",        "lvlText": "%4、"},
                     {"numFmt": "decimal",        "lvlText": "%4.%5"},
                 ]},
        "序号标题": {"字号pt": 12.0, "加粗": False, "行距倍数": 1.5, "首行缩进字符": 2, "左缩进字符": 0},
        "列表": {"字号pt": 12.0, "加粗": False, "行距倍数": 1.5, "首行缩进字符": 2, "每级左缩进字符": 2},
        "表格": {"表题前缀": "表", "表题字号pt": 10.5, "表题加粗": False,
                 "表题段前行": 0.5, "内容字号pt": 10.5, "表头加粗": True,
                 "边框粗细": 4, "宽度百分比": 100},
        "配图": {"图题前缀": "图", "图题字号pt": 10.5, "图题加粗": True,
                 "图题段后行": 0.5, "图前段前行": 0.5,
                 "默认宽度cm": 12.0, "最大宽度cm": 17.0,
                 "占位文本": "【此处为占位图，请替换为实际截图】"},
        "目录": {"标题文本": "目 录", "标题字号pt": 16.0, "级别范围": "1-3"},
        "页码": {"字号pt": 10.5, "对齐": "居中", "起始页码": 1},
    }
    for section, keys in defaults.items():
        if section not in raw:
            raw[section] = {}
        for k, v in keys.items():
            raw[section].setdefault(k, v)
    return raw


def cfg(section, key):
    """安全读取配置项。"""
    return _CONFIG.get(section, {}).get(key)


# ---------------------------------------------------------------------------
# 辅助函数
# ---------------------------------------------------------------------------

SZ_XIAOSI = Pt(12)       # 小四
SZ_WUHAO = Pt(10.5)      # 五号
SZ_SANHAO = Pt(16)       # 三号（一号标题）


def _hex_to_rgb(hex_color):
    """6位十六进制颜色 → RGBColor。"""
    hex_color = hex_color.lstrip('#')
    if len(hex_color) != 6:
        return RGBColor(0x00, 0x00, 0x00)
    return RGBColor(int(hex_color[0:2], 16),
                    int(hex_color[2:4], 16),
                    int(hex_color[4:6], 16))


def set_run_fonts(run, font_name=None, font_size=None, bold=None, color=None):
    """设置 run 的字体属性。"""
    cfg_font = _CONFIG.get("字体", {})
    fn = font_name or cfg_font.get("中文字体", "宋体")
    en = cfg_font.get("西文字体", "Times New Roman")
    run.font.name = en
    run.element.rPr.rFonts.set(qn('w:eastAsia'), fn)
    if font_size:
        run.font.size = font_size
    if bold is not None:
        run.font.bold = bold
    if color:
        run.font.color.rgb = _hex_to_rgb(color)
    else:
        fc = cfg_font.get("字体颜色", "000000")
        run.font.color.rgb = _hex_to_rgb(fc)


def set_spacing(paragraph, before_lines=0, after_lines=0, line_spacing=1.5):
    """设置段落间距（行数单位）。"""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before_lines * 12)
    pf.space_after = Pt(after_lines * 12)
    pf.line_spacing = line_spacing


def set_indent(paragraph, chars=2):
    """设置首行缩进（字符数）。"""
    if chars > 0:
        paragraph.paragraph_format.first_line_indent = Pt(chars * 12)


def set_alignment(paragraph, align_str):
    """设置段落对齐。"""
    mapping = {
        "居中": WD_ALIGN_PARAGRAPH.CENTER,
        "两端对齐": WD_ALIGN_PARAGRAPH.JUSTIFY,
        "左对齐": WD_ALIGN_PARAGRAPH.LEFT,
        "右对齐": WD_ALIGN_PARAGRAPH.RIGHT,
    }
    paragraph.alignment = mapping.get(align_str, WD_ALIGN_PARAGRAPH.JUSTIFY)


def add_field(paragraph, field_code, placeholder="请更新域"):
    """在段落中插入 Word 域代码。"""
    run = paragraph.add_run()
    fld_char_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._element.append(fld_char_begin)
    instr = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve">{field_code}</w:instrText>')
    run._element.append(instr)
    fld_char_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run._element.append(fld_char_end)
    # 占位符
    run2 = paragraph.add_run(placeholder)
    run2.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    return run2


def update_fields_on_open(document):
    """设置 Word 文档打开时自动更新域。"""
    settings = document.settings.element
    update = settings.find(qn('w:updateFields'))
    if update is None:
        update = parse_xml(f'<w:updateFields {nsdecls("w")} w:val="true"/>')
        settings.append(update)
    else:
        update.set(qn('w:val'), 'true')


# ---------------------------------------------------------------------------
# 页面设置
# ---------------------------------------------------------------------------

def set_page_layout(section):
    """设置页面布局。"""
    section.page_width = Cm(cfg("页面", "纸张宽cm"))
    section.page_height = Cm(cfg("页面", "纸张高cm"))
    section.top_margin = Cm(cfg("页面", "上边距cm"))
    section.bottom_margin = Cm(cfg("页面", "下边距cm"))
    section.left_margin = Cm(cfg("页面", "左边距cm"))
    section.right_margin = Cm(cfg("页面", "右边距cm"))
    section.header_distance = Cm(cfg("页面", "页眉距cm"))
    section.footer_distance = Cm(cfg("页面", "页脚距cm"))


def clear_footer(section):
    """清空页脚。"""
    footer = section.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        p.clear()


def add_centered_page_number(section):
    """在页脚添加居中页码。"""
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 页码域
    run = p.add_run()
    fld_char_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._element.append(fld_char_begin)
    instr = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run._element.append(instr)
    fld_char_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run._element.append(fld_char_end)
    set_run_fonts(run, font_size=Pt(cfg("页码", "字号pt")))


def restart_page_numbering(section, start=1):
    """重启页码。"""
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn('w:pgNumType'))
    if pg_num_type is None:
        pg_num_type = parse_xml(f'<w:pgNumType {nsdecls("w")} w:start="{start}"/>')
        sect_pr.append(pg_num_type)
    else:
        pg_num_type.set(qn('w:start'), str(start))


# ---------------------------------------------------------------------------
# 样式设置
# ---------------------------------------------------------------------------

def setup_base_styles(document):
    """设置文档基础样式。"""
    style = document.styles['Normal']
    style.font.name = cfg("字体", "西文字体")
    style.element.rPr.rFonts.set(qn('w:eastAsia'), cfg("字体", "中文字体"))
    style.font.size = Pt(cfg("正文", "字号pt"))
    style.paragraph_format.line_spacing = cfg("正文", "行距倍数")


def setup_heading_styles(document):
    """设置标题样式（1-5级真实 Heading 样式）。

    关键：标题一律用 Word 内置 Heading 1-5 样式，客户拿到文档后可在 Word 中
    自由改字体/对齐/颜色，不影响自动编号（编号由 numbering.xml 驱动，与样式解耦）。
    """
    cfg_title = _CONFIG.get("标题", {})
    sizes = cfg_title.get("各级字号pt", [16, 14, 12, 12, 12])
    aligns = cfg_title.get("各级对齐", ["左对齐"] * 5)
    align_map = {
        "居中": WD_ALIGN_PARAGRAPH.CENTER,
        "两端对齐": WD_ALIGN_PARAGRAPH.JUSTIFY,
        "左对齐": WD_ALIGN_PARAGRAPH.LEFT,
        "右对齐": WD_ALIGN_PARAGRAPH.RIGHT,
    }
    for level in range(1, 6):
        style_name = f'Heading {level}'
        try:
            style = document.styles[style_name]
            style.font.name = cfg("字体", "西文字体")
            style.element.rPr.rFonts.set(qn('w:eastAsia'), cfg("字体", "中文字体"))
            style.font.size = Pt(sizes[level - 1] if level - 1 < len(sizes) else 12)
            style.font.bold = cfg_title.get("加粗", True)
            style.paragraph_format.line_spacing = cfg_title.get("行距倍数", 1.5)
            style.paragraph_format.alignment = align_map.get(
                aligns[level - 1] if level - 1 < len(aligns) else "左对齐",
                WD_ALIGN_PARAGRAPH.LEFT)
        except KeyError:
            continue


def setup_heading_numbering(document):
    """创建五级标题自动编号（配置驱动，默认 第一章/一、/（一）/1、/1.1）。

    设计原则（与产品铁律一致）：
    - 标题一律使用真实 Heading 1-5 样式，客户可在 Word 中自由改字体/对齐/编号皮肤；
    - 编号由 Word numbering.xml 自动生成（numPr），**绝不手打编号文本**；
    - 具体编号格式（numFmt + lvlText）由 format_config 的「标题编号.格式」控制，
      整包替换即可切换编号风格（如改为 1/1.1/1.1.1），无需改代码。
    """
    numbering = document.part.numbering_part.numbering_definitions._numbering
    cfg_num = _CONFIG.get("标题编号", {})
    levels = cfg_num.get("格式", [
        {"numFmt": "chineseCounting", "lvlText": "第%1章"},
        {"numFmt": "chineseCounting", "lvlText": "%2、"},
        {"numFmt": "chineseCounting", "lvlText": "（%3）"},
        {"numFmt": "decimal",        "lvlText": "%4、"},
        {"numFmt": "decimal",        "lvlText": "%4.%5"},
    ])

    # 创建抽象编号
    abstract_num_id = _next_id(numbering, 'abstractNum')
    lvl_xml = []
    for i, spec in enumerate(levels):
        num_fmt = spec.get("numFmt", "decimal")
        lvl_text = spec.get("lvlText", f"%{i + 1}")
        # 标题按"靠左"排版：编号与正文均贴左边距，不逐级缩进。
        left = 0
        hanging = 0
        lvl_xml.append(
            f'<w:lvl w:ilvl="{i}"><w:start w:val="1"/>'
            f'<w:numFmt w:val="{num_fmt}"/>'
            f'<w:lvlText w:val="{lvl_text}"/>'
            f'<w:lvlJc w:val="left"/>'
            f'<w:pPr><w:ind w:left="{left}" w:hanging="{hanging}"/></w:pPr></w:lvl>'
        )
    abstract_num = parse_xml(
        f'<w:abstractNum {nsdecls("w")} w:abstractNumId="{abstract_num_id}">'
        f'<w:multiLevelType w:val="multilevel"/>'
        + "".join(lvl_xml) +
        f'</w:abstractNum>')
    numbering.append(abstract_num)

    # 创建编号实例
    num_id = _next_id(numbering, 'num')
    num = parse_xml(f'<w:num {nsdecls("w")} w:numId="{num_id}">'
                    f'<w:abstractNumId w:val="{abstract_num_id}"/></w:num>')
    numbering.append(num)
    return num_id


def _next_id(container, tag):
    """找到下一个可用的ID。"""
    max_id = 0
    for elem in container:
        tag_name = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
        if tag_name == tag:
            i = int(elem.get(qn('w:' + ('abstractNumId' if tag == 'abstractNum' else 'numId')), '0'))
            max_id = max(max_id, i)
    return max_id + 1


# ---------------------------------------------------------------------------
# 渲染函数
# ---------------------------------------------------------------------------

def render_heading(document, level, text, heading_num_id):
    """渲染标题（自动编号）。"""
    cfg_title = _CONFIG.get("标题", {})
    p = document.add_paragraph(style=f'Heading {level}')
    # 应用编号
    pPr = p._element.get_or_add_pPr()
    numPr = parse_xml(f'<w:numPr {nsdecls("w")}><w:ilvl w:val="{level - 1}"/>'
                      f'<w:numId w:val="{heading_num_id}"/></w:numPr>')
    pPr.append(numPr)
    # 清空已有文本，写入新文本
    p.clear()
    run = p.add_run(text)
    sizes = cfg_title.get("各级字号pt", [16, 14, 12, 12, 12])
    size = sizes[level - 1] if level - 1 < len(sizes) else 12
    set_run_fonts(run, font_size=Pt(size), bold=cfg_title.get("加粗", True))
    set_spacing(p, before_lines=cfg_title.get("段前行", 1.0),
                after_lines=cfg_title.get("段后行", 1.0),
                line_spacing=cfg_title.get("行距倍数", 1.5))
    return p


def render_paragraph(document, text):
    """渲染正文段落。"""
    cfg_body = _CONFIG.get("正文", {})
    p = document.add_paragraph()
    run = p.add_run(text)
    set_run_fonts(run, font_size=Pt(cfg_body.get("字号pt", 12)))
    set_spacing(p, line_spacing=cfg_body.get("行距倍数", 1.5))
    set_indent(p, chars=cfg_body.get("首行缩进字符", 2))
    set_alignment(p, cfg_body.get("对齐", "两端对齐"))
    return p


def render_hnum(document, seq, text, level=0):
    """渲染序号标题（如 （1）1）① ）。"""
    cfg_hnum = _CONFIG.get("序号标题", {})
    markers = ["（1）", "1）", "①"]
    marker = markers[min(level, 2)]
    p = document.add_paragraph()
    run = p.add_run(f"{marker} {text}")
    set_run_fonts(run, font_size=Pt(cfg_hnum.get("字号pt", 12)),
                  bold=cfg_hnum.get("加粗", False))
    set_spacing(p, line_spacing=cfg_hnum.get("行距倍数", 1.5))
    set_indent(p, chars=cfg_hnum.get("首行缩进字符", 2))
    if cfg_hnum.get("左缩进字符", 0) > 0:
        p.paragraph_format.left_indent = Pt(cfg_hnum.get("左缩进字符", 0) * 12)
    return p


def render_list(document, items):
    """渲染列表（最多三级）。"""
    cfg_list = _CONFIG.get("列表", {})
    markers = ["（1）", "1）", "①"]
    indent_per_level = cfg_list.get("每级左缩进字符", 2) * 12

    def _render_items(items, level=0):
        for i, item in enumerate(items):
            if isinstance(item, dict):
                text = item.get("t", "")
                children = item.get("children", [])
            else:
                text = item
                children = []
            marker = markers[min(level, 2)]
            seq = i + 1
            marker_text = marker.replace("1", str(seq))
            p = document.add_paragraph()
            run = p.add_run(f"{marker_text} {text}")
            set_run_fonts(run, font_size=Pt(cfg_list.get("字号pt", 12)),
                          bold=cfg_list.get("加粗", False))
            set_spacing(p, line_spacing=cfg_list.get("行距倍数", 1.5))
            set_indent(p, chars=cfg_list.get("首行缩进字符", 2))
            if level > 0:
                p.paragraph_format.left_indent = Pt(level * indent_per_level)
            if children and level < 2:
                _render_items(children, level + 1)

    _render_items(items)


def render_table(document, ch_num, tbl_num, title, header, rows):
    """渲染表格（分章编号，全框线，支持颜色配置）。"""
    cfg_tbl = _CONFIG.get("表格", {})
    cfg_tblc = _CONFIG.get("表格颜色", {})

    # 表题
    prefix = cfg_tbl.get("表题前缀", "表")
    p = document.add_paragraph()
    run = p.add_run(f"{prefix}{ch_num}-{tbl_num} {title}")
    set_run_fonts(run, font_size=Pt(cfg_tbl.get("表题字号pt", 10.5)),
                  bold=cfg_tbl.get("表题加粗", False))
    set_spacing(p, before_lines=cfg_tbl.get("表题段前行", 0.5))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if not header and not rows:
        # 空表占位
        p2 = document.add_paragraph()
        run2 = p2.add_run("【此处为表格】")
        run2.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
        return

    # 创建表格
    num_cols = len(header) if header else (len(rows[0]) if rows else 1)
    num_rows = 1 + len(rows)  # 表头 + 数据行
    table = document.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Table Grid'

    # 设置宽度
    width_pct = cfg_tbl.get("宽度百分比", 100)
    if width_pct < 100:
        page_width = cfg("页面", "纸张宽cm") - cfg("页面", "左边距cm") - cfg("页面", "右边距cm")
        total_width = Cm(page_width * width_pct / 100)
        for row in table.rows:
            for cell in row.cells:
                cell.width = Cm(total_width.cm / num_cols)

    # 颜色配置
    header_bg = cfg_tblc.get("表头背景色")
    header_fg = cfg_tblc.get("表头字体色")
    stripe = cfg_tblc.get("数据行隔行变色", False)
    alt1 = cfg_tblc.get("数据行交替色1", "FFFFFF")
    alt2 = cfg_tblc.get("数据行交替色2", "F2F2F2")

    font_size = Pt(cfg_tbl.get("内容字号pt", 10.5))
    header_bold = cfg_tbl.get("表头加粗", True)

    def _set_cell_shading(cell, color_hex):
        """设置单元格底色。"""
        if not color_hex:
            return
        tcPr = cell._tc.get_or_add_tcPr()
        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
        )
        tcPr.append(shading)

    # 表头
    for j, text in enumerate(header):
        cell = table.cell(0, j)
        cell.text = ""
        _set_cell_shading(cell, header_bg)
        p = cell.paragraphs[0]
        run = p.add_run(text)
        if header_fg:
            run.font.color.rgb = _hex_to_rgb(header_fg)
        set_run_fonts(run, font_size=font_size, bold=header_bold)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 数据行
    for i, row_data in enumerate(rows):
        bg = alt2 if (stripe and i % 2 == 1) else alt1
        for j, text in enumerate(row_data):
            if j >= num_cols:
                break
            cell = table.cell(i + 1, j)
            cell.text = ""
            _set_cell_shading(cell, bg)
            p = cell.paragraphs[0]
            run = p.add_run(str(text))
            set_run_fonts(run, font_size=font_size)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 全框线
    border_sz = cfg_tbl.get("边框粗细", 4)
    border_color = cfg_tblc.get("边框色", "000000")
    tbl_pr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="{border_sz}" w:space="0" w:color="{border_color}"/>'
        f'<w:left w:val="single" w:sz="{border_sz}" w:space="0" w:color="{border_color}"/>'
        f'<w:bottom w:val="single" w:sz="{border_sz}" w:space="0" w:color="{border_color}"/>'
        f'<w:right w:val="single" w:sz="{border_sz}" w:space="0" w:color="{border_color}"/>'
        f'<w:insideH w:val="single" w:sz="{border_sz}" w:space="0" w:color="{border_color}"/>'
        f'<w:insideV w:val="single" w:sz="{border_sz}" w:space="0" w:color="{border_color}"/>'
        f'</w:tblBorders>')
    old = tbl_pr.find(qn('w:tblBorders'))
    if old is not None:
        tbl_pr.remove(old)
    tbl_pr.append(borders)


def render_figure(document, ch_num, fig_num, title, img_path, width_cm=None):
    """渲染图片（分章编号，占位图兜底）。"""
    cfg_fig = _CONFIG.get("配图", {})

    # 图前间距
    p_pre = document.add_paragraph()
    set_spacing(p_pre, after_lines=0, before_lines=cfg_fig.get("图前段前行", 0.5))

    if img_path and os.path.isfile(img_path):
        w = width_cm or cfg_fig.get("默认宽度cm", 12.0)
        max_w = cfg_fig.get("最大宽度cm", 17.0)
        w = min(w, max_w)
        try:
            document.add_picture(img_path, width=Cm(w))
        except Exception:
            # 图片损坏时降级为占位
            p = document.add_paragraph()
            run = p.add_run(cfg_fig.get("占位文本", "【此处为占位图】"))
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            run.font.size = Pt(cfg_fig.get("图题字号pt", 10.5))
    else:
        # 占位图
        p = document.add_paragraph()
        run = p.add_run(cfg_fig.get("占位文本", "【此处为占位图，请替换为实际截图】"))
        run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
        run.font.size = Pt(cfg_fig.get("图题字号pt", 10.5))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 图题
    prefix = cfg_fig.get("图题前缀", "图")
    p = document.add_paragraph()
    run = p.add_run(f"{prefix}{ch_num}-{fig_num} {title}")
    set_run_fonts(run, font_size=Pt(cfg_fig.get("图题字号pt", 10.5)),
                  bold=cfg_fig.get("图题加粗", True))
    set_spacing(p, after_lines=cfg_fig.get("图题段后行", 0.5))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


# ---------------------------------------------------------------------------
# 封面渲染
# ---------------------------------------------------------------------------

def render_cover(document, meta):
    """渲染封面页。"""
    cfg_cover = _CONFIG.get("封面", {})
    if not cfg_cover.get("启用", True):
        return

    # 封面用一个独立节，不编页码
    section = document.sections[0]
    set_page_layout(section)

    # 上方留白（约1/3页面）
    for _ in range(8):
        p = document.add_paragraph()
        set_spacing(p, before_lines=0, after_lines=0, line_spacing=1.0)

    # 项目名称
    title = meta.get("title", "")
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    set_run_fonts(run,
                  font_size=Pt(cfg_cover.get("项目名称字号pt", 22)),
                  bold=cfg_cover.get("项目名称加粗", True))
    set_spacing(p, before_lines=1.0, after_lines=2.0, line_spacing=1.5)

    # 分隔线
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("─" * 40)
    set_run_fonts(run, font_size=Pt(12), bold=False)
    set_spacing(p, before_lines=0.5, after_lines=1.0)

    # 投标人
    bidder = meta.get("bidder", "")
    if bidder:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"投标人：{bidder}")
        set_run_fonts(run,
                      font_size=Pt(cfg_cover.get("投标人字号pt", 16)),
                      bold=cfg_cover.get("投标人加粗", False))
        set_spacing(p, before_lines=0.5, after_lines=0.5)

    # 日期
    date = meta.get("date", "")
    if date:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"日期：{date}")
        set_run_fonts(run,
                      font_size=Pt(cfg_cover.get("日期字号pt", 14)),
                      bold=False)
        set_spacing(p, before_lines=0.5, after_lines=0.5)

    # 分页到目录
    document.add_page_break()


# ---------------------------------------------------------------------------
# 主流程
# ---------------------------------------------------------------------------

def _warn_hand_typed_numbering(body):
    """渲染前门卫：h1-h5 的 text 若以编号前缀开头，警告（会导致双重编号）。

    死规矩：编号必须由引擎自动生成（numPr）；手打前缀会与自动编号叠加成
    「第一章 第一章 运营模式」，必须清除。
    """
    pats = [
        r'^第[一二三四五六七八九十百零]+章',
        r'^[一二三四五六七八九十]+、',
        r'^[（(][一二三四五六七八九十]+[）)]',
        r'^\d+[、.]',
    ]
    warned = 0
    for idx, b in enumerate(body):
        if b.get('type') in ('h1', 'h2', 'h3', 'h4', 'h5'):
            t = (b.get('text') or '').strip()
            if any(re.match(p, t) for p in pats):
                warned += 1
                print("\u26a0\ufe0f  手打编号风险[块#%d]: %s" % (idx + 1, t[:30]))
    if warned:
        print("   提示：标题文本请勿包含编号前缀，编号由引擎自动生成（numPr）；"
              "以上会出现双重编号，请删除前缀后重渲。")


# content.json 允许的内容块类型

def _check_list_items(items, level=0):
    """校验列表项：字符串，或 {t, children} 嵌套结构（与 render_list 的三级渲染对齐）。"""
    if level > 2:
        return "列表嵌套最多三级"
    for x in items:
        if isinstance(x, str):
            continue
        if isinstance(x, dict) and isinstance(x.get("t"), str) \
                and isinstance(x.get("children", []), list):
            err = _check_list_items(x.get("children", []), level + 1)
            if err:
                return err
        else:
            return "列表项必须是字符串，或 {t: 文本, children: [...]} 对象"
    return None

ALLOWED_BLOCK_TYPES = {"h1", "h2", "h3", "h4", "h5", "p", "list", "table", "figure", "hnum"}


def validate_content_json(data, path="content.json"):
    """校验 content.json 结构，返回 (ok, errors_list)。

    纯规则校验，不依赖 python-docx，可在渲染前提前暴露错误，
    避免 json.load 遇到坏数据直接抛栈。
    """
    errors = []
    if not isinstance(data, dict):
        return False, [f"{path}: 根节点必须是 JSON 对象（含 meta/body），实际为 {type(data).__name__}"]

    body = data.get("body")
    if body is None:
        return False, [f"{path}: 缺少必需字段 'body'（内容块数组）"]
    if not isinstance(body, list):
        return False, [f"{path}: 'body' 必须是数组，实际为 {type(body).__name__}"]

    for i, block in enumerate(body):
        where = f"{path} body[{i}]"
        if not isinstance(block, dict):
            errors.append(f"{where}: 内容块必须是对象，实际为 {type(block).__name__}")
            continue
        btype = block.get("type")
        if btype is None:
            errors.append(f"{where}: 缺少 'type' 字段")
            continue
        if btype not in ALLOWED_BLOCK_TYPES:
            errors.append(f"{where}: 未知 type='{btype}'（允许: {', '.join(sorted(ALLOWED_BLOCK_TYPES))}）")
            continue

        if btype in ("h1", "h2", "h3", "h4", "h5"):
            t = block.get("text")
            if not isinstance(t, str) or not t.strip():
                errors.append(f"{where} ({btype}): 'text' 必须是非空字符串")
        elif btype == "p":
            if not isinstance(block.get("text"), str):
                errors.append(f"{where} (p): 'text' 必须是字符串")
        elif btype == "list":
            items = block.get("items")
            if not isinstance(items, list) or not items:
                errors.append(f"{where} (list): 'items' 必须是非空数组")
            else:
                _err = _check_list_items(items)
                if _err:
                    errors.append(f"{where} (list): {_err}")
        elif btype == "hnum":
            t = block.get("text")
            if not isinstance(t, str) or not t.strip():
                errors.append(f"{where} (hnum): 'text' 必须是非空字符串")
        elif btype == "table":
            header, rows = block.get("header"), block.get("rows")
            if not isinstance(header, list) or not isinstance(rows, list):
                errors.append(f"{where} (table): 'header' 和 'rows' 必须是数组")
            elif not header:
                errors.append(f"{where} (table): 'header' 不能为空")
            else:
                for r in rows:
                    if not isinstance(r, list):
                        errors.append(f"{where} (table): 每行必须是数组")
                        break
        elif btype == "figure":
            if not isinstance(block.get("img"), str):
                errors.append(f"{where} (figure): 'img' 必须是字符串（图片路径，可为空串走占位）")

    return (len(errors) == 0), errors


def build(content_path, output_path, format_config=None):
    """主构建函数。"""
    global _CONFIG
    config_path = format_config

    # 自动发现配置
    if not config_path:
        # 相对 content.json 的目录找
        base_dir = os.path.dirname(os.path.abspath(content_path))
        for name in ['format_config.yaml', 'format_config.json']:
            p = os.path.join(base_dir, name)
            if os.path.isfile(p):
                config_path = p
                break
        # 默认模板
        if not config_path:
            pkg_dir = os.path.dirname(os.path.abspath(__file__))
            for name in ['format_config.yaml', 'format_config.json']:
                p = os.path.join(pkg_dir, name)
                if os.path.isfile(p):
                    config_path = p
                    break
            # 包内 templates/ 找（pip install 后默认配置在此）
            if not config_path:
                for name in ['format_config.yaml', 'format_config.json']:
                    p = os.path.join(pkg_dir, 'templates', name)
                    if os.path.isfile(p):
                        config_path = p
                        break
            # 兼容源码/可编辑安装：仓库 templates/ 目录
            if not config_path:
                pkg_parent = os.path.dirname(pkg_dir)
                for name in ['templates/format_config.yaml', 'templates/format_config.json']:
                    p = os.path.join(pkg_parent, name)
                    if os.path.isfile(p):
                        config_path = p
                        break
        if not config_path:
            print("⚠️  未找到排版配置，使用内置默认值")
            config_path = ""

    _CONFIG = _load_config(config_path) if config_path else _load_config("/dev/null")

    # 加载 content.json
    try:
        with open(content_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
    except json.JSONDecodeError as e:
        raise ValueError(
            f"❌ content.json 解析失败（第 {e.lineno} 行第 {e.colno} 列）：{e.msg}\n"
            f"   路径: {content_path}"
        )
    except FileNotFoundError:
        raise ValueError(f"❌ 找不到 content.json: {content_path}")

    ok, errs = validate_content_json(data, content_path)
    if not ok:
        raise ValueError(
            "❌ content.json 结构校验未通过：\n"
            + "\n".join(f"   • {e}" for e in errs)
        )

    meta = data.get('meta', {})
    body = data.get('body', [])
    _warn_hand_typed_numbering(body)

    # 创建文档
    document = Document()

    # 基础样式
    setup_base_styles(document)
    heading_num_id = setup_heading_numbering(document)

    # 封面页（第1节，不编页码）
    render_cover(document, meta)
    cover_page = cfg("封面", "启用") if cfg("封面", "启用") is not False else False

    # 目录页（第2节，不编页码，如果封面已占用第1节则目录为第2节）
    sec0 = document.sections[-1]
    set_page_layout(sec0)
    _build_toc(document)

    # 正文节（页码从1重启）
    document.add_section(WD_SECTION.NEW_PAGE)
    sec1 = document.sections[-1]
    set_page_layout(sec1)
    restart_page_numbering(sec1, start=cfg("页码", "起始页码"))

    clear_footer(sec0) if not cover_page else ()

    add_centered_page_number(sec1)

    # 编号模式
    heading_num_mode = cfg("编号模式", "标题编号") or "全文连续"
    tbl_num_mode = cfg("编号模式", "表格编号") or "每章重置"
    fig_num_mode = cfg("编号模式", "图片编号") or "每章重置"

    # 遍历 body 渲染
    h1_counter = 0
    tbl_counter = {}
    fig_counter = {}
    first_h1 = True

    for block in body:
        btype = block.get('type')
        if btype in ('h1', 'h2', 'h3', 'h4', 'h5'):
            level = int(btype[1])
            if level == 1:
                if heading_num_mode == "每章重置":
                    h1_counter = 1
                    tbl_counter = {}
                    fig_counter = {}
                else:
                    h1_counter += 1
            p = render_heading(document, level, block.get('text', ''), heading_num_id)
            if level == 1:
                if not first_h1 and cfg("标题", "一级标题另起页"):
                    p.paragraph_format.page_break_before = True
                first_h1 = False
        elif btype == 'hnum':
            render_hnum(document, block.get('seq', 1), block.get('text', ''),
                        block.get('level', 0))
        elif btype == 'p':
            render_paragraph(document, block.get('text', ''))
        elif btype == 'list':
            render_list(document, block.get('items', []))
        elif btype == 'table':
            ch = h1_counter if h1_counter > 0 else 1
            if tbl_num_mode == "每章重置":
                tbl_counter[ch] = tbl_counter.get(ch, 0) + 1
            else:
                # 全文连续：用全局计数器
                tbl_counter[0] = tbl_counter.get(0, 0) + 1
                ch = 0
            render_table(document, ch, tbl_counter.get(ch, 1),
                         block.get('title', ''), block.get('header', []),
                         block.get('rows', []))
        elif btype == 'figure':
            ch = h1_counter if h1_counter > 0 else 1
            if fig_num_mode == "每章重置":
                fig_counter[ch] = fig_counter.get(ch, 0) + 1
            else:
                fig_counter[0] = fig_counter.get(0, 0) + 1
                ch = 0
            img = block.get('img')
            base_dir = os.path.dirname(os.path.abspath(content_path))
            if img and not os.path.isabs(img):
                img = os.path.join(base_dir, img)
            w = block.get('width_cm')
            render_figure(document, ch, fig_counter.get(ch, 1),
                          block.get('title', ''), img,
                          float(w) if w is not None else None)

    update_fields_on_open(document)
    document.save(output_path)
    print(f"✅ 已生成: {output_path}")
    print(f"   标题数: {sum(1 for b in body if b.get('type','') in ('h1','h2','h3','h4','h5'))} | "
          f"表: {sum(v for v in tbl_counter.values())} | "
          f"图: {sum(v for v in fig_counter.values())}")
    print("   提示: 用 Word 打开后按 Ctrl+A 再按 F9 更新目录与页码域。")


def _build_toc(document):
    """构建目录页。"""
    cfg_toc = _CONFIG.get("目录", {})
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_indent(title, 0)
    set_spacing(title, before_lines=0.5, after_lines=1.0, line_spacing=1.5)
    r = title.add_run(cfg_toc.get("标题文本", "目 录"))
    r.bold = True
    r.font.size = Pt(cfg_toc.get("标题字号pt", 16))
    set_run_fonts(r)

    p = document.add_paragraph()
    set_indent(p, 0)
    set_spacing(p, line_spacing=1.5)
    lvl_range = cfg_toc.get("级别范围", "1-3")
    add_field(p, f' TOC \\\\o "{lvl_range}" \\\\h \\\\z \\\\u ', '右键此处选择"更新域"以生成目录')


# ---------------------------------------------------------------------------
# CLI 入口
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(description="擎标排版引擎 — content.json → .docx")
    parser.add_argument("content", help="content.json 路径")
    parser.add_argument("output", help="输出 .docx 路径")
    parser.add_argument("--config", "-c", default=None, help="排版配置 YAML/JSON 路径")
    args = parser.parse_args()

    build(args.content, args.output, format_config=args.config)


if __name__ == '__main__':
    main()