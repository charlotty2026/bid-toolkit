#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
排版规则提取器 v0.1 (骨架)
============================
从招标文件（PDF/TXT/MD）中提取排版格式要求，输出结构化规则表。
与 parse_bid.py 和 bid_engine.py 对接，补全「提取→执行」链条。

调用方式:
  python format_rule_extractor.py 招标文件.pdf -o rules.yaml
  python format_rule_extractor.py 招标文件.txt --json

输出:
  {
    "页边距": { "上": "2.54cm", "下": "2.54cm", "左": "2.00cm", "右": "2.00cm" },
    "字体字号": [ { "对象": "正文", "字体": "宋体", "字号": "小四" }, ... ],
    "行距": "1.5倍",
    "编号格式": { "一级": "一、", "二级": "（一）", "三级": "1.", "四级": "1.1" },
    "暗标配置": { "是否暗标": false, "暗标限制": [] },
    "页眉页脚": { "页眉": "项目编号+项目名称", "页脚": "第X页 共Y页" },
    "装订要求": "胶装，A4",
    "盖章要求": "骑缝章+正本公章",
    "章节框架": [ "一、商务部分", "二、技术部分", ... ],
    "内容要求": { "页码从1开始": true, "目录": "自动生成", ... },
    "原文来源": { "页边距": "L42", "字体字号": "L55-L60", ... }   # 行号回溯
  }

v0.1 骨架 — 目前仅合并现有 extract_format_requirements()，后续逐步完善
"""

import re
import sys
import json
import argparse
from pathlib import Path

# DOCX 格式提取器（通用版）
try:
    from .format_rule_docx_extractor import extract_docx_format_rules
    HAS_DOCX_EXTRACTOR = True
except ImportError:
    HAS_DOCX_EXTRACTOR = False


# ============================================================
#  规则表结构
# ============================================================
RULE_TEMPLATE = {
    "页边距": {},            # 上下左右
    "字体字号": [],          # 对象/字体/字号
    "行距": "",              # 1.5倍/固定值28磅
    "编号格式": {},          # 一/（一）/1./1.1
    "暗标配置": {
        "是否暗标": False,
        "暗标限制": []       # ["无页眉","无水印","正本不出现公司名"]
    },
    "页眉页脚": {},          # 页眉内容/页脚内容
    "装订要求": "",          # 胶装/A4/打孔
    "盖章要求": "",          # 骑缝章/公章
    "章节框架": [],          # 投标文件必需的章节结构
    "内容要求": {},          # 页码/目录/双面打印等
    "原文来源": {},          # 每条规则映射到原文行号，便于人工核查
    "原始文本摘要": ""       # 提取规则时匹配到的原文片段
}


# ============================================================
#  单个维度提取函数
# ============================================================

def extract_margins(text, line_map=None):
    """提取页边距要求"""
    result = {}
    source = {}

    patterns = [
        (r'页边距[：:]\s*上下[\s]*([\d.]+)\s*(?:cm|厘米|MM|mm|毫米)', "上下"),
        (r'页边距[：:]\s*左右[\s]*([\d.]+)\s*(?:cm|厘米|MM|mm|毫米)', "左右"),
        (r'上边距[\s]*([\d.]+)\s*(?:cm|厘米)', "上"),
        (r'下边距[\s]*([\d.]+)\s*(?:cm|厘米)', "下"),
        (r'左边距[\s]*([\d.]+)\s*(?:cm|厘米)', "左"),
        (r'右边距[\s]*([\d.]+)\s*(?:cm|厘米)', "右"),
    ]

    # 先找上下左右分别定义的
    for pattern, key in patterns:
        m = re.search(pattern, text)
        if m:
            if key in ("上下",):
                result["上"] = result["下"] = m.group(1) + "cm"
            elif key in ("左右",):
                result["左"] = result["右"] = m.group(1) + "cm"
            else:
                result[key] = m.group(1) + "cm"
            if line_map:
                source[key] = _find_line(text, m.start(), line_map)

    # 再找统一定义的"页边距：上下左右均为Xcm"
    unified = re.search(r'页边距[：:]\s*([\d.]+)\s*(?:cm|厘米)', text)
    if unified and not result:
        val = unified.group(1) + "cm"
        result = {"上": val, "下": val, "左": val, "右": val}

    # P1兜底：如果全空，搜A4相关描述看有没有页边距信息
    if not result:
        a4_margin = re.search(r'A4.{0,80}([\d.]+)\s*(?:cm|厘米)', text)
        if a4_margin:
            val = a4_margin.group(1) + "cm"
            result = {"上": val, "下": val, "左": val, "右": val, "_来源": "A4上下文推断"}

    return result, source


def extract_fonts(text, line_map=None):
    """提取字体字号要求"""
    result = []
    source = {}

    # 定位所有格式声明
    font_pattern = re.compile(
        r'(正文|一级标题|二级标题|三级标题|标题[一二三]|表格|页眉|页脚|标题)[：:]\s*'
        r'(宋体|黑体|仿宋|楷体|Times New Roman|Arial|微软雅黑|方正[^。，；]*?)'
        r'[\s，,；;]*'
        r'([一二三四五六七八九十\d]+号|[0-9.]+pt|小[一二三四五六七八九十]+)',
        re.MULTILINE
    )

    for m in font_pattern.finditer(text):
        obj = m.group(1)
        font = m.group(2).strip()
        size = m.group(3).strip()

        # 处理"标题：黑体，小二号"这种含逗号的
        if len(font) > 10:  # 提取异常，可能逗号截错了
            font = font.split("，")[0]

        entry = {"对象": obj, "字体": font, "字号": size}
        result.append(entry)

        if line_map:
            source[f"{obj}({font}_{size})"] = _find_line(text, m.start(), line_map)

    # 兜底1：找"正文/标题/全文 统一 宋体/黑体/仿宋"这类描述
    if not result:
        fallback = re.findall(r'[。\n]*(?:正文|标题|全文).{0,20}?(?:统一|采用|为|使用|用).{0,10}?(宋体|黑体|仿宋|楷体|Times New Roman)', text)
        for font in fallback:
            result.append({"对象": "正文(兜底)", "字体": font, "字号": "待确认(提取到字体未提取到字号)"})

    # 兜底2：全文搜关键词"宋体/黑体/仿宋"附近的句子，不管格式声明结构
    if not result:
        font_kw = re.search(r'(?:字体|字号|格式).{0,100}(宋体|黑体|仿宋|楷体|小[一二三四五]|五号|四号|三号)', text)
        if font_kw:
            ctx = text[max(0, font_kw.start()-20):font_kw.end()+20]
            result.append({"对象": "全文(兜底2)", "字体": font_kw.group(1), "字号": "待确认", "匹配上下文": ctx[:60]})

    return result, source


def extract_line_spacing(text, line_map=None):
    """提取行距要求"""
    result = ""
    source = ""

    patterns = [
        r'行距[：:]\s*([\d.]+)\s*倍',
        r'行距[：:]\s*固定值[\s]*([\d.]+)\s*(?:磅|pt)',
        r'行距[：:]\s*(单倍|1\.5倍|2倍)',
        r'固定值[\s]*([\d.]+)\s*(?:磅|pt)',
    ]

    for i, pattern in enumerate(patterns):
        m = re.search(pattern, text)
        if m:
            if i == 1:
                result = f"固定值{m.group(1)}磅"
            else:
                result = m.group(1) if m.lastindex else m.group(0)
            if result and line_map:
                source = _find_line(text, m.start(), line_map)
            break

    return result, source


def split_by_headings(chapter_text):
    """按编号标题行切分文本块，返回 [(标题, 正文), ...]
    注意：PDF提取后编号和标题可能挤在同一行（如"1.投标函"），
    取整行作为标题而非仅抓编号前缀。
    """
    pattern = re.compile(
        r'^[  \u3000]*('
        r'[一二三四五六七八九十]+[、.．]|'
        r'[（(][一二三四五六七八九十]+[）)]|'
        r'\d+-\d+[.]|'
        r'\d+[.、]|'
        r'\d+\.\d+'
        r')',
        re.MULTILINE
    )
    blocks = []
    matches = list(pattern.finditer(chapter_text))
    for i, m in enumerate(matches):
        # 从匹配位置取到行尾作为完整标题
        line_end = chapter_text.find('\n', m.start())
        if line_end < 0:
            line_end = len(chapter_text)
        full_title = chapter_text[m.start():line_end].strip()
        # 去掉尾部的页码点填充（如"投标函.............26"→"投标函"）
        full_title = re.sub(r'\.{3,}\d*$', '', full_title).strip()

        start = m.end()
        end = matches[i+1].start() if i+1 < len(matches) else len(chapter_text)
        body = chapter_text[start:end].strip()
        # 过滤：标题长度 3~80 字符
        if 3 <= len(full_title) <= 80:
            blocks.append((full_title, body))
    return blocks


def detect_format_numbering(title):
    """识别一个标题用了什么编号格式，返回编号模式字符串
    优先级：多级编号 > 横杠编号 > 中文序号 > 纯数字（避免"6.1"被识别为"6."）
    """
    # 多级数字: 6.1 6.2（必须在纯数字前匹配）
    m = re.match(r'(\d+\.\d+)', title)
    if m:
        return m.group(1)
    # 横杠数字: 2-2. 2-3. 
    m = re.match(r'(\d+-\d+)[.]', title)
    if m:
        return m.group(1) + '.'
    # 中文序号: 一、二、三、
    m = re.match(r'([一二三四五六七八九十]+)[、.．]', title)
    if m:
        return m.group(1) + '、'
    # 带括号中文序号: （一）（二）
    m = re.match(r'[（(]([一二三四五六七八九十]+)[）)]', title)
    if m:
        return '（' + m.group(1) + '）'
    # 纯数字: 1. 2. 3.
    m = re.match(r'(\d+)[.、]', title)
    if m:
        return m.group(1) + '.'
    return None


def _find_c6_listing(text):
    """从全文定位第六章的格式文件列表区（仅目录部分，不含正文）
    返回 chapter_6_text 或 None
    """
    # 找"第六章 投标文件格式"
    idx = text.find("第六章 投标文件格式")
    if idx < 0:
        idx = text.find("投标文件格式")
    if idx < 0:
        return None

    # 从第六章位置向后找下一个"第.*章"作为结束边界
    next_chapter = re.search(r'第[一二三四五六七八九十]+章', text[idx+5:])
    if next_chapter:
        end = idx + 5 + next_chapter.start()
    else:
        end = min(len(text), idx + 1500)  # 兜底
    return text[idx:end]


def extract_numbering_format(text, line_map=None):
    """
    正确做法：只读第六章的格式文件列表区，逐格式文件记录编号写法。
    不再全文扫频，避免混入招标文件自身的章节编号和正文行。
    """
    result = {}
    source = {}

    chapter_6 = _find_c6_listing(text)
    if not chapter_6:
        return {}, {}

    blocks = split_by_headings(chapter_6)
    if not blocks:
        return {}, {}

    for title, body in blocks:
        fmt = detect_format_numbering(title)
        if fmt:
            # 标题精简：去点填充和多余空格
            short = re.sub(r'\.{3,}\d*$', '', title).strip()
            short = short[:40]
            result[short] = fmt
            if line_map:
                pos = text.find(title)
                if pos >= 0:
                    source[short] = _find_line(text, pos, line_map)

    return result, source


def extract_sealed_bid_config(text, line_map=None):
    """提取暗标/明标要求"""
    result = {"是否暗标": False, "暗标限制": []}
    source = {}

    # 检测是否为暗标
    if re.search(r'暗标|匿名|不得出现.{0,10}(公司|投标人|名称|标识|标志)', text):
        result["是否暗标"] = True

    # 暗标的具体限制
    restrictions = {
        "页眉": r'页眉.{0,30}不得|不得.{0,6}页眉',
        "公司名称": r'(不得|禁止|不允许).{0,15}(公司名称|单位名称|投标人名称)',
        "水印": r'不得.{0,6}(水印|背景图|底纹)',
        "页码": r'页码.{0,20}不出现',
        "封面": r'(封面|封皮).{0,20}不得',
        "装订": r'装订.{0,20}不得|不得.{0,15}装订',
    }

    for res_name, pattern in restrictions.items():
        if re.search(pattern, text):
            result["暗标限制"].append(res_name)
            if line_map:
                m = re.search(pattern, text)
                if m:
                    source[res_name] = _find_line(text, m.start(), line_map)

    return result, source


def extract_header_footer(text, line_map=None):
    """提取页眉页脚要求"""
    result = {}
    source = {}

    for hf in ["页眉", "页脚"]:
        m = re.search(rf'{hf}[：:]\s*([^。\n]+?)(?:。|\n|$)', text)
        if m:
            result[hf] = m.group(1).strip()
            if line_map:
                source[hf] = _find_line(text, m.start(), line_map)

    return result, source


def extract_binding_requirements(text, line_map=None):
    """提取装订要求"""
    result = ""
    source = ""

    m = re.search(r'(装订|胶装|打孔|活页|精装|简装).{0,30}?(要求|方式|形式)?', text)
    if m:
        # 取匹配行附近的整句
        ctx = text[m.start():m.end()]
        # 扩大范围到句号
        ctx_start = text.rfind("。", 0, m.start())
        ctx_end = text.find("。", m.end())
        if ctx_start < 0:
            ctx_start = max(0, m.start() - 30)
        if ctx_end < 0:
            ctx_end = min(len(text), m.end() + 30)
        result = text[max(0, ctx_start):min(len(text), ctx_end)].strip()
        if line_map:
            source = _find_line(text, m.start(), line_map)

    # 检查A4等纸张要求
    if not result:
        paper = re.search(r'(A4|A3|纸张|幅面).{0,20}?(要求|标准|规格)', text)
        if paper:
            result = f"纸张: A4 (默认)"

    return result, source


def extract_seal_requirements(text, line_map=None):
    """提取盖章要求"""
    result = ""
    source = ""

    m = re.search(r'(盖章|公章|签章|骑缝章|盖章要求|印章).{0,50}?(公章|骑缝|正本|副本)?', text)
    if m:
        ctx_start = text.rfind("。", 0, m.start())
        ctx_end = text.find("。", m.end())
        if ctx_start < 0: ctx_start = max(0, m.start() - 20)
        if ctx_end < 0: ctx_end = min(len(text), m.end() + 20)
        result = text[max(0, ctx_start):min(len(text), ctx_end)].strip()
        if line_map:
            source = _find_line(text, m.start(), line_map)

    return result, source


def extract_chapter_framework(text, line_map=None):
    """
    只取第六章的格式文件名作为章节框架。
    使用与 extract_numbering_format 相同的定位方法。
    """
    result = []
    source = {}

    chapter_6 = _find_c6_listing(text)
    if not chapter_6:
        return [], {}

    # 匹配编号开头的标题行（支持数字./数字-数字./中文序号）
    title_pattern = re.compile(
        r'^[  \u3000]*('
        r'[一二三四五六七八九十]+[、.．]'
        r'|[（(][一二三四五六七八九十]+[）)]'
        r'|\d+-\d+[.]'
        r'|\d+[.、]'
        r')([^\n]+)',
        re.MULTILINE
    )

    seen = set()
    for m in title_pattern.finditer(chapter_6):
        prefix = m.group(1).strip()
        title_text = m.group(2).strip() if m.lastindex and m.lastindex >= 2 else ''
        # 去点填充和多余空格
        title_text = re.sub(r'\.{3,}\d*$', '', title_text).strip()
        title_text = re.sub(r'\s{2,}', ' ', title_text)
        # 过滤
        if not title_text or len(title_text) > 60:
            continue
        if re.match(r'^\d+$', title_text):
            continue
        entry = f"{prefix} {title_text}"
        if entry not in seen:
            seen.add(entry)
            result.append(entry)
            if line_map:
                source[entry] = _find_line(text, m.start(), line_map)

    return result, source


def extract_content_requirements(text, line_map=None):
    """提取内容要求（页码、目录、双面打印等）"""
    result = {}
    source = {}

    checks = {
        "页码从1开始": r'页码.{0,10}(?:从[01]|起始).{0,5}(?:开始|起)',
        "目录自动生成": r'(?:目录|目次).{0,10}(?:自动|生成|包含)',
        "双面打印": r'(?:双面|正反面).{0,10}(?:打印|印刷)',
        "单面打印": r'(?:单面).{0,10}(?:打印|印刷)',
        "分册装订": r'(?:分册|分卷|分本).{0,10}(?:装订|编制)',
        "页码位置": r'页码.{0,20}(?:位置|居中|右对齐|下方)',
    }

    for key, pattern in checks.items():
        m = re.search(pattern, text)
        if m:
            result[key] = True
            if line_map:
                source[key] = _find_line(text, m.start(), line_map)

    return result, source


# ============================================================
#  辅助
# ============================================================

def _find_line(text, char_pos, line_map):
    """从字符位置找到行号（用于原文回溯）"""
    if not line_map:
        return str(char_pos)
    for line_num, start, end in line_map:
        if start <= char_pos < end:
            return f"L{line_num}"
    return f"~L{len(line_map)}"


def build_line_map(text):
    """构建行号映射 [(行号, 起始字符, 结束字符), ...]"""
    lines = text.splitlines(keepends=True)
    line_map = []
    pos = 0
    for i, line in enumerate(lines, 1):
        line_map.append((i, pos, pos + len(line)))
        pos += len(line)
    return line_map


# ============================================================
#  主提取流程
# ============================================================

def extract_all_rules(text):
    """从招标文件文本中提取全部排版规则"""
    line_map = build_line_map(text)
    rules = dict(RULE_TEMPLATE)  # 深拷贝

    # 逐个维度提取
    rules["页边距"], s_margin = extract_margins(text, line_map)
    rules["字体字号"], s_font = extract_fonts(text, line_map)
    rules["行距"], s_spacing = extract_line_spacing(text, line_map)
    rules["编号格式"], s_num = extract_numbering_format(text, line_map)
    rules["暗标配置"], s_sealed = extract_sealed_bid_config(text, line_map)
    rules["页眉页脚"], s_hf = extract_header_footer(text, line_map)
    rules["装订要求"], s_binding = extract_binding_requirements(text, line_map)
    rules["盖章要求"], s_seal = extract_seal_requirements(text, line_map)
    rules["章节框架"], s_chapter = extract_chapter_framework(text, line_map)
    rules["内容要求"], s_content = extract_content_requirements(text, line_map)

    # 原文来源
    source = {}
    for d in [s_margin, s_font, s_spacing, s_num, s_sealed, s_hf,
              s_binding, s_seal, s_chapter, s_content]:
        if isinstance(d, dict):
            source.update(d)
    rules["原文来源"] = source

    return rules


# ============================================================
#  与现有工具的对接
# ============================================================


def _merge_docx_rules(docx_rules, text_rules):
    """
    DOCX提取的样式规则 → 合并到文本提取规则表中。
    策略：DOCX精确格式优先覆盖，文本提取语义规则补充。
    字体字号：DOCX检测到的正文/标题/表格各保留，
    文本提取到的其他语义规则（如编号格式/暗标/章节框架）仍保留原文结果。
    """
    from copy import deepcopy
    merged = deepcopy(text_rules)

    legacy = docx_rules.to_legacy_dict()

    # 字体字号 — DOCX精确检测优先
    font_entries = []
    if "字体" in legacy:
        font_entries.append({
            "对象": "正文（DOCX自动检测）",
            "字体": legacy["字体"],
            "字号": legacy.get("字号", "")
        })
    if "标题字体" in legacy:
        font_entries.append({
            "对象": "标题（DOCX自动检测）",
            "字体": legacy["标题字体"],
            "字号": legacy.get("标题字号", "")
        })
    if "表格字体" in legacy:
        font_entries.append({
            "对象": "表格（DOCX自动检测）",
            "字体": legacy["表格字体"],
            "字号": legacy.get("表格字号", "")
        })

    # 补充文本提取到的字体规则，去重
    seen_combos = set(
        (e["对象"], e["字体"], e.get("字号", ""))
        for e in font_entries
    )
    for entry in merged.get("字体字号", []):
        key = (entry["对象"], entry["字体"], entry.get("字号", ""))
        if key not in seen_combos:
            seen_combos.add(key)
            font_entries.append(entry)

    if font_entries:
        merged["字体字号"] = font_entries

    # 页边距 — DOCX精确值覆盖
    if "页边距" in legacy:
        merged["页边距"] = legacy["页边距"]
    # 行距
    if "行距" in legacy:
        merged["行距"] = legacy["行距"]
    # 纸张
    if "纸张" in legacy:
        merged["纸张"] = legacy["纸张"]
    return merged



def to_engine_config(rules):
    """将提取的规则表 → bid_engine.py 的 config.yaml 格式"""
    config = {}

    # 字体字号 → heading/body specs
    for entry in rules.get("字体字号", []):
        obj = entry.get("对象", "")
        font = entry.get("字体", "宋体")
        size_pt = _cn_size_to_pt(entry.get("字号", "小四"))

        if "正文" in obj:
            config["body"] = {
                "font": font, "size": size_pt,
                "first_indent_chars": 2, "line_spacing": 1.5
            }
        elif "标题" in obj or "一级" in obj:
            config["headings"] = config.get("headings", {})
            level = 1 if "一" in obj else 2 if "二" in obj or "二级" in obj else 3
            config["headings"][level] = {
                "font": font, "size": size_pt, "bold": True
            }

    # 行距
    if rules.get("行距"):
        config["行距"] = rules["行距"]

    # 页边距
    if rules.get("页边距"):
        margins = rules["页边距"]
        config["page"] = {
            "margin_top": _cm_to_cm(margins.get("上", "2.54")),
            "margin_bottom": _cm_to_cm(margins.get("下", "2.54")),
            "margin_left": _cm_to_cm(margins.get("左", "2.00")),
            "margin_right": _cm_to_cm(margins.get("右", "2.00")),
        }

    # 暗标
    if rules.get("暗标", {}).get("是否暗标", False):
        config["anonymous"] = {"enabled": True}

    return config


def _cn_size_to_pt(size_str):
    """中文字号 → pt（近似）"""
    mapping = {
        "初号": 42, "小初": 36,
        "一号": 28, "小一": 24,
        "二号": 22, "小二": 18,
        "三号": 16, "小三": 15,
        "四号": 14, "小四": 12,
        "五号": 10.5, "小五": 9,
        "六号": 7.5, "小六": 6.5,
        "七号": 5.5, "八号": 5,
    }
    if size_str in mapping:
        return mapping[size_str]
    # 尝试提取数字pt
    m = re.search(r'([\d.]+)\s*pt', size_str)
    if m:
        return float(m.group(1))
    return 12  # 默认小四


def _cm_to_cm(cell_value):
    """统一页面边距为cm值"""
    if isinstance(cell_value, str):
        m = re.match(r'([\d.]+)', cell_value)
        if m:
            return float(m.group(1))
    return float(cell_value)


# ============================================================
#  CLI
# ============================================================

def main():
    parser = argparse.ArgumentParser(
        description="排版规则提取器 — 从招标文件提取格式要求",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
使用示例:
  python format_rule_extractor.py 招标文件.pdf
  python format_rule_extractor.py 招标文件.txt -o rules.yaml
  python format_rule_extractor.py 招标文件.md --json

与bid-toolkit集成:
  bid format 招标文件.pdf   # 命令行别名
        """
    )
    parser.add_argument("input", help="招标文件路径 (PDF/TXT/MD)")
    parser.add_argument("-o", "--output", help="输出YAML文件路径")
    parser.add_argument("--json", action="store_true", help="输出JSON格式")
    parser.add_argument("--engine-config", action="store_true",
                        help="同时输出bid_engine.py可用的config.yaml")
    args = parser.parse_args()

    # 读取文件
    path = Path(args.input)
    if not path.exists():
        print(f"❌ 文件不存在: {args.input}")
        sys.exit(1)

    suffix = path.suffix.lower()

    # ── DOCX 分支：优先用样式提取器 ──
    if suffix == ".docx" and HAS_DOCX_EXTRACTOR:
        print("🔍 DOCX样式提取中...")
        docx_rules = extract_docx_format_rules(str(path), verbose=True)

        # 提取纯文本用于语义规则
        from docx import Document as DocxDocument
        doc = DocxDocument(str(path))
        text = "\n".join(p.text for p in doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += "\n" + "\n".join(p.text for p in cell.paragraphs)

        text_rules = extract_all_rules(text)
        merged = _merge_docx_rules(docx_rules, text_rules)

        if args.json or args.output:
            output = json.dumps(merged, ensure_ascii=False, indent=2, default=str)
        else:
            output = _format_pretty(merged)
        if args.output:
            with open(args.output, "w", encoding="utf-8") as f:
                f.write(output)
            print(f"✅ 规则表已导出: {args.output}")
        else:
            print(output)
        if args.engine_config:
            config = to_engine_config(merged)
            engine_path = Path(args.output or ".").parent / "format_config.yaml"
            try:
                import yaml
                with open(engine_path, "w", encoding="utf-8") as f:
                    yaml.dump(config, f, allow_unicode=True, default_flow_style=False)
                print(f"✅ 引擎配置已导出: {engine_path}")
            except ImportError:
                print("⚠️  需要 pyyaml 才能导出engine config")
        return merged

    # 如果是PDF，尝试用parse_bid提取文本
    if suffix == ".pdf":
        try:
            # 优先使用 odl-pdf
            import subprocess
            result = subprocess.run(
                ["odl-pdf", str(path)],
                capture_output=True, text=True, timeout=60
            )
            text = result.stdout
        except Exception:
            # 兜底用pymupdf
            try:
                import fitz
                doc = fitz.open(str(path))
                text = "\n".join(page.get_text() for page in doc)
            except ImportError:
                print("❌ PDF解析需要 odl-pdf 或 pymupdf")
                sys.exit(1)
    else:
        text = path.read_text("utf-8", errors="ignore")

    # 提取规则
    rules = extract_all_rules(text)

    # 输出
    if args.json or args.output:
        output = json.dumps(rules, ensure_ascii=False, indent=2, default=str)
    else:
        output = _format_pretty(rules)

    if args.output:
        with open(args.output, "w", encoding="utf-8") as f:
            f.write(output)
        print(f"✅ 规则表已导出: {args.output}")
    else:
        print(output)

    # 引擎配置文件
    if args.engine_config:
        config = to_engine_config(rules)
        engine_path = Path(args.output or ".").parent / "format_config.yaml"
        try:
            import yaml
            with open(engine_path, "w", encoding="utf-8") as f:
                yaml.dump(config, f, allow_unicode=True, default_flow_style=False)
            print(f"✅ 引擎配置已导出: {engine_path}")
        except ImportError:
            print("⚠️  需要 pyyaml 才能导出engine config")


def _format_pretty(rules):
    """美化打印"""
    lines = []
    lines.append("=" * 60)
    lines.append("📋 排版规则提取结果")
    lines.append("=" * 60)

    sections = [
        ("页边距", "页边距"),
        ("字体字号", "字体字号"),
        ("行距", "行距"),
        ("编号格式", "编号格式"),
        ("暗标配置", "暗标配置"),
        ("页眉页脚", "页眉页脚"),
        ("装订要求", "装订要求"),
        ("盖章要求", "盖章要求"),
        ("章节框架", "章节框架"),
        ("内容要求", "内容要求"),
    ]

    for key, label in sections:
        val = rules.get(key, {})
        if not val:
            continue
        lines.append(f"\n📌 {label}:")
        if isinstance(val, dict):
            for k, v in val.items():
                lines.append(f"  {k}: {v}")
        elif isinstance(val, list):
            for item in val:
                lines.append(f"  - {item}")
        else:
            lines.append(f"  {val}")

    # 来源
    src = rules.get("原文来源", {})
    if src:
        lines.append(f"\n📎 原文来源:")
        for k, v in list(src.items())[:10]:
            lines.append(f"  {k}: {v}")

    lines.append("\n" + "=" * 60)
    lines.append("💡 使用 `-o rules.yaml` 导出结构化规则")
    lines.append("💡 使用 `--engine-config` 同时导出bid_engine可用的配置")
    lines.append("=" * 60)

    return "\n".join(lines)


if __name__ == "__main__":
    main()
