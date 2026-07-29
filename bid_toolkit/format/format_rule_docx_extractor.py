"""
bid-toolkit: DOCX 格式规则提取器（通用版 v2）
===============================================
从 .docx 招标文件中自动提取字体/字号/页边距/行距等格式规则。
设计原则：
  1. 自动检测正文字体（取 run 级最频繁字体/字号，跳过加粗 run）
  2. 自动检测标题字体（三层策略：不同字体→加粗+大字号→纯字号→Heading样式）
  3. 读取 section 页边距，取多数 section 的共同值
  4. 行距两层回退：Normal 样式 → run 级行距统计
  5. 所有推断均可通过参数配置

用法:
  from format_rule_docx_extractor import extract_docx_format_rules
  rules = extract_docx_format_rules("招标文件.docx")
"""

import logging
from collections import Counter
from dataclasses import dataclass, field, asdict
from typing import Dict, Optional, List, Tuple

try:
    from docx import Document
    from docx.shared import Pt, Emu
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# 可配置常量
# ---------------------------------------------------------------------------

DEFAULT_CONFIG = {
    # 正文识别：run 长度阈值（跳过太短的无意义 run）
    "body_run_min_length": 5,
    # 正文采样段数上限
    "body_sample_paragraphs": 500,
    # 标题字号阈值（pt）：>=此值视为标题 run
    "title_font_size_threshold_pt": 12,
    # 标题检测：纯字号判定时，字号必须 > body pt + 此差值 才可信
    "title_pure_size_delta": 2,
    # 标题检测：最小累加文本长度（避免零星空 run 干扰）
    "title_min_total_len": 10,
    # 页边距：取 section 的多数值（模式），单位为 EMU
    "margin_unit": "cm",
    # 页边距：section 为空时的默认值。中文标书通常无默认值，设为 0 让下游决策。
    "margin_default_cm": 0,
    # table 采样数上限
    "table_sample_tables": 20,
    "table_sample_rows_per_table": 10,
    # 行距回退：最少段落数才可信
    "line_spacing_min_para_count": 3,
}


# ---------------------------------------------------------------------------
# 数据结构
# ---------------------------------------------------------------------------

@dataclass
class FontRule:
    font_name: str = ""
    font_size_pt: float = 0.0
    bold: Optional[bool] = None
    source: str = ""  # "body_runs" | "style_normal" | "detected"


@dataclass
class MarginRule:
    left_cm: float = 0.0
    right_cm: float = 0.0
    top_cm: float = 0.0
    bottom_cm: float = 0.0
    source: str = ""  # "section_mode"


@dataclass
class LineSpacingRule:
    """行距规则，描述行距类型和数值"""
    spacing: str = ""        # 例如 "1.5倍" 或 "固定值24磅"
    source: str = ""         # "style_normal" | "runs_inferred"


@dataclass
class DocxFormatRules:
    """docx 格式规则聚合"""
    body_font: FontRule = field(default_factory=FontRule)
    title_font: FontRule = field(default_factory=FontRule)
    table_font: FontRule = field(default_factory=FontRule)
    margins: MarginRule = field(default_factory=MarginRule)
    line_spacing: str = ""       # 例如 "1.5倍" 或 "固定值24磅"（兼容旧调用方）
    page_size: str = ""          # 例如 "A4 (21.0x29.7cm)"
    line_spacing_detail: LineSpacingRule = field(default_factory=LineSpacingRule)

    def to_legacy_dict(self) -> Dict:
        """转换为与 PDF 提取器兼容的旧格式 dict"""
        result = {}
        # 字体字号
        if self.body_font.font_name:
            result["字体"] = self.body_font.font_name
            result["字号"] = f"{self.body_font.font_size_pt:.0f}pt" if self.body_font.font_size_pt else ""

        # 标题字体 (独立 key，区别于正文)
        if self.title_font.font_name:
            result["标题字体"] = self.title_font.font_name
            if self.title_font.font_size_pt:
                result["标题字号"] = f"{self.title_font.font_size_pt:.0f}pt"

        # 表格字体
        if self.table_font.font_name and self.table_font.font_name not in (self.body_font.font_name, ""):
            result["表格字体"] = self.table_font.font_name
            if self.table_font.font_size_pt:
                result["表格字号"] = f"{self.table_font.font_size_pt:.0f}pt"

        # 页边距
        if any([self.margins.top_cm, self.margins.left_cm]):
            result["页边距"] = (
                f"上{self.margins.top_cm:.1f}cm "
                f"下{self.margins.bottom_cm:.1f}cm "
                f"左{self.margins.left_cm:.1f}cm "
                f"右{self.margins.right_cm:.1f}cm"
            )

        # 行距
        if self.line_spacing:
            result["行距"] = self.line_spacing

        if self.page_size:
            result["纸张"] = self.page_size

        return result


# ---------------------------------------------------------------------------
# 辅助函数
# ---------------------------------------------------------------------------

def _emu_to_cm(v: int) -> float:
    """EMU 转 cm"""
    return v / 360000.0 if v else 0.0


def _pt_from_run(run) -> Optional[float]:
    """从 run 获取字号(pt)，处理 None 和继承"""
    if run.font.size is None:
        return None
    return run.font.size / 12700.0


def _is_bold_run(run) -> bool:
    """判断 run 是否加粗（True=加粗，False/None=未加粗）"""
    return run.font.bold is True


# ---------------------------------------------------------------------------
# Run 数据采集（一次性遍历，避免重复扫描）
# ---------------------------------------------------------------------------

@dataclass
class RunRecord:
    """单条 run 的采集记录"""
    font_name: Optional[str] = None
    font_size_pt: Optional[int] = None  # round(pt)
    bold: bool = False
    text_len: int = 0


def _collect_runs(doc: "Document", config: dict) -> List[RunRecord]:
    """遍历段落，收集非空 run 信息。返回 [RunRecord, ...]"""
    records = []
    min_len = config["body_run_min_length"]
    for para in doc.paragraphs[: config["body_sample_paragraphs"]]:
        text = para.text.strip()
        if not text or len(text) < min_len:
            continue
        for run in para.runs:
            t = run.text.strip()
            if not t or len(t) < 2:
                continue
            fn = run.font.name
            fs_raw = _pt_from_run(run)
            fs = round(fs_raw) if fs_raw is not None else None
            records.append(RunRecord(
                font_name=fn,
                font_size_pt=fs,
                bold=_is_bold_run(run),
                text_len=len(t),
            ))
    return records


# ---------------------------------------------------------------------------
# 核心提取逻辑
# ---------------------------------------------------------------------------

def _detect_body_font(
    doc: "Document",
    runs_data: List[RunRecord],
    config: dict,
) -> FontRule:
    """
    正文字体检测策略：
    1. 优先从 Normal 样式读取字体/字号
    2. 若样式为空（直接格式场景），从正文 run 中统计（跳过加粗 run）
    3. 分离统计回退（font / size 分别 count）
    """
    # 策略1：样式
    normal_style = None
    try:
        normal_style = doc.styles['Normal']
    except KeyError:
        pass

    if normal_style and normal_style.font and normal_style.font.name:
        fs = normal_style.font.size / 12700.0 if normal_style.font.size else 0.0
        return FontRule(
            font_name=normal_style.font.name,
            font_size_pt=fs,
            bold=normal_style.font.bold,
            source="style_normal"
        )

    # 策略2：run 级统计（跳过加粗 run）
    run_pairs = Counter()
    pure_fonts = Counter()
    pure_sizes = Counter()

    for rec in runs_data:
        if rec.bold:
            continue  # 正文通常不加粗
        if rec.font_name and rec.font_size_pt is not None:
            key = (rec.font_name, rec.font_size_pt)
            run_pairs[key] += rec.text_len  # 用文本长度加权
        if rec.font_name:
            pure_fonts[rec.font_name] += rec.text_len
        if rec.font_size_pt is not None:
            pure_sizes[rec.font_size_pt] += rec.text_len

    if run_pairs:
        (best_font, best_size), _ = run_pairs.most_common(1)[0]
        return FontRule(
            font_name=best_font,
            font_size_pt=float(best_size),
            source="body_runs_paired"
        )

    # 策略3：分离统计回退
    best_fn = pure_fonts.most_common(1)[0][0] if pure_fonts else ""
    best_sz = float(pure_sizes.most_common(1)[0][0]) if pure_sizes else 10.0
    return FontRule(font_name=best_fn, font_size_pt=best_sz, source="body_runs_split")


def _detect_title_font(
    doc: "Document",
    runs_data: List[RunRecord],
    body_font: FontRule,
    config: dict,
) -> FontRule:
    """
    标题字体检测，三层策略：
      1. 不同字体+大字号 → 明显标题（最可信）
      2. 同字体+加粗+大字号 → 加粗辅助判定
      3. 同字体+不加粗+大字号 → 纯字号判定（阈值提高 config["title_pure_size_delta"] pt）
      4. 回退 → Heading 样式
    """
    threshold = config["title_font_size_threshold_pt"]
    pure_delta = config["title_pure_size_delta"]
    min_total = config["title_min_total_len"]

    diff_counter = Counter()     # (font, pt) → total_len（不同字体）
    same_bold_counter = Counter()   # 同字体+加粗
    same_plain_counter = Counter()  # 同字体+不加粗

    for rec in runs_data:
        if rec.font_size_pt is None or rec.font_size_pt < threshold:
            continue
        key = (rec.font_name, rec.font_size_pt)
        fn = rec.font_name or ""
        if fn != body_font.font_name:
            diff_counter[key] += rec.text_len
        elif rec.bold:
            same_bold_counter[key] += rec.text_len
        else:
            same_plain_counter[key] += rec.text_len

    # 1. 不同字体
    if diff_counter:
        (best_fn, best_sz), total_len = diff_counter.most_common(1)[0]
        if total_len >= min_total:
            return FontRule(
                font_name=best_fn,
                font_size_pt=float(best_sz),
                bold=None,
                source="title_runs_diff"
            )

    # 2. 加粗+大字号
    if same_bold_counter:
        (best_fn, best_sz), total_len = same_bold_counter.most_common(1)[0]
        if total_len >= min_total:
            return FontRule(
                font_name=best_fn,
                font_size_pt=float(best_sz),
                bold=True,
                source="title_runs_same_bold"
            )

    # 3. 纯大字号（必须 > body pt + delta 才可信）
    if same_plain_counter:
        body_pt = body_font.font_size_pt
        for (fn, sz), total_len in same_plain_counter.most_common(3):
            if body_pt and sz >= round(body_pt) + pure_delta and total_len >= min_total:
                return FontRule(
                    font_name=fn,
                    font_size_pt=float(sz),
                    bold=None,
                    source="title_runs_same_plain"
                )
        # 兜底：取最常出现的
        (best_fn, best_sz), total_len = same_plain_counter.most_common(1)[0]
        if total_len >= min_total:
            return FontRule(
                font_name=best_fn,
                font_size_pt=float(best_sz),
                bold=None,
                source="title_runs_same_plain_fallback"
            )

    # 4. 回退：Heading 1 样式
    for hname in ['Heading 1', '标题 1', 'heading 1']:
        try:
            hs = doc.styles[hname]
            fn = hs.font.name
            fs = hs.font.size / 12700.0 if hs.font.size else 0.0
            if fn:
                return FontRule(font_name=fn, font_size_pt=fs, source="style_heading1")
        except KeyError:
            pass

    return FontRule(source="undetected")


def _detect_table_font(
    doc: "Document",
    runs_data: List[RunRecord],
    body_font: FontRule,
    config: dict,
) -> FontRule:
    """表格字体：采样表格内 run"""
    run_pairs = Counter()
    n_tables = min(len(doc.tables), config["table_sample_tables"])
    n_rows = config["table_sample_rows_per_table"]

    for table in doc.tables[:n_tables]:
        for row in table.rows[:n_rows]:
            for cell in row.cells:
                for para in cell.paragraphs[:2]:
                    for run in para.runs:
                        if not run.text.strip():
                            continue
                        fn = run.font.name
                        fs_raw = _pt_from_run(run)
                        fs = round(fs_raw) if fs_raw else None
                        if fn and fs:
                            run_pairs[(fn, fs)] += 1

    if run_pairs:
        (best_fn, best_sz), _ = run_pairs.most_common(1)[0]
        return FontRule(font_name=best_fn, font_size_pt=float(best_sz), source="table_runs")

    return FontRule(source="undetected")


def _detect_margins(doc: "Document", config: dict) -> MarginRule:
    """页边距：取 section 的多数值。空 section 返回 config 中定义的默认值。"""
    from statistics import mode

    lefts, rights, tops, bottoms = [], [], [], []
    for sec in doc.sections:
        if sec.left_margin:
            lefts.append(round(_emu_to_cm(sec.left_margin), 2))
        if sec.right_margin:
            rights.append(round(_emu_to_cm(sec.right_margin), 2))
        if sec.top_margin:
            tops.append(round(_emu_to_cm(sec.top_margin), 2))
        if sec.bottom_margin:
            bottoms.append(round(_emu_to_cm(sec.bottom_margin), 2))

    default_cm = config.get("margin_default_cm", 0)

    def _safe_mode(vals):
        if not vals:
            return default_cm
        try:
            return mode(vals)
        except Exception:
            return Counter(vals).most_common(1)[0][0]

    return MarginRule(
        left_cm=_safe_mode(lefts),
        right_cm=_safe_mode(rights),
        top_cm=_safe_mode(tops),
        bottom_cm=_safe_mode(bottoms),
        source="section_mode"
    )


def _detect_line_spacing(doc: "Document", config: dict) -> LineSpacingRule:
    """
    行距检测，两层回退：
      1. Normal 样式行距
      2. 正文段落的 paragraph_format.line_spacing 高频值推理
    """
    min_para = config.get("line_spacing_min_para_count", 3)

    # 策略1：Normal 样式
    try:
        ns = doc.styles['Normal']
        pf = ns.paragraph_format
        if pf and pf.line_spacing is not None:
            ls = pf.line_spacing
            if hasattr(ls, 'pt'):
                return LineSpacingRule(spacing=f"固定值{ls.pt:.0f}磅", source="style_normal")
            else:
                val = f"{ls:.1f}"
                if val.endswith(".0"):
                    val = val[:-2]
                return LineSpacingRule(spacing=f"{val}倍", source="style_normal")
    except Exception:
        pass

    # 策略2：run 级行距推理
    spacing_counter = Counter()
    for para in doc.paragraphs[:200]:
        pf = para.paragraph_format
        if pf and pf.line_spacing is not None:
            ls = pf.line_spacing
            if hasattr(ls, 'pt'):
                label = f"固定值{ls.pt:.0f}磅"
            else:
                val = f"{ls:.1f}"
                if val.endswith(".0"):
                    val = val[:-2]
                label = f"{val}倍"
            spacing_counter[label] += 1

    if spacing_counter:
        best, cnt = spacing_counter.most_common(1)[0]
        if cnt >= min_para:
            return LineSpacingRule(spacing=best, source="runs_inferred")

    return LineSpacingRule(source="undetected")


def _detect_page_size(doc: "Document") -> str:
    """纸张：A4/Letter 等"""
    if not doc.sections:
        return ""
    sec = doc.sections[0]
    w = _emu_to_cm(sec.page_width)
    h = _emu_to_cm(sec.page_height)
    if abs(w - 21.0) < 0.5 and abs(h - 29.7) < 0.5:
        return "A4"
    if abs(w - 21.59) < 0.5 and abs(h - 27.94) < 0.5:
        return "Letter"
    return f"{w:.1f}x{h:.1f}cm"


# ---------------------------------------------------------------------------
# 主入口
# ---------------------------------------------------------------------------

def extract_docx_format_rules(
    docx_path: str,
    config: Optional[dict] = None,
    verbose: bool = False,
) -> DocxFormatRules:
    """
    从 docx 提取格式规则。

    Args:
        docx_path: .docx 文件绝对路径
        config: 可选配置字典，覆盖 DEFAULT_CONFIG
        verbose: 是否打印调试信息

    Returns:
        DocxFormatRules 对象

    Raises:
        ImportError: python-docx 未安装时抛出
        FileNotFoundError: 文件不存在
    """
    if not HAS_DOCX:
        raise ImportError("python-docx not installed. Run: pip install python-docx")

    cfg = {**DEFAULT_CONFIG, **(config or {})}

    if verbose:
        logging.basicConfig(level=logging.DEBUG)
        logger.debug(f"Extracting from: {docx_path}")

    import os
    if not os.path.isfile(docx_path):
        raise FileNotFoundError(f"File not found: {docx_path}")

    doc = Document(docx_path)

    # 一次性收集 run 数据，减少遍历次数
    runs_data = _collect_runs(doc, cfg)

    # 正文字体
    body_font = _detect_body_font(doc, runs_data, cfg)
    if verbose:
        logger.debug(f"Body font: {body_font}")

    # 标题字体
    title_font = _detect_title_font(doc, runs_data, body_font, cfg)
    if verbose:
        logger.debug(f"Title font: {title_font}")

    # 表格字体
    table_font = _detect_table_font(doc, runs_data, body_font, cfg)
    if verbose:
        logger.debug(f"Table font: {table_font}")

    # 页边距
    margins = _detect_margins(doc, cfg)
    if verbose:
        logger.debug(f"Margins: {margins}")

    # 行距（两层回退）
    ls_rule = _detect_line_spacing(doc, cfg)
    if verbose:
        logger.debug(f"Line spacing: {ls_rule.spacing} (source: {ls_rule.source})")

    # 纸张
    page_size = _detect_page_size(doc)
    if verbose:
        logger.debug(f"Page size: {page_size}")

    return DocxFormatRules(
        body_font=body_font,
        title_font=title_font,
        table_font=table_font,
        margins=margins,
        line_spacing=ls_rule.spacing,
        page_size=page_size,
        line_spacing_detail=ls_rule,
    )


# ---------------------------------------------------------------------------
# 测试入口
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    import sys, json

    path = sys.argv[1] if len(sys.argv) > 1 else None
    if not path:
        print("Usage: python format_rule_docx_extractor.py <docx_path>")
        sys.exit(1)

    rules = extract_docx_format_rules(path, verbose=True)

    print("\n=== Legacy dict ===")
    print(json.dumps(rules.to_legacy_dict(), ensure_ascii=False, indent=2))

    print("\n=== Full dataclass ===")
    print(json.dumps(asdict(rules), ensure_ascii=False, indent=2, default=str))

    print(f"\n=== 行距详情 ===")
    print(f"  值: {rules.line_spacing}")
    print(f"  来源: {rules.line_spacing_detail.source}")
