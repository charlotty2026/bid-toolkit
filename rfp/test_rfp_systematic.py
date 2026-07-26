#!/usr/bin/env python3
"""
RFP Generator 系统性测试脚本
测试维度：
1. 占位符问题（目标：降到10个以内）
2. 评分表/废标条款/采购需求三个核心模板质量
3. 三类差异化（货物/服务/工程结构是否正确区分）
4. 合规检查（6条红线检出能力）
"""
import sys
import os
import re
import json
import tempfile

# 加入rfp模块路径（动态获取，兼容不同部署环境）
RFP_DIR = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, RFP_DIR)

from rfp_generator import generate_markdown, SCORING_TEMPLATES, DETAILED_SCORING, REJECTION_CLAUSES
from rfp_compliance import run_all_checks, check_exclusionary, check_scoring, check_rejection_clauses, check_required_sections
from rfp_structure import get_chapters, PROJECT_TYPES

# 测试用项目信息
SAMPLE_PROJECTS = {
    "goods": {
        "project_name": "上海某医院医疗设备采购项目",
        "project_id": "SH-2026-MED-001",
        "purchaser": "上海市某医院",
        "budget": 5000000,
        "fund_source": "财政资金",
    },
    "services": {
        "project_name": "上海某园区物业管理服务项目",
        "project_id": "SH-2026-PM-001",
        "purchaser": "上海某园区管理委员会",
        "budget": 3000000,
        "fund_source": "财政资金",
    },
    "engineering": {
        "project_name": "某学校校舍焕新改造工程",
        "project_id": "ENG-2026-RN-001",
        "purchaser": "某县教育局",
        "budget": 8000000,
        "fund_source": "财政资金",
    },
}

# 6条合规红线测试用例
RED_LINE_TESTS = [
    {
        "name": "红线1: 资格条件排斥潜在投标人",
        "text": """
        投标人资格要求：
        1. 投标人须为上海市本地企业
        2. 注册资本不低于5000万元
        3. 近三年同类项目业绩金额不低于1000万元
        4. 具有ISO9001质量管理体系认证
        5. 须为大型企业
        """,
        "expected_hits": ["本地企业", "注册资本", "业绩", "ISO", "大型企业"],
    },
    {
        "name": "红线2: 评审标准模糊不可量化",
        "text": """
        评标办法：
        评分因素：
        - 技术方案：优/良/中/差（40分）
        - 服务承诺：优/良/中/差（20分）
        - 企业实力：优/良/中/差（20分）
        - 价格：20分
        合计：100分
        """,
        "expected_issue": "评分标准使用了优/良/中/差，应改为量化打分",
    },
    {
        "name": "红线3: 废标条款缺失",
        "text": """
        第一章 投标邀请
        第二章 投标人须知
        第三章 评标办法
        （全文无废标条款）
        """,
        "expected_issue": "应检测到废标条款缺失",
    },
    {
        "name": "红线4: 等标期不足",
        "text": """
        招标文件获取时间：2026年8月1日至8月5日
        投标截止时间：2026年8月8日
        开标时间：2026年8月8日
        """,
        "expected_issue": "等标期不足20日，应检测到",
    },
    {
        "name": "红线5: 中小企业政策缺失",
        "text": """
        第一章 投标邀请
        第二章 投标人须知
        第三章 资格审查
        第四章 采购需求
        第五章 评标办法
        第六章 合同条款
        第七章 投标文件格式
        （全文无中小企业政策内容）
        """,
        "expected_issue": "政府采购项目应含中小企业政策",
    },
    {
        "name": "红线6: 地域限制",
        "text": """
        投标人资格要求：
        1. 注册地在上海
        2. 须在上海设立分支机构
        3. 仅限上海地区供应商参与
        """,
        "expected_hits": ["注册地", "设立", "仅限"],
    },
]

# ========== 测试1: 占位符统计 ==========
def test_placeholders():
    print("\n" + "=" * 70)
    print("  测试1: 占位符统计（目标：每类≤10个）")
    print("=" * 70)

    results = {}
    for ptype, proj_info in SAMPLE_PROJECTS.items():
        md = generate_markdown(proj_info, ptype)

        # 统计占位符
        placeholders = re.findall(r'【[^】]*】', md)
        # 去重
        unique_placeholders = list(set(placeholders))

        # 统计字符数
        char_count = len(md)

        # 统计非占位符内容比例
        placeholder_chars = sum(len(p) for p in placeholders)
        real_content_chars = char_count - placeholder_chars
        content_ratio = real_content_chars / char_count * 100 if char_count > 0 else 0

        results[ptype] = {
            "total_placeholders": len(placeholders),
            "unique_placeholders": len(unique_placeholders),
            "char_count": char_count,
            "real_content_chars": real_content_chars,
            "content_ratio": round(content_ratio, 1),
            "placeholder_list": unique_placeholders[:20],  # 前20个
        }

        print(f"\n  【{PROJECT_TYPES[ptype]}】")
        print(f"    总字符数: {char_count}")
        print(f"    占位符总数: {len(placeholders)}")
        print(f"    去重占位符: {len(unique_placeholders)}")
        print(f"    实际内容占比: {content_ratio:.1f}%")
        print(f"    占位符列表(前10):")
        for p in unique_placeholders[:10]:
            print(f"      - {p}")

    return results


# ========== 测试2: 核心模板质量 ==========
def test_core_templates():
    print("\n" + "=" * 70)
    print("  测试2: 评分表/废标条款/采购需求模板质量")
    print("=" * 70)

    results = {}

    for ptype, proj_info in SAMPLE_PROJECTS.items():
        md = generate_markdown(proj_info, ptype)
        ptype_results = {}

        # 2a: 评分表
        scoring_section = ""
        if "评分标准" in md:
            start = md.find("评分标准")
            scoring_section = md[start:start+2000]

        scoring_items = DETAILED_SCORING.get(ptype, [])
        total_score = sum(s for _, s, _, _ in scoring_items)

        ptype_results["scoring"] = {
            "items_count": len(scoring_items),
            "total_score": total_score,
            "is_100": total_score == 100,
            "has_price_score": any("报价" in item for item, _, _, _ in scoring_items),
            "has_material_column": all(m for _, _, _, m in scoring_items),
            "sample": scoring_section[:500] if scoring_section else "未找到",
        }

        # 2b: 废标条款
        rejection_section = ""
        if "废标条款" in md or "一票否决" in md:
            for kw in ["废标条款", "一票否决"]:
                idx = md.find(kw)
                if idx >= 0:
                    rejection_section = md[idx:idx+2000]
                    break

        rejection_count = 0
        for cat, items in REJECTION_CLAUSES.items():
            rejection_count += len(items)

        ptype_results["rejection"] = {
            "categories": len(REJECTION_CLAUSES),
            "total_clauses": rejection_count,
            "has_legal_basis": "87号令" in md or "第60条" in md,
            "sample": rejection_section[:500] if rejection_section else "未找到",
        }

        # 2c: 采购需求
        chapters = get_chapters(ptype)
        procurement_chapter = None
        for ch in chapters:
            if ch["title"] == "采购需求":
                procurement_chapter = ch
                break

        ptype_results["procurement"] = {
            "found": procurement_chapter is not None,
            "sections": procurement_chapter["sections"] if procurement_chapter else [],
            "key_fields": procurement_chapter.get("key_fields", []) if procurement_chapter else [],
            "has_content": "【待填写" in md if procurement_chapter else False,
        }

        results[ptype] = ptype_results

        print(f"\n  【{PROJECT_TYPES[ptype]}】")
        print(f"    评分表: {len(scoring_items)}项, 总分={total_score}, 是否100分={'✅' if total_score == 100 else '❌'}")
        print(f"      含价格分: {'✅' if ptype_results['scoring']['has_price_score'] else '❌'}")
        print(f"      含材料列: {'✅' if ptype_results['scoring']['has_material_column'] else '❌'}")
        print(f"    废标条款: {len(REJECTION_CLAUSES)}类{rejection_count}条, 引用87号令: {'✅' if ptype_results['rejection']['has_legal_basis'] else '❌'}")
        print(f"    采购需求: {'✅ 找到' if procurement_chapter else '❌ 缺失'}")
        if procurement_chapter:
            print(f"      子节数: {len(procurement_chapter['sections'])}")
            print(f"      全部为占位符: {'❌ 是' if ptype_results['procurement']['has_content'] else '✅ 有内容'}")

    return results


# ========== 测试3: 三类差异化 ==========
def test_differentiation():
    print("\n" + "=" * 70)
    print("  测试3: 三类差异化（货物/服务/工程结构是否正确区分）")
    print("=" * 70)

    results = {}

    for ptype in ["goods", "services", "engineering"]:
        chapters = get_chapters(ptype)
        chapter_titles = [ch["title"] for ch in chapters]
        total_sections = sum(len(ch["sections"]) for ch in chapters)
        total_key_fields = sum(len(ch.get("key_fields", [])) for ch in chapters)

        # 评分模板差异化
        scoring = SCORING_TEMPLATES.get(ptype, {})

        # 特有章节检查
        has_gov_policy = "政府采购政策落实" in chapter_titles
        has_eng_extras = any("工程量清单" in t or "技术条件" in t or "图纸" in t or "最高投标限价" in t for t in chapter_titles)

        results[ptype] = {
            "chapter_count": len(chapters),
            "chapter_titles": chapter_titles,
            "total_sections": total_sections,
            "total_key_fields": total_key_fields,
            "scoring_template": scoring,
            "has_gov_policy_chapter": has_gov_policy,
            "has_engineering_extras": has_eng_extras,
        }

        print(f"\n  【{PROJECT_TYPES[ptype]}】")
        print(f"    章节数: {len(chapters)}")
        print(f"    子节总数: {total_sections}")
        print(f"    关键字段总数: {total_key_fields}")
        print(f"    评分模板: {scoring}")
        print(f"    政府采购政策章: {'✅' if has_gov_policy else '❌'}")
        print(f"    工程特有章节: {'✅' if has_eng_extras else '—'}")

    # 交叉验证：三类应有差异
    goods_ch = set(results["goods"]["chapter_titles"])
    services_ch = set(results["services"]["chapter_titles"])
    eng_ch = set(results["engineering"]["chapter_titles"])

    diff_1 = services_ch - goods_ch  # 服务类应比货物类多（实际两者相同）
    diff_2 = eng_ch - services_ch    # 工程类应有额外章节
    diff_3 = goods_ch - eng_ch       # 货物类应有政府采购政策，工程类没有

    print(f"\n  差异分析:")
    print(f"    服务类-货物类额外章节: {diff_1 if diff_1 else '无差异'}")
    print(f"    工程类-服务类额外章节: {diff_2 if diff_2 else '无差异'}")
    print(f"    货物类-工程类额外章节: {diff_3 if diff_3 else '无差异'}")

    return results


# ========== 测试4: 合规检查（6条红线） ==========
def test_compliance_redlines():
    print("\n" + "=" * 70)
    print("  测试4: 6条合规红线检出能力")
    print("=" * 70)

    results = {}

    for test in RED_LINE_TESTS:
        text = test["text"]
        ptype = "services"

        # 运行全部检查
        report = run_all_checks(text, ptype)

        # 提取排他性条款检测结果
        exclusionary = check_exclusionary(text)
        rejection = check_rejection_clauses(text)
        required_sections = check_required_sections(text, ptype)

        # 统计检出情况
        ex_hits = len(exclusionary)
        ex_reasons = [e["reason"] for e in exclusionary]
        rejection_fail = any(r["severity"] == "fail" for r in rejection)
        required_fail = any(r["severity"] == "fail" for r in required_sections)

        # 红线4: 等标期检测（当前compliance模块没有这个检查）
        has_time_check = any(t["type"] == "时间节点" for t in report["checks"]["time_nodes"])

        # 红线5: 中小企业政策检测
        has_sme_check = "政府采购政策落实" in text or "中小企业" in text

        results[test["name"]] = {
            "exclusionary_hits": ex_hits,
            "exclusionary_reasons": ex_reasons[:5],
            "rejection_detected_missing": rejection_fail,
            "required_sections_missing": sum(1 for r in required_sections if r["severity"] == "fail"),
            "time_nodes_warns": sum(1 for t in report["checks"]["time_nodes"] if t["severity"] == "warn"),
            "overall_fail": report["summary"]["fail"],
            "overall_warn": report["summary"]["warn"],
            "overall_pass": report["summary"]["pass"],
            "verdict": report["summary"]["verdict"],
        }

        print(f"\n  {test['name']}")
        print(f"    排他性条款检出: {ex_hits}处")
        if ex_reasons:
            for r in ex_reasons[:3]:
                print(f"      → {r}")
        print(f"    废标条款缺失检测: {'✅ 检出' if rejection_fail else '❌ 未检出'}")
        print(f"    必备子节缺失: {sum(1 for r in required_sections if r['severity'] == 'fail')}个")
        print(f"    时间节点告警: {sum(1 for t in report['checks']['time_nodes'] if t['severity'] == 'warn')}个")
        print(f"    总体: ✅{report['summary']['pass']} ⚠️{report['summary']['warn']} ❌{report['summary']['fail']} → {report['summary']['verdict']}")

    # 额外测试：对生成器自身的输出跑合规检查
    print(f"\n  --- 对rfp_generator生成结果跑合规检查 ---")
    for ptype, proj_info in SAMPLE_PROJECTS.items():
        md = generate_markdown(proj_info, ptype)
        report = run_all_checks(md, ptype)
        s = report["summary"]
        print(f"\n  【{PROJECT_TYPES[ptype]}】生成文件合规检查:")
        print(f"    文本长度: {len(md)}字符")
        print(f"    ✅{s['pass']} ⚠️{s['warn']} ❌{s['fail']} → {s['verdict']}")
        # 列出fail项
        for check_name, items in report["checks"].items():
            if isinstance(items, list):
                for item in items:
                    if item.get("severity") == "fail":
                        print(f"    ❌ [{check_name}] {item.get('message', '')}")

    return results


# ========== 测试5: 真实文件对比 ==========
def test_real_file_comparison():
    print("\n" + "=" * 70)
    print("  测试5: 真实招标文件对比（从样本池抽样）")
    print("=" * 70)

    import subprocess

    # 选代表性文件（每类2-3个，去重后取第一个版本）
    sample_files = {
        "services": [
            "/tmp/sample-bids/服务类/上海华铁列车库内保洁服务_招标_2025.docx",
            "/tmp/sample-bids/服务类/南大地区基地物业管理服务_招标_2026.docx",
            "/tmp/sample-bids/服务类/厦门国际银行零售业务劳务外包_招标_2026.docx",
        ],
        "goods": [
            "/tmp/sample-bids/货物类/喀什地区维吾尔医医院医疗设备采购_公开招标_2026.docx",
            "/tmp/sample-bids/货物类/昆明海关_手持式X荧光光谱仪_招标公告.docx",
        ],
        "engineering": [
            "/tmp/sample-bids/工程类/银川国芳外墙logo字制作安装_招标_2026.docx",
            "/tmp/sample-bids/工程类/工程项目供应商目录建立_公开招标_2026.docx",
        ],
    }

    results = {}

    for ptype, files in sample_files.items():
        print(f"\n  【{PROJECT_TYPES[ptype]}】")
        ptype_results = []

        for fpath in files:
            if not os.path.exists(fpath):
                print(f"    ⚠️ 文件不存在: {fpath}")
                continue

            fname = os.path.basename(fpath)

            # 解析docx
            text = ""
            try:
                from docx import Document
                doc = Document(fpath)
                text = '\n'.join(p.text for p in doc.paragraphs)
                # 也提取表格
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text += '\n' + cell.text
            except Exception as e:
                # 尝试PDF
                try:
                    import fitz
                    doc = fitz.open(fpath)
                    text = '\n'.join(page.get_text() for page in doc)
                except Exception as e2:
                    print(f"    ❌ 解析失败 {fname}: {e}")
                    continue

            if not text.strip():
                print(f"    ⚠️ 文件为空: {fname}")
                continue

            # 统计
            char_count = len(text)
            # 找章节标题
            chapter_pattern = re.findall(r'第[一二三四五六七八九十\d]+章[^\n]{0,30}', text)
            chapters_found = list(set(chapter_pattern))[:15]

            # 找评分表
            has_scoring = "评分" in text or "评标" in text or "评审" in text
            # 找废标条款
            has_rejection = any(kw in text for kw in ["废标", "否决投标", "无效投标", "投标无效"])
            # 找前附表
            has_front_sheet = "前附表" in text or "须知前附表" in text
            # 找资格要求
            has_qualification = "资格" in text and ("要求" in text or "条件" in text)
            # 找采购需求
            has_procurement = "采购需求" in text or "技术要求" in text or "服务要求" in text

            # 跑合规检查
            report = run_all_checks(text, ptype)
            s = report["summary"]

            file_result = {
                "filename": fname,
                "char_count": char_count,
                "chapters_found": chapters_found,
                "has_scoring": has_scoring,
                "has_rejection": has_rejection,
                "has_front_sheet": has_front_sheet,
                "has_qualification": has_qualification,
                "has_procurement": has_procurement,
                "compliance": {
                    "pass": s["pass"],
                    "warn": s["warn"],
                    "fail": s["fail"],
                    "verdict": s["verdict"],
                },
            }
            ptype_results.append(file_result)

            print(f"    📄 {fname[:40]}...")
            print(f"       字符数: {char_count} (vs 生成器: ~{len(generate_markdown(SAMPLE_PROJECTS[ptype], ptype))})")
            print(f"       章节数: {len(chapters_found)}")
            print(f"       评分表: {'✅' if has_scoring else '❌'} 废标: {'✅' if has_rejection else '❌'} 前附表: {'✅' if has_front_sheet else '❌'}")
            print(f"       合规: ✅{s['pass']} ⚠️{s['warn']} ❌{s['fail']} → {s['verdict']}")

        results[ptype] = ptype_results

    return results


# ========== 主函数 ==========
def main():
    print("=" * 70)
    print("  RFP Generator 系统性测试报告")
    print("  测试时间: 2026-07-25")
    print("  测试样本: 34份招标文件（去重后约30份）")
    print("=" * 70)

    all_results = {}

    # 测试1: 占位符
    all_results["placeholders"] = test_placeholders()

    # 测试2: 核心模板
    all_results["core_templates"] = test_core_templates()

    # 测试3: 三类差异化
    all_results["differentiation"] = test_differentiation()

    # 测试4: 合规红线
    all_results["compliance"] = test_compliance_redlines()

    # 测试5: 真实文件对比
    all_results["real_comparison"] = test_real_file_comparison()

    # 保存完整结果
    output_path = "/tmp/bid-toolkit/rfp/test_results.json"
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(all_results, f, ensure_ascii=False, indent=2, default=str)
    print(f"\n\n✅ 完整测试结果已保存: {output_path}")

    # 输出总结
    print("\n" + "=" * 70)
    print("  测试总结")
    print("=" * 70)

    # 占位符总结
    print("\n📊 占位符统计:")
    for ptype, data in all_results["placeholders"].items():
        status = "✅ 达标" if data["total_placeholders"] <= 10 else "❌ 超标"
        print(f"  {PROJECT_TYPES[ptype]}: {data['total_placeholders']}个占位符 / {data['char_count']}字符 → {status}")

    # 核心模板总结
    print("\n📊 核心模板:")
    for ptype, data in all_results["core_templates"].items():
        s_ok = "✅" if data["scoring"]["is_100"] else "❌"
        r_ok = "✅" if data["rejection"]["has_legal_basis"] else "❌"
        p_ok = "✅" if data["procurement"]["found"] and not data["procurement"]["has_content"] else "❌"
        print(f"  {PROJECT_TYPES[ptype]}: 评分表{s_ok} 废标条款{r_ok} 采购需求{p_ok}")

    # 差异化总结
    print("\n📊 三类差异化:")
    diff = all_results["differentiation"]
    for ptype in ["goods", "services", "engineering"]:
        d = diff[ptype]
        print(f"  {PROJECT_TYPES[ptype]}: {d['chapter_count']}章 / {d['total_sections']}节 / 评分={d['scoring_template']}")

    # 合规红线总结
    print("\n📊 6条合规红线检出:")
    for name, data in all_results["compliance"].items():
        if "exclusionary_hits" in data:
            print(f"  {name}: 排他{data['exclusionary_hits']}处, 废标缺失{'✅' if data['rejection_detected_missing'] else '❌'}")

    print("\n" + "=" * 70)
    print("  测试完成")
    print("=" * 70)


if __name__ == '__main__':
    main()
