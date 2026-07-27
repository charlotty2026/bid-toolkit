#!/usr/bin/env python3
"""
招标文件生成器 v2.0
v2.0 变更（第二轮审核 P0 补完）：
  - 前附表自动生成表格（不再用占位符）
  - 政府采购政策落实章节自动生成政策内容
  - 保证金/澄清/质疑/资格审查等新节加法律依据提示
  - 章节ID自动适配 rfp_structure v2.0 的新编号
从项目信息自动生成标准招标文件（Markdown + Word）。

用法：
  python rfp_generator.py --type services --project "XX后勤服务项目" --budget 500000
  python rfp_generator.py --type services --project "XX项目" --budget 500000 --docx
  python rfp_generator.py --config project.json -o 招标文件.md
"""

import argparse
import json
import os
import sys
from datetime import datetime, timedelta

# 导入标准结构
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from rfp_structure import (
    PROJECT_TYPES, STANDARD_CHAPTERS, ENGINEERING_EXTRA_CHAPTERS,
    LEGAL_BASIS, get_chapters,
)

import yaml as _yaml_module

# ===== 评分标准模板（默认值，可通过 --scoring-config 加载外部YAML覆盖）=====
# 格式：{ 项目类型: [ {name, score, detail}, ... ] }
# 总分必须=100，否则合规检查器会报错

def _validate_scoring(scoring_dict):
    """校验评分模板合法性：必须是dict，每个类型的分值合计必须=100"""
    if not isinstance(scoring_dict, dict):
        raise ValueError(f"评分模板必须是字典格式，实际类型：{type(scoring_dict).__name__}")
    for ptype, items in scoring_dict.items():
        if not isinstance(items, list):
            raise ValueError(f"评分模板 '{ptype}' 的值必须是列表，实际类型：{type(items).__name__}")
        total = sum(item.get('score', 0) for item in items)
        if total != 100:
            raise ValueError(f"评分模板 '{ptype}' 总分={total}，必须等于100")

def _load_default_scoring():
    """从 templates/scoring.yaml 加载默认评分模板"""
    scoring_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', 'templates', 'scoring.yaml')
    if os.path.exists(scoring_path):
        try:
            with open(scoring_path, 'r', encoding='utf-8') as f:
                data = _yaml_module.safe_load(f)
            if data and 'default' in data:
                _validate_scoring(data['default'])
                return data['default']
        except Exception as e:
            print(f"⚠️ 加载默认评分模板失败（{e}），使用内置硬编码兜底")
    # 兜底硬编码
    return {
        "goods": [
            {"name": "价格分", "score": 30, "detail": "最低有效报价得满分，其他按公式计算"},
            {"name": "技术分", "score": 50, "detail": "技术参数响应、方案质量、技术能力"},
            {"name": "商务分", "score": 20, "detail": "资质、业绩、信誉、售后服务"},
        ],
        "services": [
            {"name": "价格分", "score": 10, "detail": "最低有效报价得满分，其他按公式计算"},
            {"name": "技术分", "score": 60, "detail": "服务方案、人员配置、管理制度、应急预案"},
            {"name": "商务分", "score": 30, "detail": "资质、业绩、信誉、本地化服务能力"},
        ],
        "engineering": [
            {"name": "价格分", "score": 20, "detail": "最低有效报价得满分，其他按公式计算"},
            {"name": "技术分", "score": 60, "detail": "施工组织设计、施工方案、安全文明施工"},
            {"name": "商务分", "score": 20, "detail": "资质、业绩、信誉、项目管理人员"},
        ],
    }

SCORING_TEMPLATES = _load_default_scoring()

def load_scoring_config(config_path):
    """从外部YAML文件加载自定义评分模板
    
    YAML格式与 templates/scoring.yaml 相同。
    支持两种写法：
    1. 顶层直接是类型->列表（覆盖对应类型）
    2. 用 default/custom_example 分组（取 default 或指定组名）
    
    Returns: dict，格式同 SCORING_TEMPLATES
    
    Raises: FileNotFoundError, ValueError
    """
    if not os.path.exists(config_path):
        raise FileNotFoundError(f"评分配置文件不存在：{config_path}")
    
    try:
        with open(config_path, 'r', encoding='utf-8') as f:
            data = _yaml_module.safe_load(f)
    except _yaml_module.YAMLError as e:
        raise ValueError(f"YAML格式错误：{e}")
    
    if not data:
        raise ValueError(f"评分配置文件为空：{config_path}")
    
    if not isinstance(data, dict):
        raise ValueError(f"评分配置必须是字典格式，实际类型：{type(data).__name__}")
    
    # 如果有 'default' 键，取 default 子字典
    if 'default' in data:
        result = data['default']
    else:
        # 否则直接用顶层（过滤掉非类型键如 price_ranges）
        valid_types = {'goods', 'services', 'engineering'}
        result = {k: v for k, v in data.items() if k in valid_types}
    
    if not result:
        raise ValueError(f"评分配置中未找到有效的项目类型（goods/services/engineering）")
    
    _validate_scoring(result)
    return result

# P0 新节法律依据提示
SECTION_HINTS = {
    "招标文件的澄清与修改": "法律依据：87号令第27条，采购人可在投标截止15个工作日前发出澄清/修改",
    "投标保证金": "法律依据：87号令第33条，不超过预算金额的2%，且不得超过法定上限",
    "质疑与投诉": "法律依据：政府采购法第52-58条，质疑答复时限7个工作日，投诉处理时限30个工作日",
    "资格审查方式（资格预审/资格后审）": "法律依据：87号令第20条，资格审查由采购人或采购代理机构依法开展",
}

def generate_front_sheet(project_info, project_type):
    """P0-2: 生成投标人须知前附表（自动填充已知信息）"""
    rows = [
        ("1", "项目名称", str(project_info.get("project_name", "【待填写】"))),
        ("2", "项目编号", str(project_info.get("project_id", "【待填写】"))),
        ("3", "采购人", str(project_info.get("purchaser", "【待填写】"))),
        ("4", "预算金额", f"{project_info.get('budget', '【待填写】')}元"),
        ("5", "项目类型", PROJECT_TYPES.get(project_type, "服务类")),
        ("6", "投标截止时间", "【待填写：YYYY-MM-DD HH:MM】"),
        ("7", "开标地点", "【待填写】"),
        ("8", "投标有效期", "【待填写：自投标截止日起XX日】"),
        ("9", "投标保证金", "不超过预算金额的2%（87号令第33条）"),
        ("10", "投标文件份数", "正本1份，副本【待填写】份"),
        ("11", "评标方法", "综合评分法"),
        ("12", "资金来源", str(project_info.get("fund_source", "【待填写】"))),
    ]
    lines = ["| 序号 | 条款名称 | 内容 |", "|------|---------|------|"]
    for row in rows:
        lines.append(f"| {row[0]} | {row[1]} | {row[2]} |")
    lines.append("")
    lines.append("> 前附表与正文条款不一致时，以前附表为准。")
    return "\n".join(lines)

def generate_gov_policy_content():
    """P0-6: 生成政府采购政策落实章节内容"""
    policies = [
        ("促进中小企业发展",
         "对小型、微型企业产品的价格给予 **6%-10%的扣除**，用扣除后的价格参与评审。",
         "《政府采购促进中小企业发展管理办法》"),
        ("监狱企业扶持政策",
         "监狱企业视同小型、微型企业，享受相同的价格扣除政策。",
         "《关于政府采购支持监狱企业发展有关问题的通知》"),
        ("残疾人福利性单位政策",
         "残疾人福利性单位视同小型、微型企业，享受相同的价格扣除政策。",
         "《关于促进残疾人就业政府采购政策的通知》"),
        ("节能产品政府采购",
         "属于强制采购节能产品范围的，应当采购列入 **节能产品政府采购清单** 的产品。",
         "《关于调整节能产品政府采购清单的通知》"),
        ("环境标志产品政府采购",
         "属于环境标志产品政府采购清单中的产品，在同等条件下 **优先采购**。",
         "《关于调整环境标志产品政府采购清单的通知》"),
    ]
    lines = ["本项目落实以下政府采购政策：\n"]
    for title, content, basis in policies:
        lines.append(f"### {title}\n")
        lines.append(f"{content}\n")
        lines.append(f"> 法律依据：{basis}\n")
    return "\n".join(lines)

def generate_markdown(project_info, project_type="services", scoring_override=None):
    """生成Markdown格式招标文件
    
    Args:
        project_info: 项目信息字典
        project_type: 项目类型
        scoring_override: 自定义评分模板，格式 {类型: [{name, score, detail}, ...]}
                     传入则覆盖默认评分模板
    """
    chapters = get_chapters(project_type)
    type_name = PROJECT_TYPES.get(project_type, "服务类")
    
    # 评分模板：优先用传入的 > 默认
    if scoring_override and project_type in scoring_override:
        scoring_items = scoring_override[project_type]
    else:
        scoring_items = SCORING_TEMPLATES.get(project_type, SCORING_TEMPLATES["services"])

    lines = []
    # 文件头
    lines.append(f"# {project_info.get('project_name', '【项目名称】')}招标文件\n")
    lines.append(f"> 项目类型：{type_name}")
    lines.append(f"> 项目编号：{project_info.get('project_id', '【项目编号】')}")
    lines.append(f"> 采购人：{project_info.get('purchaser', '【采购人】')}")
    lines.append(f"> 预算金额：{project_info.get('budget', '【预算金额】')}元")
    lines.append(f"> 生成时间：{datetime.now().strftime('%Y-%m-%d')}\n")
    lines.append("---\n")

    # 法律法规依据
    lines.append("## 编制依据\n")
    for law_name, law_info in LEGAL_BASIS.items():
        lines.append(f"- **{law_info['full_name']}**（{law_info['effective_date']}起施行）")
        for art in law_info["key_articles"]:
            lines.append(f"  - {art}")
    lines.append("")

    # 各章内容
    for ch in chapters:
        lines.append(f"\n---\n")
        lines.append(f"# 第{ch['id']}章 {ch['title']}\n")
        lines.append(f"> {ch.get('说明', '')}\n")

        for section in ch["sections"]:
            lines.append(f"## {section}\n")

            # P0-2: 投标人须知前附表 -> 自动生成表格
            if section == "投标人须知前附表":
                lines.append(generate_front_sheet(project_info, project_type))
                lines.append("")
                continue

            # P0-6: 政府采购政策落实 -> 自动生成政策内容
            if ch["title"] == "政府采购政策落实":
                lines.append(generate_gov_policy_content())
                lines.append("")
                continue

            # 其他P0新节加法律依据提示
            lines.append(f"【待填写：{section}具体内容】\n")
            hint = SECTION_HINTS.get(section)
            if hint:
                lines.append(f"> {hint}\n")

        # 关键字段提示
        if ch.get("key_fields"):
            lines.append("### 关键信息字段\n")
            for field in ch["key_fields"]:
                lines.append(f"- [ ] {field}")
            lines.append("")

    # 评分标准
    lines.append("\n---\n")
    lines.append("# 附件：评分标准（参考）\n")
    lines.append(f"> 项目类型：{type_name}，总分100分\n")
    lines.append("| 评分因素 | 分值 | 评分细则 |")
    lines.append("|---------|------|----------|")
    for item in scoring_items:
        name = item.get('name', '未知')
        score = item.get('score', 0)
        detail = item.get('detail', '【待填写评分细则】')
        lines.append(f"| {name} | {score} | {detail} |")
    lines.append(f"| **合计** | **{sum(item.get('score', 0) for item in scoring_items)}** | |")

    # 废标条款
    lines.append("\n---\n")
    lines.append("# 附件：废标条款（⚠️ 投标人必读）\n")
    lines.append("以下情形之一的，按废标处理：\n")
    standard_rejection = [
        "投标文件未按招标文件要求密封的",
        "投标文件未按要求签章的",
        "投标报价超过最高投标限价的",
        "投标有效期不足的",
        "投标人未提交投标保证金的（如要求）",
        "投标文件存在重大偏差的",
        "投标人不符合资格要求的",
        "存在串通投标行为的",
    ]
    for i, item in enumerate(standard_rejection, 1):
        lines.append(f"{i}. ⚠️ {item}")
    lines.append(f"\n> 法律依据：{LEGAL_BASIS['财政部87号令']['full_name']}第60条")

    return "\n".join(lines)


def generate_docx(markdown_text, output_path):
    """将Markdown转为Word文档，标题用原生Heading样式，插入TOC域"""
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        print("错误：需要python-docx库，请运行 pip install python-docx")
        return False

    doc = Document()

    # 页面设置
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # 初始化样式
    style_normal = doc.styles['Normal']
    style_normal.font.name = '宋体'
    style_normal.font.size = Pt(12)
    style_normal.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    for level in range(1, 5):
        style = doc.styles[f'Heading {level}']
        style.font.name = '黑体'
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        style.font.color.rgb = RGBColor(0, 0, 0)

    # 先插入TOC域（占位，打开Word后F9更新）
    _insert_toc_field(doc)

    # 解析Markdown并写入
    lines = markdown_text.split('\n')
    in_table = False
    table_rows = []

    for line in lines:
        stripped = line.strip()

        # 跳过空行
        if not stripped:
            if in_table and table_rows:
                _write_table(doc, table_rows)
                table_rows = []
                in_table = False
            continue

        # 表格行
        if stripped.startswith('|'):
            in_table = True
            if not stripped.startswith('|--') and not stripped.startswith('| -'):
                cells = [c.strip() for c in stripped.split('|')[1:-1]]
                table_rows.append(cells)
            continue
        elif in_table and table_rows:
            _write_table(doc, table_rows)
            table_rows = []
            in_table = False

        # 标题
        if stripped.startswith('# ') and not stripped.startswith('## '):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith('## '):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith('### '):
            doc.add_heading(stripped[4:], level=3)
        # 分隔线
        elif stripped == '---':
            doc.add_page_break()
        # 引用
        elif stripped.startswith('> '):
            p = doc.add_paragraph()
            run = p.add_run(stripped[2:])
            run.italic = True
            run.font.size = Pt(10)
        # 列表
        elif stripped.startswith('- '):
            doc.add_paragraph(stripped[2:], style='List Bullet')
        elif stripped.startswith('- [ ] '):
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(f"☐ {stripped[6:]}")
        elif stripped[0:2].rstrip('.').isdigit() and '. ' in stripped[:5]:
            idx = stripped.index('. ')
            doc.add_paragraph(stripped[idx+2:], style='List Number')
        else:
            doc.add_paragraph(stripped)

    # 收尾：如果还有表格没写
    if in_table and table_rows:
        _write_table(doc, table_rows)

    doc.save(output_path)
    return True


def _insert_toc_field(doc):
    """在文档开头插入Word目录域代码。打开Word后按Ctrl+A -> F9更新域即可生成目录。"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    
    body = doc.element.body
    first = body[0] if len(body) > 0 else None

    # 目录标题段落（用Heading 1样式）
    toc_title = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), 'Heading1')
    pPr.append(pStyle)
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    pPr.append(jc)
    toc_title.append(pPr)
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = '目  录'
    r.append(t)
    toc_title.append(r)

    # TOC域段落
    toc_para = OxmlElement('w:p')
    r1 = OxmlElement('w:r')
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    r1.append(fld_begin)
    toc_para.append(r1)

    r2 = OxmlElement('w:r')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    r2.append(instr)
    toc_para.append(r2)

    r3 = OxmlElement('w:r')
    fld_sep = OxmlElement('w:fldChar')
    fld_sep.set(qn('w:fldCharType'), 'separate')
    r3.append(fld_sep)
    toc_para.append(r3)

    r4 = OxmlElement('w:r')
    t4 = OxmlElement('w:t')
    t4.text = '（打开Word后按 Ctrl+A 全选 -> F9 更新域，目录自动生成）'
    r4.append(t4)
    toc_para.append(r4)

    r5 = OxmlElement('w:r')
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    r5.append(fld_end)
    toc_para.append(r5)

    # 分页符
    page_break = OxmlElement('w:p')
    r_pb = OxmlElement('w:r')
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    r_pb.append(br)
    page_break.append(r_pb)

    if first is not None:
        body.insert(0, page_break)
        body.insert(0, toc_para)
        body.insert(0, toc_title)
    else:
        body.append(toc_title)
        body.append(toc_para)
        body.append(page_break)


def _write_table(doc, rows):
    """将收集的表格行写入Word表格"""
    if not rows:
        return
    from docx.shared import Pt
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Table Grid'
    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j < len(table.rows[i].cells):
                cell = table.rows[i].cells[j]
                cell.text = cell_text
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(10)


def main():
    parser = argparse.ArgumentParser(description='招标文件生成器')
    parser.add_argument('--type', choices=['goods', 'services', 'engineering'],
                        default='services', help='项目类型（默认services）')
    parser.add_argument('--project', help='项目名称')
    parser.add_argument('--project-id', help='项目编号')
    parser.add_argument('--purchaser', help='采购人')
    parser.add_argument('--budget', type=int, help='预算金额（元，纯数字）')
    parser.add_argument('--config', help='从JSON文件读取项目信息')
    parser.add_argument('--scoring-config', dest='scoring_config',
                        help='从YAML文件加载自定义评分模板（覆盖默认评分维度/分值/细则）')
    parser.add_argument('-o', '--output', default='招标文件.md', help='输出文件名')
    parser.add_argument('--docx', action='store_true', help='同时生成Word文档')
    args = parser.parse_args()

    # 收集项目信息
    if args.config:
        with open(args.config, 'r', encoding='utf-8') as f:
            project_info = json.load(f)
    else:
        project_info = {
            'project_name': args.project or '',
            'project_id': args.project_id or '',
            'purchaser': args.purchaser or '',
            'budget': args.budget or '',
        }

    # 加载自定义评分模板（可选）
    scoring_override = None
    if args.scoring_config:
        try:
            scoring_override = load_scoring_config(args.scoring_config)
            print(f"✅ 已加载自定义评分模板：{args.scoring_config}")
            for ptype, items in scoring_override.items():
                total = sum(item.get('score', 0) for item in items)
                print(f"   {ptype}: {len(items)}个评分维度，总分{total}")
        except (FileNotFoundError, ValueError) as e:
            print(f"❌ 加载评分模板失败：{e}")
            sys.exit(1)

    # 生成Markdown
    md_text = generate_markdown(project_info, args.type, scoring_override)

    # 扩展名校验：不带 --docx 时输出内容就是 Markdown，强制 .md 后缀防误导
    if not args.docx and args.output.lower().endswith('.docx'):
        args.output = args.output.rsplit('.', 1)[0] + '.md'
        print(f"⚠️ 未加 --docx，输出内容为 Markdown，文件名已自动改为 {args.output}")

    # 写入文件
    with open(args.output, 'w', encoding='utf-8') as f:
        f.write(md_text)
    print(f"✅ Markdown已生成：{args.output}")

    # 可选：生成Word
    if args.docx:
        docx_path = args.output.rsplit('.', 1)[0] + '.docx'
        if generate_docx(md_text, docx_path):
            print(f"✅ Word已生成：{docx_path}")
            print(f"   ⚠️ 打开Word后按 Ctrl+A 全选 -> F9 更新目录域")


if __name__ == '__main__':
    main()
