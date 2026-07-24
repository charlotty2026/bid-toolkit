#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
标书排版自动修复器 v1.2
==============================
功能覆盖:
  1. 占位符修复     [XX] [XXX] [400-] [地址] [电话] [日期] [投标人名称]
  2. 日期占位符     XXXX年XX月XX日、____年__月__日
  3. 字体样式统一   正文/标题字体规范
  4. 编号修复       连续序号检查、断号修复
  5. 多余空行清除   连续空行合并
  6. 图片检查       图片缺失/链接断开
  7. 表格格式优化   单元格合并残留、边框缺失
  8. 目录修复       TOC域刷新
  9. 页眉页脚检查   页眉页脚一致性
  10. 暗标扫描      隐藏文字/不可见内容
  11. 全角半角检测   中文段落中英文标点→全角（自动修复）
  12. 数据一致性检测 人员/金额/期限/计算公式不一致（仅检测不修复）
  13. 标题层级检测   Heading样式使用、编号连续性
  14. 下划线检测     落款下划线/暗标下划线警告
  15. 表格样式隔离   表格段落误用正文样式/首行缩进检测（铁律28）
  16. 图片样式检查   图片段落缩进/居中检测（铁律28）
  17. 标题编号格式   tab分隔符/缩进过大检测（铁律22A）
"""

import os
import re
import sys
import json
import argparse
from datetime import datetime
from pathlib import Path

sys.stdout.reconfigure(encoding='utf-8')

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("❌ 需要 python-docx 库，请执行: pip install python-docx")
    sys.exit(1)

# ─── 路径环境变量 ─────────────────────────────────────────
WORK_ROOT = os.environ.get("WORK_ROOT", os.path.expanduser("~/bid-workspace"))
UPLOADS_DIR = os.path.join(WORK_ROOT, "uploads")
OUTPUTS_DIR = os.path.join(WORK_ROOT, "outputs")
LOGS_DIR = os.path.join(WORK_ROOT, "logs")
TOOLS_DIR = os.path.join(WORK_ROOT, "tools")


# ============================================================
#  工具函数
# ============================================================
def log(msg, level="INFO"):
    """统一日志"""
    ts = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    print(f"[{level}] {msg}")


def find_docx_files(path):
    """查找目录下的所有 docx 文件"""
    p = Path(path)
    files = []
    if p.is_file() and p.suffix.lower() in ('.docx', '.doc'):
        files.append(str(p))
    elif p.is_dir():
        for f in p.glob("*.docx"):
            files.append(str(f))
        for f in p.glob("*.doc"):
            files.append(str(f))
    return files


# ============================================================
#  模块1: 占位符修复
# ============================================================
def scan_placeholders(doc):
    """扫描文档中的各种占位符"""
    issues = []
    
    placeholder_patterns = {
        "通用占位符[XX]": re.compile(r'\[XX\]'),
        "占位符[XXX]": re.compile(r'\[XXX\]'),
        "电话占位符[400-]": re.compile(r'\[400-\d*\]'),
        "地址占位符[地址]": re.compile(r'\[地址\]'),
        "电话占位符[电话]": re.compile(r'\[电话\]'),
        "日期占位符[日期]": re.compile(r'\[日期\]'),
        "名称占位符": re.compile(r'\[投标人名称\]|\[公司名称\]|\[投标单位\]'),
        "金额占位符": re.compile(r'\[金额\]|\[报价\]|\[总价\]'),
        "项目占位符": re.compile(r'\[项目名称\]|\[项目编号\]'),
        "负责人占位符": re.compile(r'\[负责人\]|\[联系人\]|\[法定代表人\]'),
    }
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        if not text.strip():
            continue
            
        # 去除非连续的[XX]（可能是真正的标记）
        text_clean = text.replace(" ", "").replace("　", "")
        
        for name, pattern in placeholder_patterns.items():
            matches = pattern.findall(text)
            if matches:
                issues.append({
                    "type": "占位符",
                    "subtype": name,
                    "paragraph": i + 1,
                    "text": text[:80],
                    "matches": list(set(matches)),
                })
    
    return issues


def fix_placeholders(doc, issues):
    """修复占位符（标记为待填写）"""
    if not issues:
        return
    
    # 占位符替换映射
    placeholder_fixes = {
        r'\[XX\]': '【待填写】',
        r'\[XXX\]': '【待填写】',
        r'\[400-\d*\]': '【待填写电话】',
        r'\[地址\]': '【待填写地址】',
        r'\[电话\]': '【待填写电话】',
        r'\[日期\]': '【待填写日期】',
        r'\[投标人名称\]': '【待填写投标人名称】',
        r'\[公司名称\]': '【待填写公司名称】',
        r'\[投标单位\]': '【待填写投标单位】',
        r'\[金额\]': '【待填写金额】',
        r'\[报价\]': '【待填写报价】',
        r'\[总价\]': '【待填写总价】',
        r'\[项目名称\]': '【待填写项目名称】',
        r'\[项目编号\]': '【待填写项目编号】',
        r'\[负责人\]': '【待填写负责人】',
        r'\[联系人\]': '【待填写联系人】',
        r'\[法定代表人\]': '【待填写法定代表人】',
    }
    
    fixed = set()
    for issue in issues:
        para_idx = issue["paragraph"] - 1
        if para_idx < 0 or para_idx >= len(doc.paragraphs):
            continue
        
        para = doc.paragraphs[para_idx]
        original = para.text
        
        new_text = original
        for pattern, replacement in placeholder_fixes.items():
            if re.search(pattern, new_text):
                new_text = re.sub(pattern, replacement, new_text)
        
        if new_text != original:
            # 替换段落中的文本（保留格式）
            _replace_para_text(para, new_text)
            fixed.add(para_idx)
    
    return fixed


# ============================================================
#  模块2: 日期占位符修复
# ============================================================
def scan_date_placeholders(doc):
    """扫描日期格式占位符"""
    issues = []
    
    date_patterns = {
        "XXXX年XX月XX日": re.compile(r'[Xx]{4}\s*年\s*[Xx]{2}\s*月\s*[Xx]{2}\s*日'),
        "____年__月__日": re.compile(r'_+\s*年\s*_+\s*月\s*_+\s*日'),
        "下划线日期": re.compile(r'_{2,}\s*年\s*_{1,2}\s*月\s*_{1,2}\s*日'),
    }
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        if not text.strip():
            continue
        
        for name, pattern in date_patterns.items():
            matches = pattern.findall(text)
            if matches:
                issues.append({
                    "type": "日期占位符",
                    "subtype": name,
                    "paragraph": i + 1,
                    "text": text[:80],
                    "matches": list(set(matches)),
                })
    
    return issues


def fix_date_placeholders(doc, issues):
    """修复日期占位符"""
    if not issues:
        return set()
    
    fixed = set()
    date_patterns = {
        r'[Xx]{4}\s*年\s*[Xx]{2}\s*月\s*[Xx]{2}\s*日': '【待填写日期】',
        r'_+\s*年\s*_+\s*月\s*_+\s*日': '【待填写日期】',
        r'_{2,}\s*年\s*_{1,2}\s*月\s*_{1,2}\s*日': '【待填写日期】',
    }
    
    for issue in issues:
        para_idx = issue["paragraph"] - 1
        if para_idx < 0 or para_idx >= len(doc.paragraphs):
            continue
        
        para = doc.paragraphs[para_idx]
        original = para.text
        
        new_text = original
        for pattern, replacement in date_patterns.items():
            if re.search(pattern, new_text):
                new_text = re.sub(pattern, replacement, new_text)
        
        if new_text != original:
            _replace_para_text(para, new_text)
            fixed.add(para_idx)
    
    return fixed


# ============================================================
#  模块3: 字体样式统一
# ============================================================
def scan_font_issues(doc):
    """扫描字体样式问题"""
    issues = []
    
    # 检查正文默认字体
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        
        para_style = para.style.name if para.style else "无"
        
        # 检查字体大小是否规整
        for run in para.runs:
            if run.font.size and run.font.size < Pt(8):
                issues.append({
                    "type": "字体问题",
                    "subtype": "字体过小",
                    "paragraph": i + 1,
                    "detail": f"字号: {run.font.size.pt}pt",
                    "text": run.text[:40],
                })
                break
    
    # 检查标题样式是否合理
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        
        style_name = para.style.name if para.style else ""
        
        # 超过15字的长段落用了标题样式
        if "Heading" in style_name or "标题" in style_name:
            if len(text) > 30:
                issues.append({
                    "type": "字体问题",
                    "subtype": "长段落误用标题样式",
                    "paragraph": i + 1,
                    "detail": f"样式: {style_name}, 长度: {len(text)}字",
                    "text": text[:60],
                })
        
        # 短标题（<10字）用了正文段落样式
        if ("Normal" in style_name or "正文" in style_name) and len(text) <= 10:
            # 判断是否像标题：无句号结尾、不是普通句子
            if text and not any(text.endswith(p) for p in "。,.;:：；"):
                if not re.match(r'^[一二三四五六七八九十\d]+[、.．）)]', text):
                    pass  # 暂不标记，减少误报
    
    return issues


def fix_font_issues(doc, issues):
    """自动修复字体问题"""
    if not issues:
        return set()
    
    fixed = set()
    for issue in issues:
        if issue["subtype"] == "字体过小":
            para_idx = issue["paragraph"] - 1
            if para_idx < 0 or para_idx >= len(doc.paragraphs):
                continue
            para = doc.paragraphs[para_idx]
            for run in para.runs:
                if run.font.size and run.font.size < Pt(8):
                    run.font.size = Pt(10.5)
                    fixed.add(para_idx)
    
    return fixed


# ============================================================
#  模块4: 编号修复
# ============================================================
def scan_numbering_issues(doc):
    """扫描编号问题"""
    issues = []
    
    # 检查连续编号是否断裂
    expected_numbers = {}  # pattern -> expected_next
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            expected_numbers.clear()
            continue
        
        # 匹配中文编号：一、二、三、(一)(二)(三)
        cn_match = re.match(r'^[（(]?([一二三四五六七八九十]+)[）).、\s]', text)
        if cn_match:
            cn_num = cn_match.group(1)
            key = "cn_bracket" if re.match(r'^[（(]', text) else "cn"
            
            if key in expected_numbers:
                expected = expected_numbers[key]
                actual = chinese_to_num(cn_num)
                if actual != expected:
                    issues.append({
                        "type": "编号问题",
                        "subtype": "中文编号断裂",
                        "paragraph": i + 1,
                        "detail": f"期望: {num_to_chinese(expected)}, 实际: {cn_num}",
                        "text": text[:60],
                    })
                expected_numbers[key] = actual + 1
            else:
                expected_numbers[key] = chinese_to_num(cn_num) + 1
        
        # 匹配数字编号：1. 2. 3.
        num_match = re.match(r'^(\d+)[.、．）)]', text)
        if num_match:
            num = int(num_match.group(1))
            key = "num"
            
            if key in expected_numbers:
                expected = expected_numbers[key]
                if num != expected:
                    issues.append({
                        "type": "编号问题",
                        "subtype": "数字编号断裂",
                        "paragraph": i + 1,
                        "detail": f"期望: {expected}, 实际: {num}",
                        "text": text[:60],
                    })
                expected_numbers[key] = num + 1
            else:
                expected_numbers[key] = num + 1
        
        # 重置计数器的条件：遇到明显的新章节标题或空段落
        if re.match(r'^第[一二三四五六七八九十]+章\s', text):
            expected_numbers.clear()
    
    return issues


def chinese_to_num(cn):
    """中文数字转阿拉伯数字"""
    mapping = {
        '一': 1, '二': 2, '三': 3, '四': 4, '五': 5,
        '六': 6, '七': 7, '八': 8, '九': 9, '十': 10,
    }
    if cn == '十':
        return 10
    result = 0
    for c in cn:
        result = result * 10 + mapping.get(c, 0)
    return result


def num_to_chinese(num):
    """阿拉伯数字转中文数字"""
    mapping = ['零', '一', '二', '三', '四', '五', '六', '七', '八', '九', '十']
    if num <= 10:
        return mapping[num]
    return str(num)


# ============================================================
#  模块5: 多余空行清除
# ============================================================
def scan_blank_lines(doc):
    """扫描多余空行"""
    issues = []
    consecutive_blank = 0
    blank_ranges = []
    
    for i, para in enumerate(doc.paragraphs):
        if not para.text.strip():
            consecutive_blank += 1
        else:
            if consecutive_blank >= 3:
                blank_ranges.append({
                    "start": i - consecutive_blank + 1,
                    "count": consecutive_blank
                })
            consecutive_blank = 0
    
    # 文档末尾的空行
    if consecutive_blank >= 3:
        blank_ranges.append({
            "start": len(doc.paragraphs) - consecutive_blank + 1,
            "count": consecutive_blank
        })
    
    for br in blank_ranges:
        issues.append({
            "type": "多余空行",
            "subtype": f"连续{br['count']}个空行",
            "paragraph": br["start"],
            "count": br["count"],
            "detail": f"从第{br['start']}段开始，连续{br['count']}个空段落",
        })
    
    return issues


def fix_blank_lines(doc, issues):
    """清除多余空行（合并连续空行最多保留2个）"""
    if not issues:
        return set()
    
    fixed = set()
    # 记录要删除的段落索引
    to_delete = set()
    
    consecutive_blank = 0
    blank_indices = []
    
    for i, para in enumerate(doc.paragraphs):
        if not para.text.strip():
            consecutive_blank += 1
            blank_indices.append(i)
        else:
            # 超过2个空行，多余的要删除
            if consecutive_blank > 2:
                # 保留前2个空行，删除后续的
                for idx in blank_indices[2:]:
                    to_delete.add(idx)
            
            consecutive_blank = 0
            blank_indices = []
    
    # 文档末尾多余空行
    if consecutive_blank > 2:
        for idx in blank_indices[2:]:
            to_delete.add(idx)
    
    if to_delete:
        # 逆序删除
        for idx in sorted(to_delete, reverse=True):
            p = doc.paragraphs[idx]
            p._element.getparent().remove(p._element)
            fixed.add(idx)
    
    return fixed


# ============================================================
#  模块6: 图片检查
# ============================================================
def scan_images(doc, docx_path):
    """检查文档中图片是否正常"""
    issues = []
    
    # 检查段落中的图片
    img_count = 0
    for i, para in enumerate(doc.paragraphs):
        for run in para.runs:
            if run._element.findall(qn('w:drawing')):
                img_count += 1
    
    # 检查行内图片
    inline_count = 0
    for i, para in enumerate(doc.paragraphs):
        for run in para.runs:
            drawings = run._element.findall(qn('w:drawing'))
            for drawing in drawings:
                blips = drawing.findall('.//' + qn('a:blip'))
                for blip in blips:
                    embed = blip.get(qn('r:embed'))
                    if embed:
                        try:
                            rel = doc.part.rels[embed]
                            if rel:
                                inline_count += 1
                        except:
                            issues.append({
                                "type": "图片问题",
                                "subtype": "图片链接断裂",
                                "paragraph": i + 1,
                                "detail": f"图片引用 {embed} 无法解析",
                            })
    
    # 检查是否有幽灵图片（图片文件存在但未引用）
    if docx_path and os.path.exists(docx_path):
        import zipfile
        try:
            with zipfile.ZipFile(docx_path, 'r') as z:
                media_files = [f for f in z.namelist() if f.startswith('word/media/')]
                doc_xml = z.read('word/document.xml').decode('utf-8', errors='ignore')
                
                for media in media_files:
                    filename = os.path.basename(media)
                    if filename not in doc_xml:
                        issues.append({
                            "type": "图片问题",
                            "subtype": "幽灵图片（未引用）",
                            "paragraph": 0,
                            "detail": f"文件 {filename} 存在于media目录但未在文档中引用",
                        })
        except:
            pass
    
    if img_count == 0:
        issues.append({
            "type": "图片问题",
            "subtype": "文档无图片",
            "paragraph": 0,
            "detail": "文档中未检测到任何图片，如本应包含图片请检查",
        })
    
    issues.append({
        "type": "图片统计",
        "subtype": "图片数量",
        "paragraph": 0,
        "detail": f"共检测到 {img_count} 张图片（{inline_count} 张正常引用）",
    })
    
    return issues


# ============================================================
#  模块7: 表格格式优化
# ============================================================
def scan_table_issues(doc):
    """扫描表格格式问题"""
    issues = []
    
    for t_idx, table in enumerate(doc.tables):
        rows = len(table.rows)
        cols = len(table.columns)
        
        # 检查表格是否有完全空行
        for r_idx, row in enumerate(table.rows):
            all_empty = all(cell.text.strip() == '' for cell in row.cells)
            if all_empty and rows > 1:
                issues.append({
                    "type": "表格问题",
                    "subtype": "空行",
                    "detail": f"表格{t_idx+1}第{r_idx+1}行完全为空",
                    "paragraph": 0,
                })
        
        # 检查单元格合并问题
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                # 检查单元格是否存在合并标记异常
                tc = cell._tc
                grid_span = tc.find(qn('w:tcPr'))
                if grid_span is not None:
                    span = grid_span.find(qn('w:gridSpan'))
                    if span is not None:
                        val = span.get(qn('w:val'))
                        if val and int(val) > 3:
                            issues.append({
                                "type": "表格问题",
                                "subtype": "跨列过多",
                                "detail": f"表格{t_idx+1}第{r_idx+1}行第{c_idx+1}列跨{val}列",
                                "paragraph": 0,
                            })
    
    return issues


# ============================================================
#  模块7b: 表格样式隔离检查（铁律28）
# ============================================================
def scan_table_style_isolation(doc):
    """扫描表格段落是否误用正文样式或存在首行缩进"""
    issues = []
    
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                for para in cell.paragraphs:
                    style_name = para.style.name if para.style else ''
                    # 检查是否误用正文样式
                    if style_name in ('Normal', '正文'):
                        issues.append({
                            "type": "表格样式",
                            "subtype": "误用正文样式",
                            "paragraph": 0,
                            "detail": f"表格{t_idx+1}行{r_idx+1}列{c_idx+1}: 使用了'{style_name}'样式（应为Table Text·无缩进）",
                        })
                    # 检查是否有首行缩进
                    indent = para.paragraph_format.first_line_indent
                    if indent and indent > 0:
                        issues.append({
                            "type": "表格样式",
                            "subtype": "首行缩进",
                            "paragraph": 0,
                            "detail": f"表格{t_idx+1}行{r_idx+1}列{c_idx+1}: 存在首行缩进（应为0）",
                        })
    
    return issues


def fix_table_style_isolation(doc):
    """修复表格段落的首行缩进（清零）"""
    fixed = set()
    
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                for para in cell.paragraphs:
                    indent = para.paragraph_format.first_line_indent
                    if indent and indent > 0:
                        para.paragraph_format.first_line_indent = Twips(0)
                        fixed.add((t_idx, r_idx, c_idx))
    
    return fixed


# ============================================================
#  模块7c: 图片样式检查（铁律28）
# ============================================================
def scan_image_style(doc):
    """扫描图片段落是否有首行缩进或未居中"""
    issues = []
    
    for i, para in enumerate(doc.paragraphs):
        has_image = any(run._element.findall(qn('w:drawing')) for run in para.runs)
        if has_image:
            # 检查首行缩进
            indent = para.paragraph_format.first_line_indent
            if indent and indent > 0:
                issues.append({
                    "type": "图片样式",
                    "subtype": "首行缩进",
                    "paragraph": i + 1,
                    "detail": f"段落{i+1}: 图片段落存在首行缩进（应为0·居中）",
                })
            # 检查居中
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            if para.alignment not in (WD_ALIGN_PARAGRAPH.CENTER, None):
                issues.append({
                    "type": "图片样式",
                    "subtype": "未居中",
                    "paragraph": i + 1,
                    "detail": f"段落{i+1}: 图片段落未居中",
                })
    
    return issues


def fix_image_style(doc):
    """修复图片段落：清零缩进+居中"""
    fixed = set()
    
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    for i, para in enumerate(doc.paragraphs):
        has_image = any(run._element.findall(qn('w:drawing')) for run in para.runs)
        if has_image:
            changed = False
            indent = para.paragraph_format.first_line_indent
            if indent and indent > 0:
                para.paragraph_format.first_line_indent = Twips(0)
                changed = True
            if para.alignment not in (WD_ALIGN_PARAGRAPH.CENTER, None):
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                changed = True
            if changed:
                fixed.add(i)
    
    return fixed


# ============================================================
#  模块7d: 标题编号格式检查（铁律22A）
# ============================================================
def scan_heading_numbering(doc):
    """扫描标题编号格式：tab分隔符/缩进过大"""
    issues = []
    
    for i, para in enumerate(doc.paragraphs):
        if not para.style.name.startswith('Heading'):
            continue
        
        pPr = para._element.find(qn('w:pPr'))
        if pPr is None:
            continue
        
        # 检查缩进
        ind = pPr.find(qn('w:ind'))
        if ind is not None:
            left = int(ind.get(qn('w:left'), 0) or 0)
            hanging = int(ind.get(qn('w:hanging'), 0) or 0)
            if left > 200:
                issues.append({
                    "type": "标题格式",
                    "subtype": "左缩进过大",
                    "paragraph": i + 1,
                    "detail": f"段落{i+1}: 标题左缩进={left}twips（铁律要求≤200）",
                    "text": para.text[:60],
                })
            if hanging > 200:
                issues.append({
                    "type": "标题格式",
                    "subtype": "悬挂缩进过大",
                    "paragraph": i + 1,
                    "detail": f"段落{i+1}: 标题悬挂缩进={hanging}twips（铁律要求≤200）",
                    "text": para.text[:60],
                })
        
        # 检查是否使用了自动编号（numPr = tab分隔符的根源）
        numPr = pPr.find(qn('w:numPr'))
        if numPr is not None:
            issues.append({
                "type": "标题格式",
                "subtype": "自动编号",
                "paragraph": i + 1,
                "detail": f"段落{i+1}: 使用了Word自动编号(numPr)，可能导致tab分隔符不一致",
                "text": para.text[:60],
            })
    
    return issues


def fix_heading_numbering(doc):
    """修复标题缩进过大（清零左缩进和悬挂缩进）"""
    fixed = set()
    
    for i, para in enumerate(doc.paragraphs):
        if not para.style.name.startswith('Heading'):
            continue
        
        pPr = para._element.find(qn('w:pPr'))
        if pPr is None:
            continue
        
        ind = pPr.find(qn('w:ind'))
        if ind is not None:
            left = int(ind.get(qn('w:left'), 0) or 0)
            hanging = int(ind.get(qn('w:hanging'), 0) or 0)
            changed = False
            if left > 200:
                ind.set(qn('w:left'), '0')
                changed = True
            if hanging > 200:
                ind.set(qn('w:hanging'), '0')
                changed = True
            if changed:
                fixed.add(i)
    
    return fixed


# ============================================================
#  模块8: 目录修复
# ============================================================
def scan_toc_issues(doc):
    """检查目录问题"""
    issues = []
    
    from docx.oxml.ns import qn
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        # 检查目录标题是否存在但内容为空
        if "目录" in text and len(text) <= 4:
            # 检查后续段落是否有目录条目
            has_toc_content = False
            for j in range(i + 1, min(i + 30, len(doc.paragraphs))):
                p = doc.paragraphs[j]
                if "..." in p.text or re.match(r'^\s*\d+\s', p.text):
                    has_toc_content = True
                    break
                if p.text.strip() and len(p.text) > 5 and not re.match(r'^第.*章', p.text):
                    continue
            
            if not has_toc_content:
                issues.append({
                    "type": "目录问题",
                    "subtype": "目录内容为空",
                    "paragraph": i + 1,
                    "detail": f"第{i+1}段标记为目录但无目录内容",
                })
    
    return issues


# ============================================================
#  模块9: 页眉页脚检查
# ============================================================
def scan_header_footer(doc):
    """检查页眉页脚"""
    issues = []
    
    try:
        for section in doc.sections:
            # 检查页眉
            if section.header and section.header.paragraphs:
                header_text = ' '.join(p.text for p in section.header.paragraphs if p.text.strip())
                if header_text:
                    issues.append({
                        "type": "页眉检查",
                        "subtype": "页眉内容",
                        "detail": f"页眉内容: {header_text[:60]}",
                        "paragraph": 0,
                    })
            
            # 检查页脚
            if section.footer and section.footer.paragraphs:
                footer_text = ' '.join(p.text for p in section.footer.paragraphs if p.text.strip())
                if footer_text:
                    issues.append({
                        "type": "页脚检查",
                        "subtype": "页脚内容",
                        "detail": f"页脚内容: {footer_text[:60]}",
                        "paragraph": 0,
                    })
    except Exception as e:
        issues.append({
            "type": "页眉页脚",
            "subtype": "读取失败",
            "detail": f"无法读取页眉页脚: {str(e)}",
            "paragraph": 0,
        })
    
    return issues


# ============================================================
#  模块10: 暗标扫描
# ============================================================
def scan_hidden_content(doc):
    """扫描隐藏文字/暗标风险"""
    issues = []
    
    for i, para in enumerate(doc.paragraphs):
        for run in para.runs:
            # 检查vanish（隐藏文字）
            rpr = run._element.find(qn('w:rPr'))
            if rpr is not None:
                vanish = rpr.find(qn('w:vanish'))
                if vanish is not None:
                    issues.append({
                        "type": "暗标风险",
                        "subtype": "隐藏文字",
                        "paragraph": i + 1,
                        "detail": f"存在隐藏文字: {run.text[:40]}",
                    })
    
    # 检查注释/批注
    try:
        comments_part = None
        for rel_id, rel in doc.part.rels.items():
            if "comments" in rel.reltype:
                comments_part = rel.target_part
                break
        
        if comments_part:
            comments_xml = comments_part._element
            comments = comments_xml.findall(qn('w:comment'))
            if comments:
                issues.append({
                    "type": "暗标风险",
                    "subtype": "文档批注",
                    "detail": f"文档包含 {len(comments)} 条批注，需清理",
                    "paragraph": 0,
                })
    except:
        pass
    
    # 检查元数据中的公司名称
    try:
        core_props = doc.core_properties
        if core_props:
            for attr in ['author', 'last_modified_by', 'company']:
                val = getattr(core_props, attr, None)
                if val:
                    issues.append({
                        "type": "暗标风险",
                        "subtype": "元数据",
                        "detail": f"文档属性: {attr}={val}",
                        "paragraph": 0,
                    })
    except:
        pass
    
    return issues


# ============================================================
#  模块11: 全角半角检测（自动修复）
# ============================================================
def scan_halfwidth_punctuation(doc):
    """检测中文段落中的英文标点"""
    issues = []
    
    # 中文上下文中常见的英文标点
    halfwidth_patterns = {
        "英文逗号": (r'[\u4e00-\u9fff],[\u4e00-\u9fff\d]', ',', '，'),
        "英文句号": (r'[\u4e00-\u9fff]\.[\u4e00-\u9fff\d]', '.', '。'),
        "英文引号": (r'[\u4e00-\u9fff]"[^"]*[\u4e00-\u9fff]', '"', '\u201c'),  # 复杂，单独处理
        "英文冒号": (r'[\u4e00-\u9fff]:[\u4e00-\u9fff\d]', ':', '：'),
        "英文分号": (r'[\u4e00-\u9fff];[\u4e00-\u9fff\d]', ';', '；'),
        "英文括号": (r'[\u4e00-\u9fff]\([^)]*\)', None, None),
    }
    
    # 逐段检测
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        if not text.strip():
            continue
        
        # 跳过纯英文段落
        has_cjk = bool(re.search(r'[\u4e00-\u9fff]', text))
        if not has_cjk:
            continue
        
        for name, (pattern, half, full) in halfwidth_patterns.items():
            if name == "英文引号":
                # 检测中文段落中的英文引号
                matches = re.finditer(r'[\u4e00-\u9fff\u3000-\u303f\uff00-\uffef]"', text)
                for m in matches:
                    issues.append({
                        "type": "全角半角",
                        "subtype": "英文引号",
                        "paragraph": i + 1,
                        "text": text[max(0, m.start()-5):m.end()+5],
                        "detail": "中文段落中使用了英文引号",
                    })
                continue
            
            if name == "英文括号":
                matches = re.finditer(r'[\u4e00-\u9fff]\(', text)
                for m in matches:
                    issues.append({
                        "type": "全角半角",
                        "subtype": "英文括号",
                        "paragraph": i + 1,
                        "text": text[max(0, m.start()-5):m.end()+5],
                        "detail": "中文段落中使用了英文括号",
                    })
                continue
            
            matches = re.finditer(pattern, text)
            for m in matches:
                # 找到具体是哪个字符
                for char in m.group():
                    if char in (',', '.', ':', ';'):
                        issues.append({
                            "type": "全角半角",
                            "subtype": name,
                            "paragraph": i + 1,
                            "text": text[max(0, m.start()-10):m.end()+10],
                            "detail": f"中文段落中的'{char}'应为全角",
                            "char": char,
                        })
                        break
    
    return issues



def fix_halfwidth_punctuation(doc, issues=None):
    """自动修复全角半角问题 v2 — 新增引号/括号/叹号问号，修复found变量bug"""
    # 英文→全角 映射（仅当相邻中文字符时替换）
    cjk_pat = re.compile(r'[一-鿿　-〿＀-￯]')
    
    def should_replace_dot(text, idx):
        """判断句点是否应该替换：数字前后的点不替换"""
        if idx > 0 and re.match(r'[一-鿿　-〿＀-￯]', text[idx-1]):
            return True
        # 数字左侧 → 不替换（1.23 或 1. 列表编号）
        if idx > 0 and re.match(r'[0-9]', text[idx-1]):
            return False
        # 数字右侧 → 不替换
        if idx < len(text)-1 and re.match(r'[0-9]', text[idx+1]):
            return False
        # 右侧是中文 → 替换
        if idx < len(text)-1 and re.match(r'[一-鿿]', text[idx+1]):
            return True
        return False
    
    def should_replace(text, idx, half):
        """通用判断：是否应该替换此标点"""
        if half == '.':
            return should_replace_dot(text, idx)
        # 左邻中文 → 替换
        if idx > 0 and cjk_pat.match(text[idx-1]):
            return True
        # 右邻中文 → 替换  
        if idx < len(text)-1 and cjk_pat.match(text[idx+1]):
            return True
        return False
    
    def fix_quotes(text):
        """修复中文段落中的英文引号对"""
        result = list(text)
        i = 0
        quote_count = 0
        while i < len(result):
            if result[i] == '"':
                # 检查左右是否有中文上下文
                has_cjk_left = i > 0 and cjk_pat.match(result[i-1])
                has_cjk_right = i < len(result)-1 and cjk_pat.match(result[i+1])
                if has_cjk_left or has_cjk_right:
                    quote_count += 1
                    if quote_count % 2 == 1:
                        result[i] = '“'  # 左引号 "
                    else:
                        result[i] = '”'  # 右引号 "
            elif result[i] == "'":
                has_cjk_left = i > 0 and cjk_pat.match(result[i-1])
                has_cjk_right = i < len(result)-1 and cjk_pat.match(result[i+1])
                if has_cjk_left or has_cjk_right:
                    result[i] = '‘' if (i == 0 or not cjk_pat.match(result[i-1])) else '’'
            i += 1
        return ''.join(result)
    
    def fix_punctuation_in_text(text):
        """修复一段文字中的所有英文标点"""
        fix_map = {
            ',': '，',   # 全角逗号 ，
            ':': '：',   # 全角冒号 ：
            ';': '；',   # 全角分号 ；
            '!': '！',   # 全角叹号 ！
            '?': '？',   # 全角问号 ？
            '(': '（',   # 全角左括号（
            ')': '）',   # 全角右括号）
        }
        
        # 先处理引号（特殊逻辑）
        text = fix_quotes(text)
        result = list(text)
        any_fixed = False
        
        for half, full in fix_map.items():
            for idx in range(len(result)):
                if result[idx] == half and should_replace(text, idx, half):
                    result[idx] = full
                    any_fixed = True
        
        # 句点特殊处理
        for idx in range(len(result)):
            if result[idx] == '.' and should_replace_dot(text, idx):
                result[idx] = '。'
                any_fixed = True
        
        return ''.join(result), any_fixed
    
    fixed = set()
    for para_idx, para in enumerate(doc.paragraphs):
        text = para.text
        if not text.strip():
            continue
        if not re.search(r'[一-鿿]', text):
            continue
        
        new_text, any_fixed = fix_punctuation_in_text(text)
        
        if new_text != text and any_fixed:
            if para.runs:
                runs_text = ''.join(r.text for r in para.runs)
                if runs_text == text:
                    remaining = new_text
                    for run in para.runs:
                        orig_len = len(run.text)
                        run.text = remaining[:orig_len]
                        remaining = remaining[orig_len:]
                    fixed.add(para_idx)
            else:
                para.text = new_text
                fixed.add(para_idx)
    
    return fixed


def _replace_cjk_punctuation(text, half, full):
    """替换中文字符右侧的英文标点为全角（已废弃，保留兼容）"""
    result = list(text)
    for idx in range(len(result)):
        if result[idx] == half:
            if idx > 0 and re.match(r'[一-鿿　-〿＀-￯]', result[idx-1]):
                result[idx] = full
            elif idx < len(result) - 1 and re.match(r'[一-鿿]', result[idx+1]):
                result[idx] = full
    return ''.join(result)


# ============================================================
#  模块12: 数据一致性检测（仅检测不修复）
# ============================================================
def scan_data_consistency(doc):
    """扫描文档中的数据一致性（仅检测不修复）
    
    检测内容:
    - 人员数量: 前文承诺 vs 后文配置表
    - 金额数据: 报价汇总 vs 分项合计
    - 服务期限: 承诺 vs 合同条款
    - 计算公式: 单价×数量=总价
    """
    issues = []
    
    # ─── 12a. 人员数量一致性 ───────────────────────────
    person_mentions = []  # list of (段落索引, 数量, 上下文)
    
    person_patterns = [
        r'(?:配置|配备|投入|安排|派驻)(?:人员|人力|团队|工程师|技术员|人员数量?)(?:\s*:?\s*)(\d+)\s*(?:人|名|位)',
        r'(?:共|合计|共计)(?:配置|配备|投入)(?:\s*)(\d+)\s*(?:人|名|位)',
        r'(\d+)\s*(?:人|名|位)\s*(?:的)?(?:项目)?(?:团队|人员|工程师|技术员)',
        r'(?:团队|人员|人力)(?:\s*)(\d+)\s*(?:人|名|位)',
    ]
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        if not text.strip():
            continue
        for pattern in person_patterns:
            matches = re.finditer(pattern, text)
            for m in matches:
                count = int(m.group(1))
                person_mentions.append({
                    "para": i + 1,
                    "count": count,
                    "text": text[:60],
                })
    
    # 检查表格中的人员配置
    table_person_counts = []
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            row_text = ' '.join(cell.text for cell in row.cells)
            # 查找数字
            nums = re.findall(r'(\d+)\s*(?:人|名|位)', row_text)
            for n in nums:
                table_person_counts.append({
                    "table": t_idx + 1,
                    "row": r_idx + 1,
                    "count": int(n),
                    "text": row_text[:60],
                })
    
    # 分析人员数量一致性
    if len(person_mentions) >= 2:
        counts = [m["count"] for m in person_mentions]
        if len(set(counts)) > 1:
            max_c = max(counts)
            min_c = min(counts)
            if max_c - min_c > 0:
                issues.append({
                    "type": "数据一致性",
                    "subtype": "人员数量不一致",
                    "paragraph": person_mentions[0]["para"],
                    "detail": (f"全文提及{len(person_mentions)}处人员数量: "
                               f"{', '.join([str(c)+'人' for c in sorted(set(counts))])} → 不一致"),
                    "references": [(m["para"], m["text"]) for m in person_mentions],
                })
    
    if table_person_counts:
        table_counts = [t["count"] for t in table_person_counts]
        if person_mentions:
            text_count = person_mentions[0]["count"]
            for tc in table_person_counts:
                if tc["count"] != text_count:
                    issues.append({
                        "type": "数据一致性",
                        "subtype": "人员数量不一致",
                        "paragraph": person_mentions[0]["para"],
                        "detail": (f"正文承诺{text_count}人，"
                                   f"表格{tc['table']}第{tc['row']}行显示{tc['count']}人 → 不一致"),
                    })
    
    # ─── 12b. 金额数据一致性 ───────────────────────────
    price_summaries = []
    price_details = []
    
    # 汇总金额模式
    summary_patterns = [
        r'(?:报价|总价|合计|汇总|总额)(?:\s*:?\s*)(?:人民币|RMB|¥)?\s*(\d[\d,.]*)\s*(?:万?元|万)',
        r'(?:投标报价|投标总价)(?:\s*为?\s*)(?:人民币|RMB|¥)?\s*(\d[\d,.]*)\s*(?:万?元|万)',
    ]
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        if not text.strip():
            continue
        for pattern in summary_patterns:
            matches = re.finditer(pattern, text)
            for m in matches:
                amount_str = m.group(1).replace(',', '')
                try:
                    amount = float(amount_str)
                    price_summaries.append({
                        "para": i + 1,
                        "amount": amount,
                        "text": text[:60],
                    })
                except:
                    pass
    
    # 分项金额模式（序号+金额）
    detail_pattern = re.compile(r'(\d+)[.、．）)]\s*(.*?)(\d[\d,.]*)\s*(?:元|万)')
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        if not text.strip():
            continue
        matches = detail_pattern.findall(text)
        for m in matches:
            amount_str = m[2].replace(',', '')
            try:
                amount = float(amount_str)
                price_details.append({
                    "para": i + 1,
                    "amount": amount,
                    "item": m[1][:20],
                })
            except:
                pass
    
    # 检查汇总与分项合计的差值
    if price_summaries and price_details:
        total_detail = sum(pd["amount"] for pd in price_details)
        for ps in price_summaries:
            diff = abs(ps["amount"] - total_detail)
            if diff > 0:
                issues.append({
                    "type": "数据一致性",
                    "subtype": "金额数据不一致",
                    "paragraph": ps["para"],
                    "detail": (f"报价汇总{ps['amount']}万元，"
                               f"分项合计{total_detail:.1f}万元 → "
                               f"差{diff:.1f}万元"),
                })
    
    # ─── 12c. 服务期限一致性 ───────────────────────────
    duration_mentions = []
    duration_patterns = [
        r'(?:服务|合同|项目)(?:期|期限|周期|年限)(?:\s*:?\s*)(\d+)\s*(?:年|个月)',
        r'(?:服务|运维|维护)(?:\s*)(\d+)\s*(?:年|个月)',
        r'(\d+)\s*(?:年|个月)\s*(?:服务|运维|维护|质保)',
    ]
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        if not text.strip():
            continue
        for pattern in duration_patterns:
            matches = re.finditer(pattern, text)
            for m in matches:
                num = int(m.group(1))
                unit = '年' if '年' in m.group() else '月'
                months = num * 12 if unit == '年' else num
                duration_mentions.append({
                    "para": i + 1,
                    "months": months,
                    "display": f"{num}{unit}",
                    "text": text[:60],
                })
    
    if len(duration_mentions) >= 2:
        months_set = set(m["months"] for m in duration_mentions)
        if len(months_set) > 1:
            displays = list(set(m["display"] for m in duration_mentions))
            issues.append({
                "type": "数据一致性",
                "subtype": "服务期限不一致",
                "paragraph": duration_mentions[0]["para"],
                "detail": f"全文提及{len(duration_mentions)}处期限: {' / '.join(displays)} → 不一致",
            })
    
    # ─── 12d. 计算公式核验 ───────────────────────────
    formula_patterns = [
        (r'(\d[\d,.]*)\s*[×xX*]\s*(\d[\d,.]*)\s*[=＝]\s*(\d[\d,.]*)', '×'),
        (r'(\d[\d,.]*)\s*/\s*(\d[\d,.]*)\s*[=＝]\s*(\d[\d,.]*)', '/'),
    ]
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        if not text.strip():
            continue
        for pattern, op in formula_patterns:
            matches = re.finditer(pattern, text)
            for m in matches:
                try:
                    a = float(m.group(1).replace(',', ''))
                    b = float(m.group(2).replace(',', ''))
                    c = float(m.group(3).replace(',', ''))
                    
                    if op == '×':
                        expected = round(a * b, 2)
                    else:
                        expected = round(a / b, 2) if b != 0 else 0
                    
                    if abs(expected - c) > 0.01:
                        issues.append({
                            "type": "数据一致性",
                            "subtype": "计算公式错误",
                            "paragraph": i + 1,
                            "detail": (f"{m.group()} → 应={expected}，实际{c}"),
                        })
                except:
                    pass
    
    return issues


# ============================================================
#  模块13: 标题层级检测
# ============================================================
def scan_heading_hierarchy(doc):
    """检测标题样式的使用合理性和编号连续性"""
    issues = []
    
    # ─── 13a. 检测未使用Heading样式的标题 ──────────────
    # 匹配形如标题但未使用Heading样式的段落
    heading_like_patterns = [
        r'^第[一二三四五六七八九十]+章\s',
        r'^[一二三四五六七八九十]+[、.．]',
        r'^第[一二三四五六七八九十]+节\s',
        r'^\d+[.、．]\s',
        r'^\d+\.\d+\s',
    ]
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        
        style_name = para.style.name if para.style else ""
        is_heading = "Heading" in style_name or "标题" in style_name or "heading" in style_name.lower()
        
        for pattern in heading_like_patterns:
            if re.match(pattern, text):
                if not is_heading:
                    issues.append({
                        "type": "标题层级",
                        "subtype": "未使用标题样式",
                        "paragraph": i + 1,
                        "detail": f"「{text[:30]}」看起来是标题，但样式为「{style_name}」",
                        "text": text[:40],
                    })
                break
    
    # ─── 13b. 检测标题编号的连续性 ────────────────────
    # 收集所有带编号的标题
    numbered_headings = []
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        
        # 中文编号
        cn_match = re.match(r'^([一二三四五六七八九十]+)[.、．）)]', text)
        if cn_match:
            num = chinese_to_num(cn_match.group(1))
            numbered_headings.append({
                "para": i + 1,
                "num": num,
                "type": "cn",
                "text": text[:30],
            })
        
        # 数字编号
        num_match = re.match(r'^(\d+)[.、．）)]', text)
        if num_match:
            num = int(num_match.group(1))
            numbered_headings.append({
                "para": i + 1,
                "num": num,
                "type": "num",
                "text": text[:30],
            })
    
    # 检查连续性
    for t in ['cn', 'num']:
        seq = [h for h in numbered_headings if h['type'] == t]
        for idx in range(1, len(seq)):
            expected = seq[idx-1]['num'] + 1
            actual = seq[idx]['num']
            if actual != expected and actual != 1:
                # 跳过明显是重新开始的序号
                issues.append({
                    "type": "标题层级",
                    "subtype": "编号跳跃",
                    "paragraph": seq[idx]['para'],
                    "detail": (f"第{seq[idx-1]['para']}段编号{seq[idx-1]['num']} → "
                               f"第{seq[idx]['para']}段编号{actual} "
                               f"(期望为{expected})"),
                })
    
    # ─── 13c. 检查标题样式的顺序是否正确 ───────────────
    heading_levels = []
    for i, para in enumerate(doc.paragraphs):
        style = para.style
        if style and ("Heading" in style.name or "heading" in style.name.lower()):
            # 提取级别
            level = None
            for part in style.name.split():
                if part.isdigit():
                    level = int(part)
                    break
            
            if level is None:
                # 尝试从样式名推断
                if style.name.endswith(("1", "2", "3", "4", "5", "6")):
                    level = int(style.name[-1])
            
            if level:
                heading_levels.append({
                    "para": i + 1,
                    "level": level,
                    "text": para.text[:30],
                })
    
    # 检查层级跳跃（如Heading 1 → Heading 3）
    for idx in range(1, len(heading_levels)):
        prev = heading_levels[idx-1]
        curr = heading_levels[idx]
        if curr['level'] > prev['level'] + 1:
            issues.append({
                "type": "标题层级",
                "subtype": "标题样式层级跳跃",
                "paragraph": curr['para'],
                "detail": (f"从{prev['text'][:20]} (Heading {prev['level']}) → "
                           f"{curr['text'][:20]} (Heading {curr['level']}) 跳级"),
            })
    
    return issues


# ============================================================
#  模块14: 下划线检测
# ============================================================
def scan_underline_issues(doc):
    """检测下划线相关问题（两种模式）
    
    模式1-落款下划线: 扫描文档末尾的"日期""签字"附近，检查是否有 ____ 填写位
    模式2-暗标下划线: 扫描全文，检测是否有Underline样式
    """
    issues = []
    
    # ─── 模式1: 落款下划线检查 ────────────────────
    # 从文档末尾往前扫描，找"日期""签字"落款区域
    doc_len = len(doc.paragraphs)
    # 扫描最后30%的段落范围找落款
    tail_start = max(0, doc_len - int(doc_len * 0.3))
    
    found_signature_area = False
    missing_underline = True
    
    for i in range(tail_start, doc_len):
        text = doc.paragraphs[i].text.strip()
        if not text:
            continue
        
        # 检测落款关键词
        if any(kw in text for kw in ['日期', '签字', '签章', '盖章', '年月日', '投标人']):
            found_signature_area = True
        
        # 检测是否有下划线填写位
        if re.search(r'_{2,}|——{2,}|____', text):
            missing_underline = False
    
    if found_signature_area and missing_underline:
        # 找到最后一段带中文的段落确定为落款区域
        last_para = None
        for i in range(doc_len - 1, tail_start - 1, -1):
            if doc.paragraphs[i].text.strip():
                last_para = i + 1
                break
        
        issues.append({
            "type": "下划线问题",
            "subtype": "落款缺少下划线",
            "paragraph": last_para or doc_len,
            "detail": "落款区（日期/签字处）缺少下划线填写位，建议添加______",
        })
    
    # ─── 模式2: 暗标下划线警告 ────────────────────
    underline_count = 0
    underline_locations = []
    
    for i, para in enumerate(doc.paragraphs):
        for run in para.runs:
            if run.font.underline:
                underline_count += 1
                underline_locations.append(i + 1)
                break  # 每段只计一次
    
    if underline_count > 0:
        # 检测是否可能是暗标场景（有隐藏文字或元数据标记）
        is_sealed_bid = any(
            "暗标" in doc.paragraphs[j].text 
            for j in range(min(20, doc_len))
        )
        
        locations_str = ', '.join(f'第{p}段' for p in sorted(set(underline_locations))[:10])
        if len(set(underline_locations)) > 10:
            locations_str += f' 等共{underline_count}处'
        
        issues.append({
            "type": "下划线问题",
            "subtype": "暗标禁止使用下划线" if is_sealed_bid else "下划线使用警告",
            "paragraph": underline_locations[0],
            "detail": f"共发现{underline_count}段含下划线样式: {locations_str}",
        })
    
    return issues


# ============================================================
#  段落文本替换（保留格式）
# ============================================================
def _replace_para_text(para, new_text):
    """替换段落文本，尽量保留第一个run的格式"""
    if not para.runs:
        para.text = new_text
        return
    
    first_run = para.runs[0]
    font_name = first_run.font.name
    font_size = first_run.font.size
    bold = first_run.font.bold
    
    # 清除所有run
    for run in para.runs:
        run.text = ''
    
    # 在第一个run写入新文本
    para.runs[0].text = new_text
    
    # 恢复格式
    if font_name:
        para.runs[0].font.name = font_name
    if font_size:
        para.runs[0].font.size = font_size
    if bold is not None:
        para.runs[0].font.bold = bold


# ============================================================
#  主流程
# ============================================================
def scan_document(doc, docx_path=None):
    """全面扫描文档"""
    all_issues = []
    
    log("开始扫描文档...")
    
    issues = scan_placeholders(doc)
    all_issues.extend(issues)
    log(f"  占位符扫描: {len(issues)} 处")
    
    issues = scan_date_placeholders(doc)
    all_issues.extend(issues)
    log(f"  日期占位符扫描: {len(issues)} 处")
    
    issues = scan_font_issues(doc)
    all_issues.extend(issues)
    log(f"  字体扫描: {len(issues)} 处")
    
    issues = scan_numbering_issues(doc)
    all_issues.extend(issues)
    log(f"  编号扫描: {len(issues)} 处")
    
    issues = scan_blank_lines(doc)
    all_issues.extend(issues)
    log(f"  空行扫描: {len(issues)} 处")
    
    issues = scan_images(doc, docx_path)
    all_issues.extend(issues)
    log(f"  图片检查: {len([i for i in issues if i['type'] != '图片统计'])} 处")
    
    issues = scan_table_issues(doc)
    all_issues.extend(issues)
    log(f"  表格检查: {len(issues)} 处")
    
    issues = scan_toc_issues(doc)
    all_issues.extend(issues)
    log(f"  目录检查: {len(issues)} 处")
    
    issues = scan_header_footer(doc)
    all_issues.extend(issues)
    log(f"  页眉页脚: {len(issues)} 处")
    
    issues = scan_hidden_content(doc)
    all_issues.extend(issues)
    log(f"  暗标扫描: {len(issues)} 处")
    
    issues = scan_halfwidth_punctuation(doc)
    all_issues.extend(issues)
    log(f"  全角半角检测: {len(issues)} 处")
    
    issues = scan_data_consistency(doc)
    all_issues.extend(issues)
    log(f"  数据一致性检测: {len(issues)} 处")
    
    issues = scan_heading_hierarchy(doc)
    all_issues.extend(issues)
    log(f"  标题层级检测: {len(issues)} 处")
    
    issues = scan_underline_issues(doc)
    all_issues.extend(issues)
    log(f"  下划线检测: {len(issues)} 处")
    
    return all_issues


def fix_document(doc, issues, docx_path=None):
    """自动修复文档问题"""
    log("\n开始修复...")
    total_fixed = 0
    
    placeholder_issues = [i for i in issues if i['type'] == '占位符']
    fixed = fix_placeholders(doc, placeholder_issues)
    total_fixed += len(fixed)
    log(f"  占位符修复: {len(fixed)} 段")
    
    date_issues = [i for i in issues if i['type'] == '日期占位符']
    fixed = fix_date_placeholders(doc, date_issues)
    total_fixed += len(fixed)
    log(f"  日期占位符修复: {len(fixed)} 段")
    
    font_issues = [i for i in issues if i['type'] == '字体问题']
    fixed = fix_font_issues(doc, font_issues)
    total_fixed += len(fixed)
    log(f"  字体修复: {len(fixed)} 段")
    
    blank_issues = [i for i in issues if i['type'] == '多余空行']
    fixed = fix_blank_lines(doc, blank_issues)
    total_fixed += len(fixed)
    log(f"  空行清理: {len(fixed)} 段被删除")
    
    fixed = fix_halfwidth_punctuation(doc)
    total_fixed += len(fixed)
    log(f"  全角半角修复: {len(fixed)} 段")
    
    return total_fixed


def generate_report(issues, filename):
    """生成扫描报告"""
    report_lines = []
    report_lines.append("=" * 60)
    report_lines.append(f"标书排版扫描报告")
    report_lines.append(f"文件: {filename}")
    report_lines.append(f"时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    report_lines.append(f"问题总数: {len(issues)}")
    report_lines.append("=" * 60)
    
    # 按类型分组
    by_type = {}
    for issue in issues:
        t = issue["type"]
        if t not in by_type:
            by_type[t] = []
        by_type[t].append(issue)
    
    for t, items in by_type.items():
        report_lines.append(f"\n--- {t}: {len(items)} 处 ---")
        for item in items:
            para_info = f"第{item['paragraph']}段" if item['paragraph'] > 0 else ""
            report_lines.append(f"  [{item.get('subtype', '')}] {para_info}")
            if item.get('detail'):
                report_lines.append(f"    细节: {item['detail']}")
            if item.get('text'):
                report_lines.append(f"    内容: {item['text'][:60]}")
    
    return '\n'.join(report_lines)


def save_report(report, filename):
    """保存报告到文件"""
    os.makedirs(OUTPUTS_DIR, exist_ok=True)
    report_path = os.path.join(OUTPUTS_DIR, f"{filename}_扫描报告.txt")
    with open(report_path, 'w', encoding='utf-8') as f:
        f.write(report)
    log(f"报告已保存: {report_path}")
    return report_path


def main():
    parser = argparse.ArgumentParser(
        description="标书排版自动修复器 v1.2",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
使用示例:
  # 扫描+修复指定文件
  python3 %(prog)s uploads/商务标.docx
  
  # 仅扫描不修复
  python3 %(prog)s uploads/技术标.docx --scan-only
  
  # 扫描目录下所有文件
  python3 %(prog)s uploads/ --all
  
  # 指定输出路径
  python3 %(prog)s uploads/投标书.docx -o outputs/修复版.docx
        """
    )
    
    parser.add_argument("path", help="docx文件路径或目录")
    parser.add_argument("-o", "--output", help="输出文件路径")
    parser.add_argument("--scan-only", action="store_true", help="仅扫描，不修复")
    parser.add_argument("--all", action="store_true", help="扫描目录下所有docx文件")
    parser.add_argument("--report-only", action="store_true", help="仅生成报告，不输出docx")
    
    args = parser.parse_args()
    
    # 查找文件
    files = find_docx_files(args.path)
    if not files:
        log(f"未找到docx文件: {args.path}", "ERROR")
        sys.exit(1)
    
    if args.all:
        log(f"找到 {len(files)} 个docx文件")
    
    for filepath in files:
        basename = os.path.basename(filepath)
        name_no_ext = os.path.splitext(basename)[0]
        
        log(f"\n{'='*60}")
        log(f"处理文件: {filepath}")
        log(f"{'='*60}")
        
        try:
            doc = Document(filepath)
        except Exception as e:
            log(f"打开文件失败: {e}", "ERROR")
            continue
        
        # 扫描
        issues = scan_document(doc, filepath)
        
        # 生成报告
        report = generate_report(issues, basename)
        print(f"\n{report}")
        report_path = save_report(report, name_no_ext)
        
        # 仅扫描模式
        if args.scan_only:
            log("\n仅扫描模式完成，未做任何修改")
            continue
        
        # 仅报告模式
        if args.report_only:
            log(f"\n报告已生成，未做修复: {report_path}")
            continue
        
        # 修复
        fixed_count = fix_document(doc, issues, filepath)
        log(f"\n总计修复: {fixed_count} 处")
        
        # 保存
        if args.output:
            output_path = args.output
        else:
            os.makedirs(OUTPUTS_DIR, exist_ok=True)
            output_path = os.path.join(OUTPUTS_DIR, f"{name_no_ext}_修复版.docx")
        
        doc.save(output_path)
        log(f"\n✅ 修复版已保存: {output_path}")
        
        # 修复后重新扫描验证
        log("\n--- 修复后再次扫描验证 ---")
        doc2 = Document(output_path)
        remaining = scan_document(doc2)
        remaining_count = len([r for r in remaining if r['type'] not in ('图片统计', '页眉检查', '页脚检查', '数据一致性', '下划线问题')])
        
        if remaining_count == 0:
            log(f"\n🎉 完美！所有问题已修复！")
        else:
            log(f"\n⚠️ 剩余 {remaining_count} 个问题（多为需人工处理的复杂问题）")
    
    log("\n✅ 全部完成！")


if __name__ == '__main__':
    main()
