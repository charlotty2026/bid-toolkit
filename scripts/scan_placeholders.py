#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""占位符扫描工具"""
import re, sys, os
sys.stdout.reconfigure(encoding='utf-8')

def scan_placeholders(text):
    """扫描[XX]、[400-]等占位符"""
    patterns = [
        r'\[XX\]',
        r'\[\d{3}-\]',
        r'\[.*?XX.*?\]',
        r'\d{4}年\d{1,2}月\s+日',
    ]
    
    found = []
    for pattern in patterns:
        matches = re.findall(pattern, text)
        found.extend(matches)
    
    if found:
        print(f'⚠️ 发现{len(found)}处占位符/未填日期:')
        for f in set(found):
            print(f'  - {f}')
    else:
        print('✅ 无占位符')
    
    return found

def _read_file_text(path):
    """读取文件文本，支持 .docx / .md / .txt"""
    if path.endswith('.docx'):
        try:
            from docx import Document
        except ImportError:
            print("❌ 需要 python-docx: pip install python-docx")
            sys.exit(1)
        doc = Document(path)
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        return '\n'.join(parts)
    else:
        with open(path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()


if __name__ == '__main__':
    import argparse
    parser = argparse.ArgumentParser(description='占位符扫描工具 - 扫描标书中的[XX]、待填日期等占位符残留')
    parser.add_argument('file', help='要扫描的文件路径（支持 .docx / .md / .txt）')
    args = parser.parse_args()

    if not os.path.isfile(args.file):
        print(f'❌ 文件不存在: {args.file}')
        sys.exit(1)

    text = _read_file_text(args.file)
    found = scan_placeholders(text)
    sys.exit(1 if found else 0)
