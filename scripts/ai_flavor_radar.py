#!/usr/bin/env python3
"""
AI味雷达 (AI Flavor Radar)
==========================
AI文风检测工具。区分投标方案和自媒体两种场景，
保留专业术语和书面语，精准定位AI模板腔。

支持两种输出模式：
  - 文本/JSON/Markdown报告（检测+建议）
  - Word文档修订标记（直接在原文上标修订，逐条接受/拒绝）

用法:
    python ai_flavor_radar.py <文件路径> [--mode bid|social] [--format text|json|markdown|docx]
    python ai_flavor_radar.py --stdin [--mode bid|social]  # 从stdin读取
    cat 文件.md | python ai_flavor_radar.py --stdin --mode bid

示例:
    python ai_flavor_radar.py 投标方案.md --mode bid
    python ai_flavor_radar.py 公众号文章.md --mode social --format markdown
    python ai_flavor_radar.py 方案.docx --mode bid --format docx -o 方案_修订.docx
    echo "值得注意的是，我们将全面提升服务质量" | python ai_flavor_radar.py --stdin --mode bid

License: MIT
"""

import argparse
import json
import os
import re
import sys
from dataclasses import dataclass, field, asdict
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Optional, Tuple


# ============================================================
# 数据结构
# ============================================================

@dataclass
class Hit:
    """单条检测结果"""
    rule_id: str          # 规则ID，如 C01, B01
    rule_name: str        # 规则名称，如 "起承转合套话"
    category: str         # 分类，如 "结构套话"
    severity: str         # 严重程度: fatal/high/medium/low
    line_num: int         # 行号（1-indexed）
    col_start: int        # 匹配起始列（0-indexed）
    col_end: int          # 匹配结束列
    matched_text: str     # 匹配到的原文
    full_line: str        # 完整行原文
    suggestion: str       # 修改建议
    example: Optional[Dict] = None  # 改前→改后示范（来自规则的examples[0]）


@dataclass
class ScanResult:
    """完整扫描结果"""
    file_path: str
    mode: str             # bid / social
    total_lines: int
    total_chars: int
    hits: List[Hit] = field(default_factory=list)

    @property
    def hit_count(self) -> int:
        return len(self.hits)

    @property
    def fatal_count(self) -> int:
        return sum(1 for h in self.hits if h.severity == "fatal")

    @property
    def high_count(self) -> int:
        return sum(1 for h in self.hits if h.severity == "high")

    @property
    def medium_count(self) -> int:
        return sum(1 for h in self.hits if h.severity == "medium")

    @property
    def low_count(self) -> int:
        return sum(1 for h in self.hits if h.severity == "low")

    @property
    def score(self) -> int:
        """AI味评分 0-100，越高越像AI写的"""
        if self.total_chars == 0:
            return 0
        # 按严重程度加权
        weights = {"fatal": 15, "high": 8, "medium": 4, "low": 2}
        raw = sum(weights.get(h.severity, 2) for h in self.hits)
        # 按文本长度归一化（每1000字）
        normalized = raw / (self.total_chars / 1000) if self.total_chars > 0 else raw
        score = min(100, int(normalized * 10))
        return score

    @property
    def grade(self) -> str:
        """AI味等级"""
        s = self.score
        if s >= 70:
            return "🔴 重度AI味"
        elif s >= 40:
            return "🟡 中度AI味"
        elif s >= 15:
            return "🟢 轻度AI味"
        else:
            return "✅ 基本自然"


# ============================================================
# 规则加载
# ============================================================

RULES_DIR = Path(__file__).parent / "ai_flavor_rules"


# 模式中文名映射
MODE_NAMES = {
    "bid": "投标方案",
    "social": "自媒体",
    "xiaohongshu": "小红书",
    "email": "邮件",
    "paper": "学术论文",
}


def load_rules(mode: str) -> List[Dict]:
    """加载规则：common + 场景规则"""
    rules = []

    # 通用规则
    common_path = RULES_DIR / "common.json"
    if common_path.exists():
        with open(common_path, "r", encoding="utf-8") as f:
            data = json.load(f)
            rules.extend(data.get("rules", []))

    # 场景规则
    scene_path = RULES_DIR / f"{mode}.json"
    if scene_path.exists():
        with open(scene_path, "r", encoding="utf-8") as f:
            data = json.load(f)
            rules.extend(data.get("rules", []))

    return rules


def load_whitelist() -> List[str]:
    """加载全局白名单：命中白名单的文本跳过检测（误杀防护）"""
    whitelist_path = RULES_DIR / "whitelist.json"
    if not whitelist_path.exists():
        return []
    try:
        with open(whitelist_path, "r", encoding="utf-8") as f:
            data = json.load(f)
            return data.get("whitelist", [])
    except (json.JSONDecodeError, OSError):
        return []


# ============================================================
# 检测引擎
# ============================================================

class FlavorRadar:
    def __init__(self, mode: str = "bid", custom_rules_dir: Optional[str] = None):
        self.mode = mode
        if custom_rules_dir:
            global RULES_DIR
            RULES_DIR = Path(custom_rules_dir)
        self.rules = load_rules(mode)
        self.whitelist = load_whitelist()
        self._apply_mode_overrides(mode)

    def _apply_mode_overrides(self, mode: str):
        """Adjust rule severity by mode (same rule, different weight per scene)"""
        if mode == "bid":
            for rule in self.rules:
                # C03 buzzwords may be normal industry jargon in bid context
                if rule["id"] == "C03":
                    rule["severity"] = "low"

    def scan(self, text: str, file_path: str = "<stdin>") -> ScanResult:
        """扫描文本，返回结果"""
        lines = text.split("\n")
        result = ScanResult(
            file_path=file_path,
            mode=self.mode,
            total_lines=len(lines),
            total_chars=len(text),
        )

        for i, line in enumerate(lines, 1):
            for rule in self.rules:
                # C05 cross-paragraph check: skip per-line, handle separately
                if rule["id"] == "C05":
                    continue
                hits = self._check_rule(rule, line, i)
                result.hits.extend(hits)

        # C05: full-text cross-paragraph detection
        for rule in self.rules:
            if rule["id"] == "C05":
                c05_hits = self._check_c05_cross_paragraph(rule, lines)
                result.hits.extend(c05_hits)

        return result

    def _check_c05_cross_paragraph(self, rule: Dict, lines: List[str]) -> List[Hit]:
        """C05 cross-paragraph detection: count list-symbol groups across paragraphs.
        Same symbol group used in >2 paragraphs = template repetition."""
        hits = []

        symbol_groups = [
            ("\u2460\u2461\u2462\u2463\u2464", [r"\u2460", r"\u2461", r"\u2462"]),
            ("\uff081\uff09\uff082\uff09\uff083\uff09", [r"\uff081\uff09", r"\uff082\uff09", r"\uff083\uff09"]),
            ("\u7b2c\u4e00\u7b2c\u4e8c\u7b2c\u4e09", [r"\u7b2c\u4e00[\uff0c,\u3001]", r"\u7b2c\u4e8c[\uff0c,\u3001]", r"\u7b2c\u4e09[\uff0c,\u3001]"]),
            ("\u9996\u5148\u5176\u6b21\u518d\u6b21", [r"\u9996\u5148[\uff0c,]", r"\u5176\u6b21[\uff0c,]", r"\u518d\u6b21[\uff0c,]"]),
        ]

        for group_name, patterns in symbol_groups:
            paragraph_hits = []
            for i, line in enumerate(lines, 1):
                stripped = line.strip()
                if not stripped or stripped.startswith("```"):
                    continue
                matched_in_line = []
                for p in patterns:
                    if re.search(p, line):
                        matched_in_line.append(p)
                if len(matched_in_line) >= 2:
                    paragraph_hits.append((i, line.rstrip(), matched_in_line))

            if len(paragraph_hits) > 2:
                for line_num, full_line, matched in paragraph_hits:
                    hits.append(Hit(
                        rule_id=rule["id"],
                        rule_name=rule["name"],
                        category=rule["category"],
                        severity=rule["severity"],
                        line_num=line_num,
                        col_start=0,
                        col_end=len(full_line),
                        matched_text=f"\u8de8\u6bb5\u91cd\u590d\u4f7f\u7528{group_name}\uff08\u5168\u6587\u5171{len(paragraph_hits)}\u6bb5\u4f7f\u7528\uff09",
                        full_line=full_line,
                        suggestion=rule.get("suggestion", ""),
                    ))

        return hits

    def _is_whitelisted(self, matched_text: str) -> bool:
        """判断命中文本是否在白名单内（误杀防护）"""
        for wl in self.whitelist:
            if wl and wl in matched_text:
                return True
        return False

    def _check_rule(self, rule: Dict, line: str, line_num: int) -> List[Hit]:
        """检查单行是否命中规则"""
        hits = []

        # 跳过Markdown代码块内的内容
        stripped = line.strip()
        if stripped.startswith("```") or stripped.startswith("| `"):
            return hits

        # 自定义检查（长难句、长段落等）
        if rule.get("custom_check"):
            hits.extend(self._custom_check(rule, line, line_num))
            return hits

        # 正则检查
        patterns = rule.get("patterns", [])
        examples = rule.get("examples", [])
        example = examples[0] if examples else None
        for pattern in patterns:
            try:
                for match in re.finditer(pattern, line):
                    matched_text = match.group()
                    # 误杀防护：命中文本含白名单词则跳过
                    if self._is_whitelisted(matched_text):
                        continue
                    hits.append(Hit(
                        rule_id=rule["id"],
                        rule_name=rule["name"],
                        category=rule["category"],
                        severity=rule["severity"],
                        line_num=line_num,
                        col_start=match.start(),
                        col_end=match.end(),
                        matched_text=matched_text,
                        full_line=line.rstrip(),
                        suggestion=rule.get("suggestion", ""),
                        example=example,
                    ))
            except re.error:
                # 正则编译失败，跳过
                continue

        return hits

    def _custom_check(self, rule: Dict, line: str, line_num: int) -> List[Hit]:
        """自定义检测逻辑"""
        check_type = rule.get("custom_check")
        threshold = rule.get("threshold", {})
        hits = []

        if check_type == "long_sentence":
            # 长难句检测：按句号/问号/叹号分句，检查每句长度
            char_limit = threshold.get("char_count", 80)
            comma_limit = threshold.get("comma_count", 4)
            de_limit = threshold.get("de_count", 4)

            # 按中文标点分句
            sentences = re.split(r'[。！？；]', line)
            for sent in sentences:
                sent = sent.strip()
                if not sent or len(sent) < char_limit:
                    continue

                comma_count = sent.count("，") + sent.count(",")
                de_count = sent.count("的")

                reasons = []
                if len(sent) >= char_limit:
                    reasons.append(f"{len(sent)}字")
                if comma_count >= comma_limit:
                    reasons.append(f"{comma_count}个逗号")
                if de_count >= de_limit:
                    reasons.append(f"{de_count}个'的'")

                if reasons:
                    # 找到句子在行中的位置
                    idx = line.find(sent)
                    hits.append(Hit(
                        rule_id=rule["id"],
                        rule_name=rule["name"],
                        category=rule["category"],
                        severity=rule["severity"],
                        line_num=line_num,
                        col_start=max(0, idx),
                        col_end=max(0, idx) + len(sent),
                        matched_text=sent[:60] + "..." if len(sent) > 60 else sent,
                        full_line=line.rstrip(),
                        suggestion=f"长难句({', '.join(reasons)})。{rule.get('suggestion', '')}",
                    ))

        elif check_type == "long_paragraph":
            # 长段落检测
            char_limit = threshold.get("char_count", 200)
            if len(line.strip()) >= char_limit:
                hits.append(Hit(
                    rule_id=rule["id"],
                    rule_name=rule["name"],
                    category=rule["category"],
                    severity=rule["severity"],
                    line_num=line_num,
                    col_start=0,
                    col_end=len(line),
                    matched_text=f"({len(line.strip())}字)",
                    full_line=line.rstrip()[:80] + "..." if len(line) > 80 else line.rstrip(),
                    suggestion=rule.get("suggestion", ""),
                ))

        return hits


# ============================================================
# 报告输出
# ============================================================

# 终端颜色
COLORS = {
    "fatal": "\033[91m",   # 红色
    "high": "\033[93m",    # 黄色
    "medium": "\033[33m",  # 橙色
    "low": "\033[90m",     # 灰色
    "reset": "\033[0m",
    "bold": "\033[1m",
    "cyan": "\033[96m",
    "green": "\033[92m",
}


def format_text_report(result: ScanResult, use_color: bool = True) -> str:
    """终端文本格式报告"""
    lines = []
    c = COLORS if use_color else {k: "" for k in COLORS}

    # 头部
    lines.append(f"{c['bold']}{'='*60}")
    lines.append(f"  AI味雷达扫描报告")
    lines.append(f"{'='*60}{c['reset']}")
    lines.append(f"  文件: {result.file_path}")
    mode_name = MODE_NAMES.get(result.mode, result.mode)
    lines.append(f"  模式: {mode_name}")
    lines.append(f"  行数: {result.total_lines}  字数: {result.total_chars}")
    lines.append("")

    # 评分
    lines.append(f"  AI味评分: {result.score}/100  {result.grade}")
    lines.append(f"  命中: {result.hit_count}条 "
                 f"(致命{result.fatal_count} / 高{result.high_count} / "
                 f"中{result.medium_count} / 低{result.low_count})")
    lines.append(f"{c['bold']}{'='*60}{c['reset']}")
    lines.append("")

    if result.hit_count == 0:
        lines.append(f"  {c['green']}✅ 未检测到AI味，文章看起来很自然！{c['reset']}")
        return "\n".join(lines)

    # 按严重程度分组
    severity_order = ["fatal", "high", "medium", "low"]
    severity_label = {
        "fatal": "致命",
        "high": "高危",
        "medium": "中危",
        "low": "低危",
    }

    for sev in severity_order:
        sev_hits = [h for h in result.hits if h.severity == sev]
        if not sev_hits:
            continue

        lines.append(f"{c['bold']}  【{severity_label[sev]}】{c['reset']}")

        for hit in sev_hits:
            color = c.get(sev, "")
            reset = c["reset"]

            lines.append(f"")
            lines.append(f"  {color}▸ 第{hit.line_num}行 [{hit.rule_id}] {hit.rule_name}{reset}")
            lines.append(f"    原文: {hit.matched_text}")
            lines.append(f"    建议: {hit.suggestion}")
            # 展示改前→改后示范
            example = hit.example
            if example:
                lines.append(f"    {c['cyan']}示范:{c['reset']}")
                lines.append(f"      {c['reset']}改前: {example.get('before', '')}")
                lines.append(f"      {c['reset']}改后: {example.get('after', '')}")

        lines.append("")

    # 分类统计
    lines.append(f"{c['cyan']}{'─'*60}")
    lines.append(f"  分类统计:{c['reset']}")
    categories = {}
    for h in result.hits:
        cat = h.category
        if cat not in categories:
            categories[cat] = 0
        categories[cat] += 1
    for cat, count in sorted(categories.items(), key=lambda x: -x[1]):
        bar = "█" * min(count, 30)
        lines.append(f"    {cat:8s} {bar} {count}")
    lines.append("")

    return "\n".join(lines)


def format_json_report(result: ScanResult) -> str:
    """JSON格式报告"""
    data = {
        "file_path": result.file_path,
        "mode": result.mode,
        "total_lines": result.total_lines,
        "total_chars": result.total_chars,
        "score": result.score,
        "grade": result.grade,
        "hit_count": result.hit_count,
        "fatal_count": result.fatal_count,
        "high_count": result.high_count,
        "medium_count": result.medium_count,
        "low_count": result.low_count,
        "hits": [asdict(h) for h in result.hits],
    }
    return json.dumps(data, ensure_ascii=False, indent=2)


def format_markdown_report(result: ScanResult) -> str:
    """Markdown格式报告"""
    lines = []
    lines.append(f"# AI味雷达扫描报告")
    lines.append(f"")
    lines.append(f"| 项目 | 值 |")
    lines.append(f"|------|-----|")
    lines.append(f"| 文件 | `{result.file_path}` |")
    mode_name = MODE_NAMES.get(result.mode, result.mode)
    lines.append(f"| 模式 | {mode_name} |")
    lines.append(f"| AI味评分 | **{result.score}/100** {result.grade} |")
    lines.append(f"| 总命中 | {result.hit_count}条 |")
    lines.append(f"| 致命/高/中/低 | {result.fatal_count}/{result.high_count}/{result.medium_count}/{result.low_count} |")
    lines.append(f"")
    lines.append(f"## 检测详情")
    lines.append(f"")

    severity_order = ["fatal", "high", "medium", "low"]
    severity_emoji = {"fatal": "🔴", "high": "🟡", "medium": "🟠", "low": "⚪"}

    for sev in severity_order:
        sev_hits = [h for h in result.hits if h.severity == sev]
        if not sev_hits:
            continue
        lines.append(f"### {severity_emoji[sev]} {sev.upper()}")
        lines.append(f"")
        for hit in sev_hits:
            lines.append(f"**第{hit.line_num}行** [{hit.rule_id}] {hit.rule_name}")
            lines.append(f"")
            lines.append(f"> 原文: `{hit.matched_text}`")
            lines.append(f">")
            lines.append(f"> 建议: {hit.suggestion}")
            if hit.example:
                lines.append(f">")
                lines.append(f"> 改前: {hit.example.get('before', '')}")
                lines.append(f"> 改后: {hit.example.get('after', '')}")
            lines.append(f"")

    return "\n".join(lines)


# ============================================================
# Word文档修订标记输出
# ============================================================

def _default_docx_output(input_path: str) -> str:
    """生成默认输出文件名"""
    if input_path == "<stdin>" or not input_path:
        return "ai_flavor_radar_output.docx"
    base, ext = os.path.splitext(input_path)
    return f"{base}_AI味雷达修订.docx"


def _generate_docx_with_track_changes(result: 'ScanResult', text: str, output_path: str):
    """生成带Word修订标记的docx文件。
    
    每条命中规则在原文对应位置生成一条修订标记：
    - w:del 标记删除命中文字
    - w:ins 插入修改建议（带规则编号）
    用户在Word中右键接受/拒绝即可。
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        print("错误: 生成docx需要python-docx，请运行 pip install python-docx", file=sys.stderr)
        sys.exit(1)

    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # 标题
    title = doc.add_heading('AI味雷达修订报告', level=1)
    title.alignment = 1  # 居中
    
    # 概要信息
    info = doc.add_paragraph()
    info.add_run(f'评分: {result.score}/100 {result.grade}  |  '
                 f'命中: {result.hit_count}条 '
                 f'(致命{result.fatal_count}/高{result.high_count}/'
                 f'中{result.medium_count}/低{result.low_count})').bold = True
    
    doc.add_paragraph('以下为原文，AI味雷达检测到的问题已用Word修订标记标出。'
                      '请在Word中右键「接受」或「拒绝」每条修订建议。')
    
    # 分割线
    doc.add_paragraph('—' * 30)
    
    # 按行处理原文
    lines = text.split('\n')
    
    # 按行号分组命中
    hits_by_line: Dict[int, List['Hit']] = {}
    for hit in result.hits:
        if hit.line_num not in hits_by_line:
            hits_by_line[hit.line_num] = []
        hits_by_line[hit.line_num].append(hit)
    
    # 每行内按col_start排序（从后往前处理，避免位置偏移）
    for line_num, line_hits in hits_by_line.items():
        line_hits.sort(key=lambda h: h.col_start, reverse=True)
    
    rev_id = [0]  # 修订ID计数器
    now = datetime.now().strftime('%Y-%m-%dT%H:%M:%SZ')
    AUTHOR = 'AI味雷达'
    
    for i, line in enumerate(lines, 1):
        p = doc.add_paragraph()
        
        if i not in hits_by_line:
            # 无命中，直接写入原文
            p.add_run(line)
            continue
        
        # 有命中：从后往前插入修订标记
        modified_line = line
        line_hit_list = hits_by_line[i]
        
        # 记录各段文字和修订标记的顺序
        segments = []  # [(text, hit_or_None)]
        cursor = 0
        
        # 按col_start正向排序，分段
        sorted_hits = sorted(line_hit_list, key=lambda h: h.col_start)
        
        for hit in sorted_hits:
            # 命中前的正常文字
            if hit.col_start > cursor:
                segments.append((line[cursor:hit.col_start], None))
            # 命中文字+修订建议
            segments.append((line[hit.col_start:hit.col_end], hit))
            cursor = hit.col_end
        
        # 末尾剩余文字
        if cursor < len(line):
            segments.append((line[cursor:], None))
        
        # 渲染segments到段落
        for seg_text, hit in segments:
            if hit is None:
                p.add_run(seg_text)
            else:
                _insert_track_change(p, seg_text, hit, rev_id, now, AUTHOR)
    
    # 末尾添加检测详情汇总
    doc.add_paragraph('')
    doc.add_paragraph('—' * 30)
    detail_heading = doc.add_heading('检测详情', level=2)
    
    severity_order = ["fatal", "high", "medium", "low"]
    severity_label = {"fatal": "🔴 致命", "high": "🟡 高危", 
                      "medium": "🟠 中危", "low": "⚪ 低危"}
    
    for sev in severity_order:
        sev_hits = [h for h in result.hits if h.severity == sev]
        if not sev_hits:
            continue
        
        doc.add_heading(severity_label[sev], level=3)
        for hit in sev_hits:
            item = doc.add_paragraph(style='List Bullet')
            item.add_run(f'第{hit.line_num}行 [{hit.rule_id}] {hit.rule_name}').bold = True
            item.add_run(f'\n原文: {hit.matched_text}')
            item.add_run(f'\n建议: {hit.suggestion}')

    # 开启Word原生修订模式（Track Changes）
    # 关键：settings.xml 必须有 <w:trackChanges/>，否则Word/WPS打开时
    # 不识别为"修订模式"，修订标记可能不显示或不可接受/拒绝
    _enable_track_changes(doc)

    doc.save(output_path)


def _enable_track_changes(doc) -> None:
    """在settings.xml中开启Word原生修订模式（Track Changes）。

    Word原生修订 = settings.xml 里的 <w:trackChanges/> 元素。
    没有它，w:ins/w:del 标记虽然存在但Word/WPS打开时不会进入修订模式，
    用户无法在审阅面板里接受/拒绝——这就是"伪修订"和"真修订"的区别。
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    settings = doc.settings._element  # w:settings
    # 检查是否已存在，避免重复
    existing = settings.find(qn('w:trackChanges'))
    if existing is None:
        track = OxmlElement('w:trackChanges')
        # 插入到 settings 的最前面（符合Word惯例：修订标记类元素在前）
        settings.insert(0, track)


def _insert_track_change(paragraph, matched_text: str, hit: 'Hit', 
                          rev_id: list, timestamp: str, author: str):
    """在段落中插入一条Track Change修订标记。
    
    效果：删除命中文字（红色删除线），插入建议文字（下划线）。
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    
    p_element = paragraph._element
    
    # 1. 创建删除标记 w:del
    del_elem = OxmlElement('w:del')
    rev_id[0] += 1
    del_elem.set(qn('w:id'), str(rev_id[0]))
    del_elem.set(qn('w:author'), author)
    del_elem.set(qn('w:date'), timestamp)
    
    del_run = OxmlElement('w:r')
    del_rpr = OxmlElement('w:rPr')
    del_color = OxmlElement('w:color')
    del_color.set(qn('w:val'), 'FF0000')
    del_strike = OxmlElement('w:strike')
    del_rpr.append(del_color)
    del_rpr.append(del_strike)
    del_run.append(del_rpr)
    
    del_t = OxmlElement('w:delText')
    del_t.text = matched_text
    del_t.set(qn('xml:space'), 'preserve')
    del_run.append(del_t)
    del_elem.append(del_run)
    
    # 2. 创建插入标记 w:ins（建议文字）
    ins_elem = OxmlElement('w:ins')
    rev_id[0] += 1
    ins_elem.set(qn('w:id'), str(rev_id[0]))
    ins_elem.set(qn('w:author'), author)
    ins_elem.set(qn('w:date'), timestamp)
    
    ins_run = OxmlElement('w:r')
    ins_rpr = OxmlElement('w:rPr')
    ins_color = OxmlElement('w:color')
    ins_color.set(qn('w:val'), '0066CC')
    ins_underline = OxmlElement('w:u')
    ins_underline.set(qn('w:val'), 'single')
    ins_rpr.append(ins_color)
    ins_rpr.append(ins_underline)
    ins_run.append(ins_rpr)
    
    # 截断建议文字，避免行内过长
    suggestion_short = hit.suggestion[:50] + '...' if len(hit.suggestion) > 50 else hit.suggestion
    ins_t = OxmlElement('w:t')
    ins_t.text = f' [{hit.rule_id}: {suggestion_short}] '
    ins_t.set(qn('xml:space'), 'preserve')
    ins_run.append(ins_t)
    ins_elem.append(ins_run)
    
    # 插入到段落末尾
    p_element.append(del_elem)
    p_element.append(ins_elem)

def main():
    parser = argparse.ArgumentParser(
        description="AI味雷达 - 只标不改的AI文风检测工具",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  %(prog)s 投标方案.md --mode bid
  %(prog)s 公众号文章.md --mode social --format markdown
  echo "值得注意的是，我们将全面提升服务质量" | %(prog)s --stdin --mode bid
  %(prog)s 文件.md --mode bid --format json > report.json
        """
    )
    parser.add_argument("file", nargs="?", help="待检测的文件路径（.md/.txt）")
    parser.add_argument("--stdin", action="store_true", help="从标准输入读取文本")
    parser.add_argument("--mode", choices=["bid", "social", "xiaohongshu", "email", "paper"], default="bid",
                        help="检测模式: bid=投标方案(默认), social=自媒体, xiaohongshu=小红书, email=邮件, paper=学术论文")
    parser.add_argument("--format", choices=["text", "json", "markdown", "docx"], default="text",
                        help="输出格式: text(默认), json, markdown, docx(Word修订标记)")
    parser.add_argument("-o", "--output", help="输出文件路径（docx格式必填）")
    parser.add_argument("--no-color", action="store_true", help="禁用终端颜色")
    parser.add_argument("--rules-dir", help="自定义规则目录路径")
    parser.add_argument("--no-examples", action="store_true", help="报告不显示改前→改后示范")

    args = parser.parse_args()

    # 读取文本
    if args.stdin:
        text = sys.stdin.read()
        file_path = "<stdin>"
    elif args.file:
        if not os.path.exists(args.file):
            print(f"错误: 文件不存在: {args.file}", file=sys.stderr)
            sys.exit(1)

        # docx格式：如果输入是docx，先提取文本
        if args.file.lower().endswith(".docx"):
            try:
                from docx import Document as DocxDocument
                doc = DocxDocument(args.file)
                text = "\n".join(p.text for p in doc.paragraphs)
            except ImportError:
                print("错误: 处理docx需要python-docx，请运行 pip install python-docx", file=sys.stderr)
                sys.exit(1)
        else:
            with open(args.file, "r", encoding="utf-8") as f:
                text = f.read()
        file_path = args.file
    else:
        parser.print_help()
        sys.exit(1)

    # 扫描
    radar = FlavorRadar(mode=args.mode, custom_rules_dir=args.rules_dir)
    result = radar.scan(text, file_path=file_path)

    # --no-examples：报告不显示示范（docx模式始终显示，便于教学）
    if args.no_examples:
        for h in result.hits:
            h.example = None

    # 输出
    use_color = not args.no_color and sys.stdout.isatty()
    if args.format == "json":
        print(format_json_report(result))
    elif args.format == "markdown":
        print(format_markdown_report(result))
    elif args.format == "docx":
        output_path = args.output or _default_docx_output(file_path)
        _generate_docx_with_track_changes(result, text, output_path)
        print(f"✅ Word修订文档已生成: {output_path}")
        print(f"   评分: {result.score}/100 {result.grade}")
        print(f"   命中: {result.hit_count}条 (致命{result.fatal_count}/高{result.high_count}/中{result.medium_count}/低{result.low_count})")
        print(f"   用Word打开后，逐条接受/拒绝修订建议即可。")
    else:
        print(format_text_report(result, use_color=use_color))


if __name__ == "__main__":
    main()


def run_scan(file_path: str, mode: str = "bid", fmt: str = "text", output: str = None, extra_args: list = None) -> dict:
    """CLI 调用的入口函数，直接返回扫描结果"""
    import sys
    old_argv = sys.argv
    sys.argv = ["ai_flavor_radar.py", file_path, "--mode", mode, "--format", fmt]
    if output:
        sys.argv.extend(["-o", output])
    if extra_args:
        sys.argv.extend(extra_args)
    try:
        main()
    except SystemExit:
        pass
    finally:
        sys.argv = old_argv
