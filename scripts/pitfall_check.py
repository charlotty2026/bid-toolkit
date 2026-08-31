#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
投标文件踩坑检查器 v1.0
======================
扫描投标文件，检测常见的踩坑问题。
基于 pitfalls.md 的12条规则 + 3条额外补充规则，共15条检查规则。

日期: 2026-07-27

用法：
    python scripts/pitfall_check.py check 投标文件.md
    python scripts/pitfall_check.py check 投标文件.docx
    python scripts/pitfall_check.py check 投标文件.md --json
"""

import argparse
import json
import re
import sys
from pathlib import Path
from datetime import date

# ============================================================
# 配置
# ============================================================

BASE_DIR = Path(__file__).resolve().parent.parent
PITFALLS_FILE = BASE_DIR / "company_profile" / "pitfalls.md"

# ============================================================
# 15条检查规则定义
# ============================================================

PITFALL_RULES: list[dict] = [
    {
        "id": 1,
        "name": "绝对化用语",
        "severity": "ERROR",
        "keywords": ["保证中标", "确保中标", "承诺中标", "必中"],
        "description": "投标文件中禁止出现'保证中标'等绝对化表述，会被废标",
        "suggestion": "删除'保证中标'，改为'具有丰富的同类项目经验'",
    },
    {
        "id": 2,
        "name": "绝对化用语",
        "severity": "ERROR",
        "keywords": ["100%成功率", "100%满意", "100%通过", "独家技术", "唯一选择", "独家供应商"],
        "description": "禁止出现'100%成功率''独家技术''唯一选择'等绝对化/排他性表述",
        "suggestion": "删除绝对化用语，改为客观描述实际能力和经验",
    },
    {
        "id": 3,
        "name": "无依据的领先表述",
        "severity": "WARN",
        "keywords": ["国内领先", "国际先进", "行业第一", "国内首创", "国际领先", "行业最优"],
        "description": "禁止出现'国内领先''国际先进'等无依据的领先表述",
        "suggestion": "删除领先表述，或提供权威机构认证/第三方数据支撑",
    },
    {
        "id": 4,
        "name": "人员冲突",
        "severity": "ERROR",
        "keywords": [r"同时担任.*项目经理", r"项目经理.*兼任", r"兼任.*项目经理"],
        "description": "项目经理同时担任多个项目经理，存在人员冲突",
        "suggestion": "确认项目经理无在建项目冲突，或更换有全职承诺的项目经理",
        "use_regex": True,
    },
    {
        "id": 5,
        "name": "业绩虚报",
        "severity": "ERROR",
        "keywords": [r"业绩.*\d+.*项", r"已完成.*\d+.*个.*项目"],
        "description": "业绩数量需与实际附件证明材料一致，否则视为业绩虚报",
        "suggestion": "核实业绩数量，确保每项业绩都有对应的合同/验收证明",
        "use_regex": True,
    },
    {
        "id": 6,
        "name": "虚假承诺",
        "severity": "WARN",
        "keywords": [r"\d+小时内响应", r"\d+分钟内到达", r"即时.*响应"],
        "description": "响应时间优于招标要求但无实施依据，可能构成虚假承诺",
        "suggestion": "确认响应时间承诺有人员/制度支撑，或调整为可实现的承诺",
        "use_regex": True,
    },
    {
        "id": 7,
        "name": "资质失效",
        "severity": "WARN",
        "keywords": [],
        "description": "资质证书过期，引用已失效的资质证书",
        "suggestion": "更新资质证书或删除过期证书引用",
        "custom_check": "check_cert_expiry",
    },
    {
        "id": 8,
        "name": "异常低价",
        "severity": "WARN",
        "keywords": [r"低于成本", "零利润", r"成本价.*报价"],
        "description": "报价低于成本价，可能被认定为异常低价",
        "suggestion": "核实报价构成，确保报价不低于成本价",
        "use_regex": True,
    },
    {
        "id": 9,
        "name": "暗标泄密",
        "severity": "ERROR",
        "keywords": ["我公司名称", "本公司地址", "公司电话", "法定代表人姓名"],
        "description": "暗标中包含公司信息，会导致废标",
        "suggestion": "删除暗标部分的公司名称、地址、电话等标识信息",
    },
    {
        "id": 10,
        "name": "套模板痕迹",
        "severity": "WARN",
        "keywords": ["城商行", "其他项目", "原项目", "上一项目"],
        "description": "方案内容与项目无关，存在套模板痕迹",
        "suggestion": "删除与其他项目相关的残留内容，确保方案针对本项目定制",
    },
    {
        "id": 11,
        "name": "超范围经营",
        "severity": "ERROR",
        "keywords": [r"承诺.*超出.*经营范围", r"经营范围.*不包括"],
        "description": "承诺超出经营范围，属于超范围经营",
        "suggestion": "核实承诺事项是否在营业执照经营范围内",
        "use_regex": True,
    },
    {
        "id": 12,
        "name": "法规失效",
        "severity": "WARN",
        "keywords": ["已废止", r"失效.*法规", r"原.*规定.*已.*废止"],
        "description": "引用已废止法规，会导致方案无效",
        "suggestion": "更新引用的法规为最新有效版本",
        "use_regex": True,
    },
    {
        "id": 13,
        "name": "人员复用冲突",
        "severity": "WARN",
        "keywords": [],
        "description": "同一人员出现在不同岗位，存在人员复用冲突",
        "suggestion": "检查人员配置表，确保同一人未在多个岗位重复出现",
        "custom_check": "check_personnel_reuse",
    },
    {
        "id": 14,
        "name": "时间矛盾",
        "severity": "WARN",
        "keywords": [],
        "description": "时间节点矛盾，如承诺到岗时间与合同要求不一致",
        "suggestion": "核实所有时间承诺，确保与招标文件要求一致",
        "custom_check": "check_time_conflict",
    },
    {
        "id": 15,
        "name": "金额错误",
        "severity": "ERROR",
        "keywords": [],
        "description": "金额大小写不一致，属于金额错误",
        "suggestion": "核对金额大小写是否一致",
        "custom_check": "check_amount_mismatch",
    },
]


# ============================================================
# 文件读取（支持MD/TXT和DOCX）
# ============================================================

def read_file(file_path: str) -> list[str]:
    """读取输入文件，返回行列表。支持 .md/.txt/.docx 格式。"""
    path = Path(file_path)
    if not path.exists():
        print(f"错误：文件不存在：{file_path}", file=sys.stderr)
        sys.exit(1)

    suffix = path.suffix.lower()
    if suffix in (".md", ".txt"):
        with open(path, "r", encoding="utf-8") as f:
            return f.read().split("\n")
    elif suffix == ".docx":
        return _read_docx(path)
    else:
        print(f"错误：不支持的文件格式：{suffix}", file=sys.stderr)
        print("支持：.md / .txt / .docx", file=sys.stderr)
        sys.exit(1)


def _read_docx(path: Path) -> list[str]:
    """读取DOCX文件，返回行列表（含表格内容）。"""
    try:
        from docx import Document
    except ImportError:
        print("错误：需要 python-docx：pip install python-docx", file=sys.stderr)
        sys.exit(1)

    doc = Document(str(path))
    lines: list[str] = []
    for para in doc.paragraphs:
        lines.append(para.text)
    # 读取表格中的文本
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text for cell in row.cells)
            lines.append(row_text)
    return lines


# ============================================================
# 自定义检查函数（规则7/13/14/15）
# ============================================================

def check_cert_expiry(lines: list[str]) -> list[dict]:
    """规则7: 检查资质证书有效期是否过期。"""
    issues: list[dict] = []
    today = date.today()
    date_pattern = r"(\d{4})[-/年](\d{1,2})[-/月](\d{1,2})日?"
    cert_keywords = ["ISO", "认证", "许可证", "资质", "证书", "营业执照"]

    for i, line in enumerate(lines, 1):
        if not any(kw in line for kw in cert_keywords):
            continue
        for m in re.finditer(date_pattern, line):
            try:
                y, mo, d = int(m.group(1)), int(m.group(2)), int(m.group(3))
                cert_date = date(y, mo, d)
                if cert_date < today:
                    issues.append({
                        "line": i,
                        "matched": m.group(0),
                        "context": line.strip()[:120],
                        "detail": f"发现证书有效期{m.group(0)}已过期",
                    })
            except ValueError:
                pass
    return issues


def check_personnel_reuse(lines: list[str]) -> list[dict]:
    """规则13: 检查同一人员是否出现在不同岗位。"""
    # 匹配 "姓名：XXX  岗位：YYY" 或 "XXX—项目经理" 等模式
    patterns = [
        (r"姓名[：:]\s*([\u4e00-\u9fa5]{2,4})\s*[,，\s]*岗位[：:]\s*(.+?)(?:[,，\s]|$)", "岗位"),
        (r"姓名[：:]\s*([\u4e00-\u9fa5]{2,4})\s*[,，\s]*职务[：:]\s*(.+?)(?:[,，\s]|$)", "职务"),
        (r"([\u4e00-\u9fa5]{2,4})\s*[—\-|]\s*(项目经理|技术负责人|安全员|质量员|施工员|资料员)", "岗位"),
    ]

    name_positions: dict[str, list[tuple[int, str]]] = {}
    for i, line in enumerate(lines, 1):
        for pat, _label in patterns:
            for m in re.finditer(pat, line):
                name = m.group(1)
                position = m.group(2).strip()
                name_positions.setdefault(name, []).append((i, position))

    issues: list[dict] = []
    seen_names: set[str] = set()
    for name, occurrences in name_positions.items():
        positions = {pos for _, pos in occurrences}
        if len(positions) > 1 and name not in seen_names:
            seen_names.add(name)
            first_line, first_pos = occurrences[0]
            issues.append({
                "line": first_line,
                "matched": name,
                "context": f"'{name}'出现在{len(positions)}个不同岗位: {', '.join(positions)}",
                "detail": f"同一人员'{name}'出现在不同岗位",
            })
    return issues


def check_time_conflict(lines: list[str]) -> list[dict]:
    """规则14: 检查时间节点矛盾（同类承诺出现不同天数）。"""
    issues: list[dict] = []
    time_patterns = [
        (r"(\d+)\s*个?工作?日内?\s*(?:到岗|进场|开工)", "到岗"),
        (r"(\d+)\s*个?工作?日内?\s*(?:完成|交付|竣工)", "完成"),
        (r"承诺\s*(\d+)\s*天", "承诺天数"),
    ]

    found_times: dict[str, list[tuple[int, int, str]]] = {}
    for i, line in enumerate(lines, 1):
        for pat, label in time_patterns:
            for m in re.finditer(pat, line):
                days = int(m.group(1))
                found_times.setdefault(label, []).append((i, days, line.strip()[:120]))

    for label, times in found_times.items():
        if len(times) < 2:
            continue
        day_values = [t[1] for t in times]
        if max(day_values) != min(day_values):
            first = times[0]
            issues.append({
                "line": first[0],
                "matched": label,
                "context": first[2],
                "detail": f"'{label}'时间承诺矛盾: 出现{day_values}",
            })
    return issues


def check_amount_mismatch(lines: list[str]) -> list[dict]:
    """规则15: 检查金额大小写不一致。"""
    issues: list[dict] = []
    cn_num_map = {
        "零": 0, "壹": 1, "贰": 2, "叁": 3, "肆": 4, "伍": 5,
        "陆": 6, "柒": 7, "捌": 8, "玖": 9, "拾": 10, "佰": 100,
        "仟": 1000, "万": 10000, "亿": 100000000, "圆": 0, "元": 0, "整": 0,
    }

    def _cn_to_int(cn: str) -> int:
        total = 0
        section = 0
        current = 0
        for ch in cn:
            if ch not in cn_num_map:
                continue
            val = cn_num_map[ch]
            if val >= 10000:
                section += current
                total += section * val
                section = 0
                current = 0
            elif val >= 10:
                if current == 0:
                    current = 1
                section += current * val
                current = 0
            else:
                current = val
        return total + section + current

    for i, line in enumerate(lines, 1):
        cn_match = re.search(r"[壹贰叁肆伍陆柒捌玖拾佰仟万亿圆元整]{4,}", line)
        num_match = re.search(r"(\d[\d,]*(?:\.\d+)?)\s*元", line)
        if cn_match and num_match:
            cn_val = _cn_to_int(cn_match.group(0))
            num_val = float(num_match.group(1).replace(",", ""))
            if cn_val > 0 and abs(cn_val - num_val) > 1:
                issues.append({
                    "line": i,
                    "matched": f"大写:{cn_match.group(0)} 小写:{num_match.group(1)}元",
                    "context": line.strip()[:120],
                    "detail": f"金额大小写不一致: 大写≈{cn_val}元, 小写={num_val}元",
                })
    return issues


# ============================================================
# 自定义规则加载（从 pitfalls.md）
# ============================================================

def load_custom_pitfalls() -> list[dict]:
    """从 pitfalls.md 加载用户自定义踩坑规则。

    文件格式：编号|关键词|问题描述（每行一条）
    """
    if not PITFALLS_FILE.exists():
        return []

    custom_rules: list[dict] = []
    with open(PITFALLS_FILE, "r", encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if not line or line.startswith("#") or line.startswith("<!--") or line.startswith("-->"):
                continue
            parts = line.split("|")
            if len(parts) >= 3 and parts[0].strip().isdigit():
                custom_rules.append({
                    "id": int(parts[0].strip()),
                    "keyword": parts[1].strip(),
                    "description": parts[2].strip(),
                })
    return custom_rules


def run_custom_check(lines: list[str], custom_rules: list[dict]) -> list[dict]:
    """执行用户自定义踩坑规则检查。"""
    issues: list[dict] = []
    for rule in custom_rules:
        kw = rule["keyword"]
        for i, line in enumerate(lines, 1):
            if kw in line:
                issues.append({
                    "rule_id": f"C{rule['id']}",
                    "rule_name": "自定义踩坑规则",
                    "severity": "WARN",
                    "line": i,
                    "matched": kw,
                    "context": line.strip()[:120],
                    "detail": f"第{i}行命中自定义规则'{kw}'",
                    "suggestion": rule["description"],
                })
    return issues


# ============================================================
# 核心扫描逻辑
# ============================================================

def run_check(lines: list[str]) -> list[dict]:
    """执行全部15条内置规则检查，返回问题列表。"""
    all_issues: list[dict] = []

    for rule in PITFALL_RULES:
        # 自定义检查函数
        if rule.get("custom_check"):
            check_fn = globals().get(rule["custom_check"])
            if check_fn:
                for iss in check_fn(lines):
                    all_issues.append({
                        "rule_id": rule["id"],
                        "rule_name": rule["name"],
                        "severity": rule["severity"],
                        "line": iss["line"],
                        "matched": iss.get("matched", ""),
                        "context": iss.get("context", ""),
                        "detail": iss.get("detail", rule["description"]),
                        "suggestion": rule["suggestion"],
                    })
            continue

        # 关键词匹配
        keywords = rule.get("keywords", [])
        use_regex = rule.get("use_regex", False)
        for i, line in enumerate(lines, 1):
            for kw in keywords:
                if use_regex:
                    try:
                        matches = list(re.finditer(kw, line))
                    except re.error:
                        matches = []
                    for m in matches:
                        all_issues.append({
                            "rule_id": rule["id"],
                            "rule_name": rule["name"],
                            "severity": rule["severity"],
                            "line": i,
                            "matched": m.group(0),
                            "context": line.strip()[:120],
                            "detail": f"第{i}行发现'{m.group(0)}'",
                            "suggestion": rule["suggestion"],
                        })
                else:
                    if kw in line:
                        all_issues.append({
                            "rule_id": rule["id"],
                            "rule_name": rule["name"],
                            "severity": rule["severity"],
                            "line": i,
                            "matched": kw,
                            "context": line.strip()[:120],
                            "detail": f"第{i}行发现'{kw}'",
                            "suggestion": rule["suggestion"],
                        })

    return all_issues


# ============================================================
# 报告输出
# ============================================================

def print_text_report(file_path: str, issues: list[dict], custom_count: int) -> None:
    """打印文本格式报告。"""
    total_rules = len(PITFALL_RULES) + custom_count
    error_count = sum(1 for i in issues if i["severity"] == "ERROR")
    warn_count = sum(1 for i in issues if i["severity"] == "WARN")

    print("踩坑检查报告")
    print("=" * 60)
    print(f"检查文件: {file_path}")
    print(f"检查规则: {total_rules}条")
    print()

    if not issues:
        print("未发现踩坑问题，检查通过！")
        print()
        print(f"汇总: 0个问题")
        return

    print("问题列表:")
    print()
    for iss in issues:
        tag = f"[{iss['severity']}]"
        print(f"{tag} 规则{iss['rule_id']}-{iss['rule_name']}: {iss['detail']}")
        if iss.get("context"):
            print(f"  原文: {iss['context'][:80]}")
        if iss.get("suggestion"):
            print(f"  建议: {iss['suggestion']}")
        print()

    print(f"汇总: {len(issues)}个问题 ({error_count}个ERROR, {warn_count}个WARN)")


def print_json_report(file_path: str, issues: list[dict], custom_count: int) -> None:
    """打印JSON格式报告。"""
    total_rules = len(PITFALL_RULES) + custom_count
    error_count = sum(1 for i in issues if i["severity"] == "ERROR")
    warn_count = sum(1 for i in issues if i["severity"] == "WARN")

    report = {
        "检查文件": file_path,
        "检查规则数": total_rules,
        "问题数": len(issues),
        "ERROR数": error_count,
        "WARN数": warn_count,
        "问题列表": issues,
    }
    print(json.dumps(report, ensure_ascii=False, indent=2))


# ============================================================
# CLI入口
# ============================================================

def main() -> None:
    parser = argparse.ArgumentParser(
        description="投标文件踩坑检查器 v1.0",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
    python scripts/pitfall_check.py check 投标文件.md
    python scripts/pitfall_check.py check 投标文件.docx
    python scripts/pitfall_check.py check 投标文件.md --json
        """,
    )
    subparsers = parser.add_subparsers(dest="command", help="子命令")

    check_parser = subparsers.add_parser("check", help="检查投标文件")
    check_parser.add_argument("file", help="投标文件路径（.md/.txt/.docx）")
    check_parser.add_argument("--json", action="store_true", help="输出JSON格式报告")

    args = parser.parse_args()

    if args.command == "check":
        lines = read_file(args.file)

        # 内置15条规则检查
        issues = run_check(lines)

        # 用户自定义规则检查（从 pitfalls.md 加载）
        custom_rules = load_custom_pitfalls()
        if custom_rules:
            issues.extend(run_custom_check(lines, custom_rules))

        if args.json:
            print_json_report(args.file, issues, len(custom_rules))
        else:
            print_text_report(args.file, issues, len(custom_rules))
    else:
        parser.print_help()


if __name__ == "__main__":
    main()
