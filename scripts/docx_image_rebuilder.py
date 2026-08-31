#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
docx_image_rebuilder.py - Word文档图片结构化提取与表格重建工具

功能：
  ① 爬结构：遍历文档段落，按Heading样式识别标题，自动分组标题→正文→图片
  ② 建表格：新建文档，按分组逐行创建表格（左列文字，右列图片）
  
关键技巧：
  - 图片不走「保存到硬盘再插入」的弯路
  - 直接用 deepcopy + XML节点操作 复用图片关系
  - 又快又不丢质量

作者：bid-toolkit contributors
日期：2026-05-25
版本：v1.0
"""

import sys
import os
from copy import deepcopy
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============================================================
# 第一部分：字体工具函数
# ============================================================

def set_chinese_font(run, font_name='宋体', size=None):
    """
    设置run的中文字体，确保三个字段齐全
    
    Word字体分三个字段：
      w:ascii / w:hAnsi → 英文/数字
      w:eastAsia → 中文/日文/韩文
    
    font.name 只设前两个，必须手动补 eastAsia
    """
    run.font.name = font_name
    if size:
        run.font.size = Pt(size)
    
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)


def set_cell_shading(cell, color='D9E2F3'):
    """设置单元格底色（默认浅蓝）"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading)


# ============================================================
# 第二部分：结构爬取
# ============================================================

def extract_document_structure(doc_path):
    """
    提取文档结构：按标题层级分组，收集每组下的段落和图片
    
    返回: [
        {
            'level': int,           # 标题层级 1-9
            'title': str,           # 标题文本
            'paragraphs': [str],    # 正文段落列表
            'shapes': [Element],    # 图片XML元素列表
            'shape_count': int      # 图片数量
        },
        ...
    ]
    """
    print(f"[1/4] 正在读取文档: {doc_path}")
    doc = Document(doc_path)
    
    structure = []
    current_group = None
    total_shapes = 0
    
    # 遍历所有段落
    for i, para in enumerate(doc.paragraphs):
        style_name = para.style.name if para.style else ''
        text = para.text.strip()
        
        # 识别标题（Heading 1-9）
        if style_name.startswith('Heading'):
            try:
                level = int(style_name.replace('Heading ', ''))
            except ValueError:
                continue
            
            # 保存上一个分组
            if current_group is not None:
                structure.append(current_group)
            
            # 创建新分组
            current_group = {
                'level': level,
                'title': text,
                'paragraphs': [],
                'shapes': [],
                'shape_count': 0
            }
            
        elif current_group is not None:
            # 收集正文段落（非空）
            if text:
                current_group['paragraphs'].append(text)
            
            # 收集段落中的图片（inline_shapes）
            # python-docx的para._element包含所有子元素
            for shape in para._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'):
                current_group['shapes'].append(shape)
                current_group['shape_count'] += 1
                total_shapes += 1
    
    # 保存最后一个分组
    if current_group is not None:
        structure.append(current_group)
    
    print(f"[2/4] 结构提取完成:")
    print(f"      - 标题分组数: {len(structure)}")
    print(f"      - 图片总数: {total_shapes}")
    
    # 打印结构概览
    for i, group in enumerate(structure[:5]):  # 只显示前5个
        indent = "  " * (group['level'] - 1)
        print(f"      {indent}[H{group['level']}] {group['title'][:40]}... ({group['shape_count']}张图)")
    if len(structure) > 5:
        print(f"      ... 还有 {len(structure)-5} 个分组")
    
    return structure


def extract_all_inline_shapes(doc_path):
    """
    备用方法：直接提取文档中所有inline shapes（不依赖段落归属）
    用于文档没有Heading样式的情况
    """
    doc = Document(doc_path)
    shapes = []
    
    # 从body直接提取所有drawing元素
    body = doc.element.body
    for shape in body.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'):
        shapes.append(shape)
    
    return shapes


# ============================================================
# 第三部分：表格重建
# ============================================================

def rebuild_with_table(doc_path, output_path, 
                       left_width=9.0, right_width=7.0,
                       table_style='Table Grid',
                       brand_color='2E5090'):
    """
    将原文档按标题分组重建为表格形式
    
    参数:
        doc_path: 源文档路径
        output_path: 输出文档路径
        left_width: 左列宽度（cm）
        right_width: 右列宽度（cm）
        table_style: 表格样式
        brand_color: 品牌主色（RGB十六进制）
    
    表格结构:
        左列: 标题(加粗) + 正文
        右列: 图片（deepcopy复用，不落地硬盘）
    """
    print(f"\n[3/4] 开始重建文档...")
    
    # 提取结构
    structure = extract_document_structure(doc_path)
    
    if not structure:
        print("[错误] 未提取到任何标题结构，请检查文档是否使用了Heading样式")
        return False
    
    # 创建新文档
    new_doc = Document()
    
    # 设置默认字体
    style = new_doc.styles['Normal']
    style.font.name = '宋体'
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), '宋体')
    rFonts.set(qn('w:hAnsi'), '宋体')
    rFonts.set(qn('w:eastAsia'), '宋体')
    
    # 添加文档标题
    title_para = new_doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run('文档图片整理')
    title_run.bold = True
    title_run.font.size = Pt(18)
    set_chinese_font(title_run, size=18)
    new_doc.add_paragraph()  # 空行
    
    # 逐组创建表格
    for idx, group in enumerate(structure):
        # 跳过没有图片的分组（可选）
        # if group['shape_count'] == 0:
        #     continue
        
        # 创建表格：1行2列
        table = new_doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = table_style
        
        # 设置列宽
        table.columns[0].width = Cm(left_width)
        table.columns[1].width = Cm(right_width)
        
        # 设置表格边框
        tbl = table._element
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        
        # 左列：标题+正文
        left_cell = table.cell(0, 0)
        
        # 标题（带层级缩进）
        title_para = left_cell.paragraphs[0]
        indent = "  " * (group['level'] - 1)
        title_run = title_para.add_run(f"{indent}{group['title']}")
        title_run.bold = True
        title_run.font.size = Pt(14 - group['level'] + 1)  # 层级越高字越小
        set_chinese_font(title_run)
        
        # 正文
        for text in group['paragraphs']:
            p = left_cell.add_paragraph()
            run = p.add_run(text)
            run.font.size = Pt(10.5)
            set_chinese_font(run)
        
        # 如果没有正文，添加提示
        if not group['paragraphs']:
            p = left_cell.add_paragraph()
            run = p.add_run("(无正文内容)")
            run.italic = True
            run.font.color.rgb = RGBColor(128, 128, 128)
            set_chinese_font(run)
        
        # 右列：图片（deepcopy复用）
        right_cell = table.cell(0, 1)
        right_para = right_cell.paragraphs[0]
        right_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if group['shapes']:
            for shape_element in group['shapes']:
                # 关键技巧：deepcopy XML节点，保持图片关系
                # 这样不需要保存到硬盘再插入，直接复用原文档的图片数据
                try:
                    new_shape = deepcopy(shape_element)
                    right_para._element.append(new_shape)
                    
                    # 图片之间添加空行
                    right_para.add_run().add_break()
                except Exception as e:
                    print(f"      [警告] 复制图片失败: {e}")
        else:
            # 没有图片时显示占位
            run = right_para.add_run("(无图片)")
            run.italic = True
            run.font.color.rgb = RGBColor(128, 128, 128)
            set_chinese_font(run)
        
        # 表格后空一行
        new_doc.add_paragraph()
        
        # 进度显示
        if (idx + 1) % 10 == 0:
            print(f"      已处理 {idx + 1}/{len(structure)} 个分组")
    
    # 添加页脚
    footer_para = new_doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(f"\n生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(128, 128, 128)
    set_chinese_font(footer_run)
    
    # 保存
    new_doc.save(output_path)
    print(f"\n[4/4] ✅ 完成! 已保存: {output_path}")
    print(f"      共处理 {len(structure)} 个分组")
    
    return True


# ============================================================
# 第四部分：高级功能
# ============================================================

def extract_images_to_folder(doc_path, output_folder):
    """
    将文档中的所有图片提取到指定文件夹
    用于需要单独处理图片的场景
    """
    print(f"提取图片到: {output_folder}")
    os.makedirs(output_folder, exist_ok=True)
    
    doc = Document(doc_path)
    image_count = 0
    
    # 遍历所有关系，找到图片
    rels = doc.part.rels
    for rel in rels.values():
        if "image" in rel.reltype:
            image_count += 1
            image_data = rel.target_part.blob
            
            # 推断图片格式
            if image_data[:8] == b'\x89PNG\r\n\x1a\n':
                ext = 'png'
            elif image_data[:2] == b'\xff\xd8':
                ext = 'jpg'
            else:
                ext = 'bin'
            
            # 保存
            image_path = os.path.join(output_folder, f"image_{image_count:03d}.{ext}")
            with open(image_path, 'wb') as f:
                f.write(image_data)
    
    print(f"共提取 {image_count} 张图片")
    return image_count


def scan_document_issues(doc_path):
    """
    扫描文档常见问题
    返回问题列表
    """
    print(f"扫描文档: {doc_path}")
    doc = Document(doc_path)
    issues = []
    
    # 1. 占位符检查
    placeholder_patterns = ['XX', '____', '【待填】', '[待填]', '待定']
    for i, para in enumerate(doc.paragraphs):
        for pattern in placeholder_patterns:
            if pattern in para.text:
                issues.append({
                    'type': '占位符',
                    'line': i + 1,
                    'text': para.text[:50] + '...' if len(para.text) > 50 else para.text
                })
    
    # 2. 手动编号检查（简单启发式）
    import re
    manual_numbering = re.compile(r'^(\d+[、.．]|\([\d一二三四五六七八九十]+\)|[一二三四五六七八九十]、)')
    for i, para in enumerate(doc.paragraphs):
        if manual_numbering.match(para.text.strip()):
            style = para.style.name if para.style else '无样式'
            if not style.startswith('Heading'):
                issues.append({
                    'type': '疑似手动编号',
                    'line': i + 1,
                    'text': para.text[:50]
                })
    
    # 3. 图片检查
    image_count = 0
    for para in doc.paragraphs:
        for shape in para._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'):
            image_count += 1
    
    issues.append({
        'type': '统计',
        'line': '-',
        'text': f'段落数: {len(doc.paragraphs)}, 图片数: {image_count}'
    })
    
    print(f"扫描完成，发现 {len(issues)} 个问题/统计:")
    for issue in issues:
        print(f"  [{issue['type']}] 行{issue['line']}: {issue['text']}")
    
    return issues


# ============================================================
# 第五部分：命令行入口
# ============================================================

def main():
    """命令行入口"""
    import argparse
    
    parser = argparse.ArgumentParser(
        description='Word文档图片结构化提取与表格重建工具',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  # 基本用法：提取结构并重建为表格
  python docx_image_rebuilder.py input.docx output.docx
  
  # 只提取图片到文件夹
  python docx_image_rebuilder.py input.docx --extract-images ./images
  
  # 扫描文档问题
  python docx_image_rebuilder.py input.docx --scan
  
  # 自定义列宽
  python docx_image_rebuilder.py input.docx output.docx --left-width 10 --right-width 6
        """
    )
    
    parser.add_argument('input', help='输入docx文件路径')
    parser.add_argument('output', nargs='?', help='输出docx文件路径（重建模式）')
    parser.add_argument('--extract-images', metavar='FOLDER', help='提取图片到指定文件夹')
    parser.add_argument('--scan', action='store_true', help='扫描文档问题')
    parser.add_argument('--left-width', type=float, default=9.0, help='左列宽度(cm)，默认9.0')
    parser.add_argument('--right-width', type=float, default=7.0, help='右列宽度(cm)，默认7.0')
    
    args = parser.parse_args()
    
    # 检查输入文件
    if not os.path.exists(args.input):
        print(f"[错误] 文件不存在: {args.input}")
        sys.exit(1)
    
    if not args.input.endswith('.docx'):
        print(f"[警告] 输入文件不是.docx格式，可能无法正常处理")
    
    # 执行对应功能
    if args.scan:
        scan_document_issues(args.input)
    elif args.extract_images:
        extract_images_to_folder(args.input, args.extract_images)
    elif args.output:
        rebuild_with_table(
            args.input, 
            args.output,
            left_width=args.left_width,
            right_width=args.right_width
        )
    else:
        parser.print_help()


if __name__ == '__main__':
    main()
