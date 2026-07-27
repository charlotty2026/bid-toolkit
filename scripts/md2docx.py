#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Markdown → 标书 Word 一键生成（md2docx）
========================================
将 Markdown 标书草稿转换为符合排版规范的 Word 文档。
附带全角/半角标点扫描与自动修复（--scan / --no-auto-fix）。

注意：本工具是「转换器」，不是 docx 格式检查器。
如需检查已有 Word 文件的排版格式，请使用 bid_engine.py --check。
========================================
基于标书排版实战经验
吸收实战踩坑经验 v1.0
"""

import os, sys, re, argparse
from datetime import datetime
from pathlib import Path

sys.stdout.reconfigure(encoding='utf-8')

try:
    from docx import Document
    from docx.shared import Pt, Inches, Twips
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("❌ 需要 python-docx: pip install python-docx")
    sys.exit(1)

try:
    import markdown
    from markdown.extensions.tables import TableExtension
except ImportError:
    print("❌ 需要 markdown: pip install markdown")
    sys.exit(1)

# ============================================================
#  字体工具（三字段齐设，防MS Gothic回退）
# ============================================================

def set_run_song(run, size=12, bold=False):
    """设置宋体：三字段齐设（ascii / hAnsi / eastAsia），防MS Gothic"""
    run.font.name = '宋体'
    run.font.size = Pt(size)
    run.font.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), '宋体')
    rFonts.set(qn('w:hAnsi'), '宋体')
    rFonts.set(qn('w:eastAsia'), '宋体')  # 最关键！
    # 清除主题字体引用，防回退
    for attr in ['w:asciiTheme', 'w:hAnsiTheme', 'w:eastAsiaTheme']:
        try:
            del rFonts.attrib[qn(attr)]
        except KeyError:
            pass

def set_run_hei(run, size=16, bold=True):
    """设置黑体"""
    run.font.name = '黑体'
    run.font.size = Pt(size)
    run.font.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    for key in ['w:ascii', 'w:hAnsi', 'w:eastAsia']:
        rFonts.set(qn(key), '黑体')
    for attr in ['w:asciiTheme', 'w:hAnsiTheme', 'w:eastAsiaTheme']:
        try:
            del rFonts.attrib[qn(attr)]
        except KeyError:
            pass

def init_doc_styles(doc):
    """样式定义级别强制宋体"""
    for level in range(1, 6):
        try:
            style = doc.styles[f'Heading {level}']
            style.font.name = '宋体'
            rPr = style.element.get_or_add_rPr()
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rPr.insert(0, rFonts)
            rFonts.set(qn('w:ascii'), '宋体')
            rFonts.set(qn('w:hAnsi'), '宋体')
            rFonts.set(qn('w:eastAsia'), '宋体')
            for attr in ['w:asciiTheme', 'w:hAnsiTheme', 'w:eastAsiaTheme']:
                try:
                    del rFonts.attrib[qn(attr)]
                except KeyError:
                    pass
        except KeyError:
            pass
    try:
        style = doc.styles['Normal']
        style.font.name = '宋体'
        style.font.size = Pt(12)
    except KeyError:
        pass

# ============================================================
#  标准标题层级
# ============================================================

HEADING_MAP = {
    1: {'text': None, 'font': set_run_song, 'size': 16, 'bold': True, 'style': 'Heading 1'},
    2: {'text': None, 'font': set_run_song, 'size': 15, 'bold': True, 'style': 'Heading 2'},
    3: {'text': None, 'font': set_run_song, 'size': 14, 'bold': True, 'style': 'Heading 3'},
    4: {'text': None, 'font': set_run_song, 'size': 14, 'bold': True, 'style': 'Heading 4'},
}

def add_heading(doc, text, level):
    """添加标准标题"""
    h = HEADING_MAP.get(level, HEADING_MAP[4])
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        h['font'](run, h['size'], h['bold'])
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def add_body(doc, text):
    """添加正文：宋体小四，1.5倍行距，首行缩进2字符(480twips)，两端对齐"""
    p = doc.add_paragraph(text, style='Normal')
    for run in p.runs:
        set_run_song(run, 12, False)
    # 首行缩进2字符：2 × 12磅 × 20 = 480 twips
    p.paragraph_format.first_line_indent = Twips(480)
    p.paragraph_format.line_spacing = 1.5
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

def add_table(doc, headers, rows):
    """添加标准表格：五号宋体，居中对齐"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    # 表头
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for para in cell.paragraphs:
            for run in para.runs:
                set_run_song(run, 10.5, True)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 数据行
    for row in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = val
            for para in row_cells[i].paragraphs:
                for run in para.runs:
                    set_run_song(run, 10.5, False)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.first_line_indent = Twips(0)  # 表格无缩进
    return table

# ============================================================
#  全角半角检测+自动修复（标书铁律·致命级痛点）
# ============================================================

# 中文标点 → 全角
CN_PUNCT_HALF_TO_FULL = {
    ',': '，', '.': '。', '!': '！', '?': '？', ':': '：', ';': '；',
    '(': '（', ')': '）', '[': '【', ']': '】', '<': '《', '>': '》',
    '"': '"', "'": "'",
}
# 英文标点 → 半角（中文语境下的英文标点应为半角）
EN_PUNCT_FULL_TO_HALF = {
    '，': ',', '。': '.', '！': '!', '？': '?', '：': ':', '；': ';',
    '（': '(', '）': ')', '【': '[', '】': ']', '《': '<', '》': '>',
}
# 全角数字/字母 → 半角
FULLWIDTH_DIGITS = str.maketrans(
    '０１２３４５６７８９ＡＢＣＤＥＦＧＨＩＪＫＬＭＮＯＰＱＲＳＴＵＶＷＸＹＺａｂｃｄｅｆｇｈｉｊｋｌｍｎｏｐｑｒｓｔｕｖｗｘｙｚ',
    '0123456789ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz'
)

def fix_fullwidth_halfwidth(text):
    """
    自动修复全角半角混用：
    1. 中文上下文中，中文标点统一全角
    2. 英文/数字统一半角
    3. 返回 (修复后文本, 修改次数)
    """
    fixed = text
    changes = 0
    
    # 1. 全角数字/字母 → 半角
    old_len = len(fixed)
    fixed = fixed.translate(FULLWIDTH_DIGITS)
    if len(fixed) != old_len:
        changes += 1
    
    # 2. 逐字符判断：中文语境下的标点修复
    result = []
    in_cn = True  # 默认中文语境
    for ch in fixed:
        if ch in CN_PUNCT_HALF_TO_FULL and in_cn:
            result.append(CN_PUNCT_HALF_TO_FULL[ch])
            changes += 1
        elif ch in EN_PUNCT_FULL_TO_HALF and not in_cn:
            result.append(EN_PUNCT_FULL_TO_HALF[ch])
            changes += 1
        else:
            result.append(ch)
        # 切换语境：遇到中文字符进入中文语境
        if '\u4e00' <= ch <= '\u9fff' or '\u3000' <= ch <= '\u303f':
            in_cn = True
        elif ch.isascii() and ch.isalpha():
            in_cn = False
    
    return ''.join(result), changes

def scan_fullwidth_issues(text):
    """扫描全角半角问题，返回问题列表"""
    issues = []
    lines = text.split('\n')
    for i, line in enumerate(lines):
        # 检查中文里混入了半角标点
        for half, full in CN_PUNCT_HALF_TO_FULL.items():
            # 只在中文语境中检查（标点前后有中文）
            for m in re.finditer(rf'[\u4e00-\u9fff]{re.escape(half)}[\u4e00-\u9fff]', line):
                issues.append({
                    'line': i+1, 'type': '半角标点混入中文',
                    'char': half, 'should_be': full,
                    'context': line[max(0,m.start()-5):m.end()+5]
                })
        # 检查全角数字/字母
        for m in re.finditer(r'[\uff10-\uff19\uff21-\uff3a\uff41-\uff5a]+', line):
            issues.append({
                'line': i+1, 'type': '全角数字/字母',
                'char': m.group(),
                'context': line[max(0,m.start()-5):m.end()+5]
            })
    return issues

# ============================================================
#  Markdown 解析 → Word
# ============================================================

def md_to_docx(md_text, output_path, auto_fix=True):
    """Markdown 文本 -> 标准格式 Word 文档"""
    doc = Document()
    init_doc_styles(doc)

    # 🔧 全角半角自动修复（负责人点名的大痛点）
    if auto_fix:
        fixed_text, changes = fix_fullwidth_halfwidth(md_text)
        if changes > 0:
            print(f'🔧 自动修复 {changes} 处全角半角混用')
            md_text = fixed_text

    lines = md_text.split('\n')
    i = 0
    in_table = False
    table_headers = []
    table_rows = []

    while i < len(lines):
        line = lines[i].rstrip()

        # ── 表格 ──
        if line.startswith('|') and line.endswith('|'):
            cells = [c.strip() for c in line.strip('|').split('|')]
            if not in_table:
                table_headers = cells
                table_rows = []
                in_table = True
                # 跳过分隔行
                if i + 1 < len(lines) and re.match(r'^[\s|\-:]+$', lines[i+1]):
                    i += 1
            else:
                table_rows.append(cells)
            i += 1
            continue

        # 表格结束，输出
        if in_table:
            add_table(doc, table_headers, table_rows)
            in_table = False
            continue

        # ── 标题 ──
        h_match = re.match(r'^(#{1,4})\s+(.+)$', line)
        if h_match:
            level = len(h_match.group(1))
            text = h_match.group(2).strip()
            add_heading(doc, text, level)
            i += 1
            continue

        # ── 空行 ──
        if not line.strip():
            i += 1
            continue

        # ── 正文 ──
        add_body(doc, line)
        i += 1

    # 文档末尾未闭合的表格
    if in_table:
        add_table(doc, table_headers, table_rows)

    doc.save(output_path)
    return output_path

# ============================================================
#  命令行
# ============================================================

def main():
    parser = argparse.ArgumentParser(description='Markdown → 标书Word')
    parser.add_argument('input', help='Markdown 输入文件')
    parser.add_argument('-o', '--output', default=None, help='输出docx路径')
    parser.add_argument('--scan', action='store_true', help='仅扫描全角半角问题，不生成Word')
    parser.add_argument('--no-auto-fix', action='store_true', help='跳过全角半角自动修复')
    args = parser.parse_args()

    if not os.path.exists(args.input):
        print(f'❌ 文件不存在: {args.input}')
        sys.exit(1)

    with open(args.input, 'r', encoding='utf-8') as f:
        md_text = f.read()

    # 扫描模式
    if args.scan:
        issues = scan_fullwidth_issues(md_text)
        if issues:
            print(f'\n⚠️  发现 {len(issues)} 处全角半角问题:\n')
            for issue in issues:
                print(f"  行{issue['line']}: [{issue['type']}] \"{issue['char']}\" → \"{issue.get('should_be', '半角')}\"")
                print(f"        上下文: …{issue['context']}…\n")
        else:
            print('✅ 未发现全角半角问题')
        sys.exit(0)

    output = args.output or os.path.splitext(args.input)[0] + '_排版.docx'
    md_to_docx(md_text, output, auto_fix=not args.no_auto_fix)
    print(f'✅ 已生成: {output}')

if __name__ == '__main__':
    main()

