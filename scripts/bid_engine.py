#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
标书自动化引擎 v3.6
一行命令完成：Markdown → 全角半角修复 → Word生成 → 格式自检 → Mermaid图表 → 企业资料注入

用法:
  python bid_engine.py 标书.md                    # 生成Word
  python bid_engine.py 标书.md -o 输出.docx       # 指定输出
  python bid_engine.py 标书.md --scan             # 仅扫描问题
  python bid_engine.py 标书.md --check            # 生成后自检
  python bid_engine.py 标书.md --暗标             # 暗标模式(去公司标识)
  python bid_engine.py 标书.md --template government  # 使用政府采购模板
  python bid_engine.py 标书.md --template enterprise  # 使用企业投标模板
  python bid_engine.py 标书.md --template engineering # 使用工程类模板
  python bid_engine.py 标书.md --config my.yaml  # 使用自定义配置
  python bid_engine.py 标书.md --profile          # 自动注入企业资料库信息
  python bid_engine.py 标书.md --mermaid-api      # 使用mermaid.ink API渲染图表(不依赖mmdc)
"""

import os, sys, re, argparse, json, urllib.request, urllib.error, base64
from datetime import datetime
from pathlib import Path

sys.stdout.reconfigure(encoding='utf-8')

try:
    from docx import Document
    from docx.shared import Pt, Inches, Twips, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("❌ 需要 python-docx: pip install python-docx")
    sys.exit(1)

try:
    import yaml
except ImportError:
    print("⚠️  未安装pyyaml，将使用默认配置。安装命令: pip install pyyaml")
    yaml = None

# ===== 默认配置 =====
DEFAULT_CONFIG = {
    'body': {'font': '宋体', 'size': 12, 'bold': False, 'first_indent_chars': 2, 'line_spacing': 1.5, 'alignment': 'justify'},
    'headings': {
        1: {'font': '宋体', 'size': 16, 'bold': True, 'space_before': 12, 'space_after': 6},
        2: {'font': '宋体', 'size': 15, 'bold': True, 'space_before': 8, 'space_after': 4},
        3: {'font': '宋体', 'size': 14, 'bold': True, 'space_before': 6, 'space_after': 3},
        4: {'font': '宋体', 'size': 14, 'bold': True, 'space_before': 4, 'space_after': 2},
    },
    'table': {'font': '宋体', 'size': 10.5, 'bold_header': True, 'style': 'Table Grid'},
    'page': {'margin_top': 2.54, 'margin_bottom': 2.54, 'margin_left': 2.00, 'margin_right': 2.00},
    'anonymous': {'enabled': False, 'replace_words': ['我公司', '本公司'], 'replace_with': '投标人', 'company_names': [], 'company_addresses': []},
    'quality': {'check_placeholder': True, 'check_punctuation': True, 'check_forbidden_words': True, 'forbidden_words': ['保证中标', '100%成功率', '最优价格', '独家技术', '唯一选择', '最高质量']},
    'mermaid': {'api_fallback': False, 'api_url': 'https://mermaid.ink/img/'},
    'profile': {'enabled': False, 'dir': ''},
}


def load_config(config_path=None, template_name=None):
    """加载配置文件，优先级：config_path > template > 自动查找config.yaml > 默认"""
    import copy
    config = copy.deepcopy(DEFAULT_CONFIG)

    # 1) 加载模板
    if template_name and yaml:
        tpl_path = Path(__file__).parent.parent / 'templates' / f'{template_name}.yaml'
        if tpl_path.exists():
            with open(tpl_path, 'r', encoding='utf-8') as f:
                tpl = yaml.safe_load(f)
                if tpl:
                    _deep_merge(config, tpl)
            print(f'📋 已加载模板: {template_name}')
        else:
            print(f'⚠️  模板不存在: {tpl_path}，使用默认配置')

    # 2) 加载config.yaml（覆盖模板）
    if config_path and yaml:
        with open(config_path, 'r', encoding='utf-8') as f:
            custom = yaml.safe_load(f)
            if custom:
                _deep_merge(config, custom)
        print(f'📋 已加载配置: {config_path}')
    elif yaml:
        default_path = Path(__file__).parent.parent / 'config.yaml'
        if default_path.exists():
            with open(default_path, 'r', encoding='utf-8') as f:
                custom = yaml.safe_load(f)
                if custom:
                    _deep_merge(config, custom)

    return config


def _deep_merge(base, override):
    """递归合并字典"""
    for k, v in override.items():
        if isinstance(v, dict) and isinstance(base.get(k), dict):
            _deep_merge(base[k], v)
        else:
            base[k] = v


def build_specs(config):
    """从配置构建内部规格字典"""
    body = config.get('body', DEFAULT_CONFIG['body'])
    headings_cfg = config.get('headings', {})
    table_cfg = config.get('table', DEFAULT_CONFIG['table'])
    page_cfg = config.get('page', DEFAULT_CONFIG['page'])

    heading_spec = {}
    for level in [1, 2, 3, 4]:
        h = headings_cfg.get(level, headings_cfg.get(str(level), DEFAULT_CONFIG['headings'].get(level, {})))
        heading_spec[level] = {
            'size': h.get('size', 12),
            'bold': h.get('bold', True),
            'name': h.get('font', '宋体'),
            'space_before': h.get('space_before', 6),
            'space_after': h.get('space_after', 3),
        }

    body_spec = {
        'font': body.get('font', '宋体'),
        'size': body.get('size', 12),
        'bold': body.get('bold', False),
        'first_indent_chars': body.get('first_indent_chars', 2),
        'line_spacing': body.get('line_spacing', 1.5),
    }

    table_spec = {
        'font': table_cfg.get('font', '宋体'),
        'size': table_cfg.get('size', 10.5),
        'bold_header': table_cfg.get('bold_header', True),
        'style': table_cfg.get('style', 'Table Grid'),
    }

    page_spec = {
        'margin_top': page_cfg.get('margin_top', 2.54),
        'margin_bottom': page_cfg.get('margin_bottom', 2.54),
        'margin_left': page_cfg.get('margin_left', 2.00),
        'margin_right': page_cfg.get('margin_right', 2.00),
    }

    return heading_spec, body_spec, table_spec, page_spec


# ===== 下划线/占位符保留逻辑 =====
UNDERLINE_PATTERNS = [
    (r'(致[：:]\\s*)(_{3,})(\\s*（[^）]+）)', r'\\1{value}\\3'),
    (r'(根据[：:]\\s*)(_{3,})(\\s*（[^）]+）)', r'\\1{value}\\3'),
    (r'(项目名称[：:]\\s*)(_{3,})', r'\\1{value}'),
    (r'([：:]\\s*)(_{3,})(\\s*（[^）]+）)', r'\\1{value}\\3'),
    (r'([：:]\\s*)(_{3,})', r'\\1{value}'),
    (r'（(招标人|投标人|供应商|采购人|代理机构|甲方|乙方|丙方)）', r'（\\1）'),
]


def preserve_underlines(text, replacements=None):
    """保留下划线占位符的格式，同时支持填入实际值。"""
    if not replacements:
        return text
    lines = text.split('\\n')
    result = []
    for line in lines:
        for pattern, template in UNDERLINE_PATTERNS:
            match = re.search(pattern, line)
            if match:
                for key, value in replacements.items():
                    if key in line:
                        line = re.sub(
                            r'(' + re.escape(key) + r'\\s*)(_{3,})(\\s*（[^）]+）)',
                            lambda m: m.group(1) + value + m.group(3),
                            line
                        )
                        if '（' not in line:
                            line = re.sub(
                                r'(' + re.escape(key) + r'\\s*)(_{3,})',
                                lambda m: m.group(1) + value,
                                line
                            )
                        break
        result.append(line)
    return '\\n'.join(result)


# ===== 文本格式化工具 =====
def set_run_font(run, font_name='宋体', size=12, bold=False):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    for attr in ['w:asciiTheme', 'w:hAnsiTheme', 'w:eastAsiaTheme']:
        try: del rFonts.attrib[qn(attr)]
        except KeyError: pass


def init_doc_styles(doc, heading_spec, body_spec):
    try:
        style = doc.styles['Normal']
        style.font.name = body_spec['font']
        style.font.size = Pt(body_spec['size'])
        style.paragraph_format.line_spacing = body_spec['line_spacing']
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    except KeyError: pass
    for level, spec in heading_spec.items():
        try:
            style = doc.styles[f'Heading {level}']
            style.font.name = spec['name']
            style.font.size = Pt(spec['size'])
            style.font.bold = spec['bold']
            rPr = style.element.get_or_add_rPr()
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rPr.insert(0, rFonts)
            rFonts.set(qn('w:ascii'), spec['name'])
            rFonts.set(qn('w:hAnsi'), spec['name'])
            rFonts.set(qn('w:eastAsia'), spec['name'])
        except KeyError: pass


# ===== 全角半角检测 =====
HALF_TO_FULL = {
    ',': '，', '.': '。', '!': '！', '?': '？', ':': '：', ';': '；',
    '(': '（', ')': '）', '[': '【', ']': '】', '<': '《', '>': '》',
    '\\\"': '\\\"', "'": "'", '~': '～', '-': '—',
}
FULLWIDTH_TRANS = str.maketrans(
    '０１２３４５６７８９ＡＢＣＤＥＦＧＨＩＪＫＬＭＮＯＰＱＲＳＴＵＶＷＸＹＺ'
    'ａｂｃｄｅｆｇｈｉｊｋｌｍｎｏｐｑｒｓｔｕｖｗｘｙｚ',
    '0123456789ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz'
)


def fix_punctuation(text):
    """全角半角自动修复，跳过编号和小数中的点号"""
    fixed = text.translate(FULLWIDTH_TRANS)
    result = []
    in_cn = True
    changes = 0
    for i, ch in enumerate(fixed):
        if ch == '.' and i > 0 and i < len(fixed) - 1:
            prev_ch = fixed[i - 1]
            next_ch = fixed[i + 1]
            if prev_ch.isdigit() and next_ch.isdigit():
                result.append(ch)
                continue
        if ch in HALF_TO_FULL and in_cn:
            result.append(HALF_TO_FULL[ch])
            changes += 1
        else:
            result.append(ch)
        if '\\u4e00' <= ch <= '\\u9fff' or '\\u3000' <= ch <= '\\u303f': in_cn = True
        elif ch.isascii() and ch.isalpha(): in_cn = False
    return ''.join(result), changes


def scan_punctuation(text):
    issues = []
    for i, line in enumerate(text.split('\\n')):
        for half, full in HALF_TO_FULL.items():
            for m in re.finditer(rf'[\\u4e00-\\u9fff]{re.escape(half)}[\\u4e00-\\u9fff]', line):
                issues.append({'line': i+1, 'type': '半角标点混入中文', 'char': half, 'should_be': full,
                    'context': line[max(0,m.start()-8):m.end()+8]})
        for m in re.finditer(r'[\\uff10-\\uff19\\uff21-\\uff3a\\uff41-\\uff5a]+', line):
            issues.append({'line': i+1, 'type': '全角数字/字母', 'char': m.group(),
                'context': line[max(0,m.start()-5):m.end()+5]})
    return issues


# ===== Markdown解析 =====
def parse_md_table(lines, start_idx):
    headers, rows, idx = [], [], start_idx
    if idx < len(lines) and lines[idx].strip().startswith('|'):
        headers = [c.strip() for c in lines[idx].strip().strip('|').split('|')]
        idx += 1
    if idx < len(lines) and re.match(r'^[\\s|\\-:]+$', lines[idx]): idx += 1
    while idx < len(lines) and lines[idx].strip().startswith('|'):
        rows.append([c.strip() for c in lines[idx].strip().strip('|').split('|')])
        idx += 1
    return headers, rows, idx


# ===== Word文档构建 =====
def add_heading(doc, text, level, heading_spec):
    spec = heading_spec.get(level, heading_spec.get(4, {}))
    p = doc.add_heading(text, level=level)
    for run in p.runs: set_run_font(run, spec['name'], spec['size'], spec['bold'])
    p.paragraph_format.space_before = Pt(spec['space_before'])
    p.paragraph_format.space_after = Pt(spec['space_after'])
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


# ===== Mermaid图表渲染（支持本地mmdc + 网络API兜底） =====
def render_mermaid(code, output_png=None, cache_dir=None, use_api=False):
    """将Mermaid代码渲染为PNG图片。

    渲染方式（按优先级）：
    1. 本地mmdc命令（需要安装mermaid-cli）
    2. mermaid.ink网络API（--mermaid-api启用，无需额外依赖）
    3. 两者都失败则返回None，调用方插入占位文本

    Args:
        code: Mermaid图表代码
        output_png: 输出PNG路径（默认自动生成）
        cache_dir: 缓存目录
        use_api: 是否优先使用网络API（兜底模式）
    """
    import subprocess, tempfile, hashlib
    if cache_dir is None:
        cache_dir = Path(tempfile.gettempdir()) / 'mermaid_cache'
    cache_dir = Path(cache_dir)
    cache_dir.mkdir(parents=True, exist_ok=True)

    content_hash = hashlib.md5(code.encode()).hexdigest()[:12]
    if output_png is None:
        output_png = cache_dir / f'mermaid_{content_hash}.png'
    else:
        output_png = Path(output_png)

    # 缓存命中
    if output_png.exists() and output_png.stat().st_size > 0:
        return str(output_png)

    # 方式1：网络API渲染（mermaid.ink）
    png_path = _render_via_api(code, output_png, content_hash)
    if png_path:
        return png_path

    # 方式2：本地mmdc渲染
    png_path = _render_via_mmdc(code, output_png, content_hash, cache_dir)
    if png_path:
        return png_path

    return None


def _render_via_api(code, output_png, content_hash):
    """使用mermaid.ink网络API渲染Mermaid图表"""
    try:
        # mermaid.ink使用base64编码的mermaid代码（URL-safe）
        # 格式：https://mermaid.ink/img/{base64}
        code_bytes = code.encode('utf-8')
        encoded = base64.urlsafe_b64encode(code_bytes).decode('ascii')
        api_url = f'https://mermaid.ink/img/{encoded}'

        req = urllib.request.Request(
            api_url,
            headers={'User-Agent': 'bid-toolkit/3.6'}
        )
        with urllib.request.urlopen(req, timeout=15) as resp:
            if resp.status == 200:
                img_data = resp.read()
                if len(img_data) > 100:  # 有效图片至少100字节
                    output_png = Path(output_png)
                    output_png.write_bytes(img_data)
                    print(f'🌐 Mermaid图表已通过API渲染: {output_png.name}')
                    return str(output_png)
        return None
    except urllib.error.URLError as e:
        print(f'⚠️  Mermaid API渲染失败（网络不可达）: {e.reason}')
        return None
    except urllib.error.HTTPError as e:
        print(f'⚠️  Mermaid API返回错误: HTTP {e.code}')
        return None
    except Exception as e:
        print(f'⚠️  Mermaid API异常: {e}')
        return None


def _render_via_mmdc(code, output_png, content_hash, cache_dir):
    """使用本地mmdc命令渲染Mermaid图表"""
    import subprocess
    mmd_file = cache_dir / f'mermaid_{content_hash}.mmd'
    mmd_file.write_text(code, encoding='utf-8')
    try:
        result = subprocess.run(
            ['mmdc', '-i', str(mmd_file), '-o', str(output_png),
             '-b', 'white', '-w', 1200],
            capture_output=True, text=True, timeout=30
        )
        if result.returncode == 0 and output_png.exists():
            print(f'🔧 Mermaid图表已通过mmdc渲染: {output_png.name}')
            return str(output_png)
        else:
            print(f'⚠️  mmdc渲染失败: {result.stderr[:200]}')
            return None
    except FileNotFoundError:
        # mmdc未安装，提示用户
        return None
    except subprocess.TimeoutExpired:
        print('⚠️  mmdc渲染超时(30s)，跳过')
        return None
    except Exception as e:
        print(f'⚠️  mmdc渲染异常: {e}')
        return None


def add_mermaid_block(doc, code, body_spec, use_api=False):
    """在Word文档中插入Mermaid图表（或占位文本）"""
    png_path = render_mermaid(code, use_api=use_api)
    if png_path:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        try:
            run.add_picture(png_path, width=Inches(5.5))
        except Exception as e:
            p.clear()
            run = p.add_run(f'[Mermaid图表渲染失败: {e}]')
            set_run_font(run, body_spec['font'], body_spec['size'], False)
            run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    else:
        # 两种渲染方式都失败
        if not _check_mmdc_installed():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run('[Mermaid图表] 请安装mermaid-cli或使用 --mermaid-api 参数')
            set_run_font(run, body_spec['font'], body_spec['size'] - 1, False)
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        else:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f'[Mermaid图表渲染失败]\\n{code[:200]}')
            set_run_font(run, body_spec['font'], body_spec['size'] - 1, False)
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)


def _check_mmdc_installed():
    """检查mmdc命令是否可用"""
    import subprocess
    try:
        subprocess.run(['mmdc', '--version'], capture_output=True, timeout=5)
        return True
    except:  # noqa: E722
        return False


# ===== 企业资料库自动注入（--profile） =====
def load_company_profile(profile_dir=None):
    """从company_profile/目录加载企业信息，返回替换字典。

    Args:
        profile_dir: 企业资料库目录路径（默认使用项目内置的company_profile/）

    Returns:
        dict: 包含企业信息的字典，key为占位符名称，value为替换文本
    """
    if profile_dir is None:
        profile_dir = Path(__file__).parent.parent / 'company_profile'
    else:
        profile_dir = Path(profile_dir)

    if not profile_dir.exists():
        print(f'⚠️  企业资料库目录不存在: {profile_dir}')
        return {}

    profile = {}

    # 读取公司基本信息
    info_file = profile_dir / 'company_info.md'
    if info_file.exists():
        text = info_file.read_text(encoding='utf-8')
        # 提取公司名称
        m = re.search(r'## 公司名称\\n\\n(.+)', text)
        if m:
            name = m.group(1).strip()
            if not name.startswith('***'):
                profile['company_name'] = name
                profile['{company_name}'] = name
        # 提取统一社会信用代码
        m = re.search(r'统一社会信用代码：(.+)', text)
        if m:
            code = m.group(1).strip()
            if not code.startswith('***'):
                profile['credit_code'] = code
        # 提取注册资本
        m = re.search(r'注册资本：(.+)', text)
        if m:
            cap = m.group(1).strip()
            if not cap.startswith('***'):
                profile['registered_capital'] = cap
        # 提取公司简介
        m = re.search(r'## 公司简介\\n\\n(.+?)(?:\\n##|$)', text, re.DOTALL)
        if m:
            intro = m.group(1).strip()
            if not intro.startswith('***'):
                profile['company_intro'] = intro

    # 读取资质信息
    qual_file = profile_dir / 'qualifications.md'
    if qual_file.exists():
        text = qual_file.read_text(encoding='utf-8')
        quals = []
        in_list = False
        for line in text.split('\\n'):
            line = line.strip()
            if line.startswith('- '):
                quals.append(line[2:])
                in_list = True
            elif in_list and line:
                quals.append(line)
        if quals:
            profile['qualifications'] = '\\n'.join(quals)

    # 读取团队信息
    team_file = profile_dir / 'team.md'
    if team_file.exists():
        text = team_file.read_text(encoding='utf-8')
        members = []
        for line in text.split('\\n'):
            line = line.strip()
            if line.startswith('- 姓名') or line.startswith('- 职务'):
                members.append(line)
        if members:
            profile['team_members'] = '\\n'.join(members)

    # 读取业绩信息
    perf_file = profile_dir / 'performance.md'
    if perf_file.exists():
        text = perf_file.read_text(encoding='utf-8')
        projs = []
        for line in text.split('\\n'):
            if line.strip().startswith('- '):
                projs.append(line.strip()[2:])
        if projs:
            profile['performance'] = '\\n'.join(projs)

    # 构建占位符替换表
    placeholders = {}
    placeholders['{company_name}'] = profile.get('company_name', '')
    placeholders['{credit_code}'] = profile.get('credit_code', '')
    placeholders['{registered_capital}'] = profile.get('registered_capital', '')
    placeholders['{company_intro}'] = profile.get('company_intro', '')
    placeholders['{qualifications}'] = profile.get('qualifications', '')
    placeholders['{team_members}'] = profile.get('team_members', '')
    placeholders['{performance}'] = profile.get('performance', '')

    return placeholders


def inject_profile(text, placeholders):
    """在Markdown文本中替换企业资料占位符。

    支持的占位符：
      {company_name}     - 公司名称
      {credit_code}      - 统一社会信用代码
      {registered_capital} - 注册资本
      {company_intro}    - 公司简介
      {qualifications}   - 资质列表
      {team_members}     - 团队成员信息
      {performance}      - 历史业绩

    Args:
        text: 原始Markdown文本
        placeholders: 占位符替换字典

    Returns:
        str: 替换后的Markdown文本
    """
    if not placeholders:
        return text

    replacements = 0
    for key, value in placeholders.items():
        if value and key in text:
            text = text.replace(key, value)
            replacements += 1

    if replacements > 0:
        print(f'📋 企业资料注入: 已替换 {replacements} 处占位符')
    return text


# ===== 目录域插入 =====
def insert_toc(doc, title='目  录', toc_levels='1-3'):
    """在文档开头插入Word目录域代码(TOC field)。"""
    new_para = OxmlElement('w:p')
    body = doc.element.body
    body.insert(0, new_para)

    toc_title = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), 'Heading1')
    pPr.append(pStyle)
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    pPr.append(jc)
    toc_title.append(pPr)
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = title
    r.append(t)
    toc_title.append(r)
    body.insert(0, toc_title)

    toc_para = OxmlElement('w:p')
    r = OxmlElement('w:r')
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    r.append(fldChar_begin)
    toc_para.append(r)

    r2 = OxmlElement('w:r')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = f' TOC \\\\o "{toc_levels}" \\\\h \\\\z \\\\u '
    r2.append(instrText)
    toc_para.append(r2)

    r3 = OxmlElement('w:r')
    fldChar_sep = OxmlElement('w:fldChar')
    fldChar_sep.set(qn('w:fldCharType'), 'separate')
    r3.append(fldChar_sep)
    toc_para.append(r3)

    r4 = OxmlElement('w:r')
    t2 = OxmlElement('w:t')
    t2.text = '（打开Word后按 Ctrl+A 全选 → F9 更新域，目录自动生成）'
    r4.append(t2)
    toc_para.append(r4)

    r5 = OxmlElement('w:r')
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    r5.append(fldChar_end)
    toc_para.append(r5)

    body.insert(2, toc_para)

    page_break = OxmlElement('w:p')
    r_brk = OxmlElement('w:r')
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    r_brk.append(br)
    page_break.append(r_brk)
    body.insert(3, page_break)

    return True


# ===== 假标题检测 =====
def check_heading_styles(doc):
    """检查文档中是否有'假标题'"""
    issues = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        style = para.style.name if para.style else ''

        if style != 'Normal':
            continue
        if len(text) < 2 or len(text) > 50:
            continue
        indent = para.paragraph_format.first_line_indent
        if indent is not None and indent != 0:
            continue

        is_bold = any(run.bold for run in para.runs if run.bold is not None)
        has_numbering = bool(re.match(
            r'^(第[一二三四五六七八九十百]+[章节条款]|[一二三四五六七八九十]+[、．.]'
            r'|\\d+[\\.\\、]|\\d+\\.\\d+|[\\(（]\\d+[\\)）])',
            text
        ))

        if is_bold and not has_numbering:
            issues.append({
                'para_idx': i + 1,
                'text': text[:40],
                'type': '疑似假标题',
                'detail': 'Normal样式+加粗但无Heading样式，目录不会收录',
                'severity': 'warn',
                'fix': '选中段落 → Word样式栏改为 Heading 1/2/3'
            })
        elif has_numbering:
            issues.append({
                'para_idx': i + 1,
                'text': text[:40],
                'type': '疑似标题未设样式',
                'detail': f'检测到"{text[:8]}"开头，可能是标题但未用Heading样式，目录不会收录',
                'severity': 'fail',
                'fix': '选中段落 → Word样式栏改为对应的 Heading 级别'
            })
    return issues


def strip_md_residue(text):
    """去除Markdown残留语法标记"""
    text = re.sub(r'!\[([^\]]*)\]\(([^)]*)\)', r'[图片: \1]', text)
    text = re.sub(r'\[([^\]]+)\]\(([^)]+)\)', r'\1（\2）', text)
    text = re.sub(r'~~(.+?)~~', r'\1', text)
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'`([^`]+)`', r'\1', text)
    text = re.sub(r'^>\\s?', '', text, flags=re.MULTILINE)
    return text


def add_body(doc, text, body_spec):
    text = strip_md_residue(text)
    p = doc.add_paragraph(text, style='Normal')
    for run in p.runs: set_run_font(run, body_spec['font'], body_spec['size'], body_spec['bold'])
    p.paragraph_format.first_line_indent = Twips(480)
    p.paragraph_format.line_spacing = body_spec['line_spacing']
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def add_table(doc, headers, rows, table_spec):
    if not headers: return None
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = table_spec['style']
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header)
        set_run_font(run, table_spec['font'], table_spec['size'], table_spec['bold_header'])
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = ''
            p = row_cells[i].paragraphs[0]
            run = p.add_run(val)
            set_run_font(run, table_spec['font'], table_spec['size'], False)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Twips(0)
    return table


# ===== 暗标模式 =====
def apply_anonymous(text, anonymous_config):
    """应用暗标模式，替换公司标识"""
    if not anonymous_config.get('enabled', False):
        return text
    replace_with = anonymous_config.get('replace_with', '投标人')
    for word in anonymous_config.get('replace_words', ['我公司', '本公司']):
        text = text.replace(word, replace_with)
    for name in anonymous_config.get('company_names', []):
        text = text.replace(name, replace_with)
    for addr in anonymous_config.get('company_addresses', []):
        text = text.replace(addr, '【地址已隐藏】')
    print(f'🔒 暗标模式：已替换公司标识')
    return text


# ===== 质检规则 =====
BID_TYPO_DICT = {
    "帐号": ("账号", "财务规范用「账」"),
    "帐户": ("账户", "财务规范用「账」"),
    "帐单": ("账单", "财务规范用「账」"),
    "帐目": ("账目", "财务规范用「账」"),
    "帐务": ("账务", "财务规范用「账」"),
    "结帐": ("结账", "财务规范用「账」"),
    "对帐": ("对账", "财务规范用「账」"),
    "记帐": ("记账", "财务规范用「账」"),
    "转帐": ("转账", "财务规范用「账」"),
    "登陆系统": ("登录系统", "系统登录用「录」"),
    "登陆平台": ("登录平台", "系统登录用「录」"),
    "登陆网站": ("登录网站", "系统登录用「录」"),
    "做为": ("作为", "作为用「作」"),
    "按装": ("安装", "安装用「安」"),
    "部暑": ("部署", "部署用「署」"),
    "布署": ("部署", "部署用「署」"),
    "软体": ("软件", "大陆用语用「软件」"),
    "硬体": ("硬件", "大陆用语用「硬件」"),
    "网路": ("网络", "大陆用语用「网络」"),
    "萤幕": ("屏幕", "大陆用语用「屏幕」"),
    "滑鼠": ("鼠标", "大陆用语用「鼠标」"),
    "讯息": ("信息", "大陆用语用「信息」"),
    "资料库": ("数据库", "大陆用语用「数据库」"),
    "视窗": ("窗口", "大陆用语用「窗口」"),
    "专案管理": ("项目管理", "大陆用语用「项目」"),
    "既使": ("即使", "即使用「即」"),
    "己经": ("已经", "已经用「已」"),
    "自已": ("自己", "自己用「己」"),
}


def check_typos(text):
    issues = []
    for wrong, (correct, note) in BID_TYPO_DICT.items():
        start = 0
        while True:
            pos = text.find(wrong, start)
            if pos < 0:
                break
            ctx_start = max(0, pos - 10)
            ctx_end = min(len(text), pos + len(wrong) + 10)
            context = text[ctx_start:ctx_end].replace('\\n', ' ')
            issues.append({
                'type': f'错别字「{wrong}」应为「{correct}」',
                'word': wrong,
                'correct': correct,
                'note': note,
                'context': context,
            })
            start = pos + len(wrong)
    return issues


CONSISTENCY_PATTERNS = {
    '投标有效期': [
        r'投标有效期\\s*[:：为]?\\s*(\\d+)\\s*天',
        r'有效期\\s*[:：为]?\\s*(\\d+)\\s*个?天',
    ],
    '总报价': [
        r'(?:总报价|总[价金额计]|合计金额|报价总[价额])\\s*[:：为是]?\\s*(\\d[\\d,]*(?:\\.\\d+)?)\\s*(万元|元|万)',
        r'投标(?:总)?报价\\s*[:：为是]?\\s*(\\d[\\d,]*(?:\\.\\d+)?)\\s*(万元|元|万)',
    ],
    '人员总数': [
        r'(?:总人数|人员[总数配置量]|配备人员)\\s*[:：为共]?\\s*(\\d+)\\s*人',
        r'共\\s*(\\d+)\\s*人',
    ],
}


# ===== md_to_docx 核心转换函数 =====
def md_to_docx(md_text, output_path, auto_fix=True, dark_mode=False, config=None, no_toc=False, use_mermaid_api=False):
    """Markdown → Word转换"""
    if config is None:
        config = DEFAULT_CONFIG

    heading_spec, body_spec, table_spec, page_spec = build_specs(config)

    if auto_fix:
        fixed_text, changes = fix_punctuation(md_text)
        if changes > 0:
            print(f'🔧 自动修复 {changes} 处全角半角混用')
            md_text = fixed_text

    if dark_mode:
        anonymous_cfg = config.get('anonymous', DEFAULT_CONFIG['anonymous'])
        anonymous_cfg['enabled'] = True
        md_text = apply_anonymous(md_text, anonymous_cfg)

    doc = Document()
    init_doc_styles(doc, heading_spec, body_spec)

    for section in doc.sections:
        section.top_margin = Cm(page_spec['margin_top'])
        section.bottom_margin = Cm(page_spec['margin_bottom'])
        section.left_margin = Cm(page_spec['margin_left'])
        section.right_margin = Cm(page_spec['margin_right'])

    lines = md_text.split('\\n')
    i, in_table = 0, False
    table_headers, table_rows = [], []
    while i < len(lines):
        line = lines[i].rstrip()
        # Mermaid代码块检测
        if line.strip().startswith('```mermaid'):
            i += 1
            mermaid_code = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                mermaid_code.append(lines[i])
                i += 1
            i += 1
            add_mermaid_block(doc, '\\n'.join(mermaid_code), body_spec, use_api=use_mermaid_api)
            continue
        if line.strip().startswith('```') and not line.strip().startswith('```mermaid'):
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                i += 1
            i += 1
            continue
        if line.startswith('|') and line.endswith('|'):
            if not in_table:
                table_headers, table_rows, i = parse_md_table(lines, i)
                in_table = True
                continue
            else:
                headers, rows, i = parse_md_table(lines, i)
                if headers: add_table(doc, headers, rows, table_spec)
                continue
        if in_table:
            add_table(doc, table_headers, table_rows, table_spec)
            in_table = False
        h_match = re.match(r'^(#{1,4})\\s+(.+)$', line)
        if h_match:
            add_heading(doc, h_match.group(2).strip(), len(h_match.group(1)), heading_spec)
            i += 1; continue
        if line.startswith('>'):
            text = line.lstrip('> ').strip()
            if text:
                p = doc.add_paragraph()
                run = p.add_run(text)
                set_run_font(run, body_spec['font'], body_spec['size'], False)
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                p.paragraph_format.left_indent = Cm(1)
            i += 1; continue
        list_match = re.match(r'^(\\s*)[-*]\\s+(.+)$', line)
        if list_match:
            text = list_match.group(2).strip()
            p = doc.add_paragraph(style='List Bullet')
            p.clear()
            run = p.add_run(text)
            set_run_font(run, body_spec['font'], body_spec['size'], False)
            p.paragraph_format.line_spacing = body_spec['line_spacing']
            i += 1; continue
        if not line.strip():
            i += 1; continue
        clean_text = re.sub(r'\\*\\*(.+?)\\*\\*', r'\\1', line)
        clean_text = re.sub(r'\\*(.+?)\\*', r'\\1', clean_text)
        add_body(doc, clean_text, body_spec)
        i += 1
    if in_table: add_table(doc, table_headers, table_rows, table_spec)
    if not no_toc:
        insert_toc(doc)
    doc.save(output_path)
    return output_path


# ===== 质检函数 =====
def check_docx(docx_path, config=None):
    if config is None:
        config = DEFAULT_CONFIG
    quality_cfg = config.get('quality', DEFAULT_CONFIG['quality'])
    forbidden_words = quality_cfg.get('forbidden_words', [])

    doc = Document(docx_path)
    results = {'pass': [], 'warn': [], 'fail': [],
        'stats': {'paragraphs': len(doc.paragraphs), 'tables': len(doc.tables),
            'total_chars': sum(len(p.text) for p in doc.paragraphs)}}

    # --- 检查1: 占位符残留 ---
    if quality_cfg.get('check_placeholder', True):
        placeholder_patterns = ['***', '____', '【待填】', '{{.*?}}', '（待补充）']
        placeholder_count = 0
        for para in doc.paragraphs:
            for pattern in placeholder_patterns:
                matches = re.findall(pattern, para.text)
                if matches:
                    placeholder_count += len(matches)
                    results['fail'].append({
                        'type': '占位符残留',
                        'detail': f'发现 {len(matches)} 处未替换占位符 "{pattern}"',
                        'context': para.text[:100],
                        'severity': 'fail',
                        'fix': '替换为实际内容'
                    })
        if placeholder_count == 0:
            results['pass'].append({'type': '占位符检查', 'detail': '未发现占位符残留'})

    # --- 检查2: 禁用词 ---
    if quality_cfg.get('check_forbidden_words', True) and forbidden_words:
        found = []
        for para in doc.paragraphs:
            for word in forbidden_words:
                if word in para.text:
                    found.append({'word': word, 'context': para.text[:100]})
        if found:
            for f in found:
                results['warn'].append({
                    'type': '禁用词',
                    'detail': f'发现禁用词 "{f["word"]}"',
                    'context': f['context'],
                    'severity': 'warn',
                    'fix': f'替换为更谦逊的表达'
                })
        else:
            results['pass'].append({'type': '禁用词检查', 'detail': '未发现禁用词'})

    # --- 检查3: 错别字 ---
    full_text = '\\n'.join(p.text for p in doc.paragraphs)
    typo_issues = check_typos(full_text)
    if typo_issues:
        for t in typo_issues:
            results['fail'].append({
                'type': t['type'],
                'detail': t['note'],
                'context': t['context'],
                'severity': 'fail',
                'fix': f'将「{t["word"]}」替换为「{t["correct"]}」'
            })
    else:
        results['pass'].append({'type': '错别字检查', 'detail': '未发现标书常见错别字'})

    # --- 检查4: 假标题 ---
    heading_issues = check_heading_styles(doc)
    if heading_issues:
        for h in heading_issues:
            target = results['fail' if h['severity'] == 'fail' else 'warn']
            target.append({
                'type': h['type'],
                'detail': h['detail'],
                'context': h['text'],
                'severity': h['severity'],
                'fix': h['fix']
            })
    else:
        results['pass'].append({'type': '标题样式', 'detail': '未发现假标题'})

    # --- 检查5: 全角半角（仅扫描） ---
    punct_issues = scan_punctuation(full_text)
    if punct_issues:
        for p in punct_issues[:5]:
            results['warn'].append({
                'type': p['type'],
                'detail': f'"{p["char"]}" 应为 "{p.get("should_be", "半角")}"',
                'context': p['context'],
                'severity': 'warn',
                'fix': '在bid_engine转换时已自动修复，强烈建议重新生成'
            })
        if len(punct_issues) > 5:
            results['warn'].append({'type': '全角半角', 'detail': f'尚有 {len(punct_issues)-5} 处未显示'})
    else:
        results['pass'].append({'type': '全角半角', 'detail': '未发现问题'})

    return results


def check_scoring_coverage(matrix_file, content_file):
    """评分项覆盖矩阵检查。"""
    with open(matrix_file, 'r', encoding='utf-8') as f:
        matrix_text = f.read()
    with open(content_file, 'r', encoding='utf-8') as f:
        content_text = f.read()

    results = {'total': 0, 'covered': 0, 'uncovered': [], 'status_mismatch': []}

    # 解析评分矩阵
    scoring_items = []
    current_item = {}
    for line in matrix_text.split('\\n'):
        line = line.strip()
        if not line or line.startswith('#'):
            continue
        if line.startswith('- ') and '|' in line:
            if current_item.get('id'):
                scoring_items.append(current_item)
            parts = line[2:].split('|')
            current_item = {
                'id': parts[0].strip() if len(parts) > 0 else '',
                'content': parts[1].strip() if len(parts) > 1 else '',
                'chapter': parts[2].strip() if len(parts) > 2 else '',
                'status': parts[3].strip() if len(parts) > 3 else '',
            }
    if current_item.get('id'):
        scoring_items.append(current_item)

    results['total'] = len(scoring_items)

    for item in scoring_items:
        # 检查内容是否在正文中出现
        found = False
        if item['content']:
            keywords = item['content'].split()
            for kw in keywords:
                if len(kw) > 2 and kw in content_text:
                    found = True
                    break

        if found:
            results['covered'] += 1
        else:
            results['uncovered'].append(item)

        if item.get('status') and item['status'] != '已响应':
            results['status_mismatch'].append(item)

    return results


def check_priority_issues(content_file, dark_mode=False):
    """P0/P1/P2合规分级检查。"""
    with open(content_file, 'r', encoding='utf-8') as f:
        text = f.read()

    if dark_mode:
        text = apply_anonymous(text, {
            'enabled': True, 'replace_words': ['我公司', '本公司'],
            'replace_with': '投标人', 'company_names': [], 'company_addresses': []
        })

    p0_rules = [
        ('投标人名称缺失', r'(?:投[标标]人|供应商|承[包包]商)[\\s：:]*[（(]?\\s*[）)]?\\s*$'),
        ('项目名称缺失', r'(?:项目名称[：:])\\s*$'),
        ('投标有效期未填写', r'(?:投标有效期|有效期)[\\s：:]*[（(]?\\s*[）)]?\\s*$'),
        ('报价金额未填写', r'(?:总报价|投标报价|报价金额)[\\s：:]*[（(]?\\s*[）)]?'),
    ]
    p1_rules = [
        ('联系人信息缺失', r'(?:联系人|授权代表|项目负责人)[\\s：:]*\\s*$'),
        ('联系电话缺失', r'(?:联系电话|手机|电话)[\\s：:]*\\s*$'),
        ('邮箱地址缺失', r'(?:邮箱|E-?mail|电子邮箱)[\\s：:]*\\s*$'),
        ('服务承诺未写', r'(?:服务承诺|售后|质保期)[\\s：:]*\\s*$'),
    ]
    p2_rules = [
        ('公司简介缺失', r'(?:公司简介|企业介绍)[\\s：:]*\\s*$'),
        ('团队介绍缺失', r'(?:项目团队|人员配置|组织架构)[\\s：:]*\\s*$'),
        ('类似业绩缺失', r'(?:类似项目|相关业绩|成功案例)[\\s：:]*\\s*$'),
    ]

    issues = []
    for desc, pattern in p0_rules:
        for m in re.finditer(pattern, text, re.MULTILINE):
            issues.append({'level': 'P0', 'desc': desc, 'match': m.group()[:50], 'line': text[:m.start()].count('\\n') + 1})
    for desc, pattern in p1_rules:
        for m in re.finditer(pattern, text, re.MULTILINE):
            issues.append({'level': 'P1', 'desc': desc, 'match': m.group()[:50], 'line': text[:m.start()].count('\\n') + 1})
    for desc, pattern in p2_rules:
        for m in re.finditer(pattern, text, re.MULTILINE):
            issues.append({'level': 'P2', 'desc': desc, 'match': m.group()[:50], 'line': text[:m.start()].count('\\n') + 1})

    p0_count = sum(1 for i in issues if i['level'] == 'P0')
    p1_count = sum(1 for i in issues if i['level'] == 'P1')
    deliverable = (p0_count == 0 and p1_count <= 2)

    return {'issues': issues, 'p0_count': p0_count, 'p1_count': p1_count, 'deliverable': deliverable}


def print_check_report(results):
    """打印质检报告"""
    print('\\\\n' + '=' * 60)
    print('📋 标书质检报告')
    print('=' * 60)
    s = results['stats']
    print(f'📊 统计: {s["paragraphs"]} 段落, {s["tables"]} 表格, {s["total_chars"]} 字符')
    print()
    for item in results['pass']:
        print(f'  ✅ {item["type"]}: {item["detail"]}')
    for item in results['warn']:
        print(f'  ⚠️  {item["type"]}: {item["detail"]}')
        if 'context' in item:
            ctx = item['context'][:60]
            print(f'      上下文: …{ctx}…')
    for item in results['fail']:
        print(f'  ❌ {item["type"]}: {item["detail"]}')
        if 'context' in item:
            ctx = item['context'][:60]
            print(f'      上下文: …{ctx}…')
        if 'fix' in item:
            fix = item['fix']
            print(f'      💡 修复: {fix}')
    total_issues = len(results['warn']) + len(results['fail'])
    nf = len(results['fail'])
    nw = len(results['warn'])
    if total_issues == 0:
        print('\\\\n🎉 全部通过！没有发现任何问题。')
    else:
        print(f'\\\\n📈 共 {total_issues} 个问题（{nf} 个错误, {nw} 个警告）')


# ===== 行业自动检测 =====
def detect_industry(text, detection_path=None):
    """从招标文件文本中自动检测行业类型，返回模板名称

    使用 bid_type_detection.yaml 中的关键词规则打分，
    强信号2分，中信号1分，达到阈值即定性。
    """
    if detection_path is None:
        detection_path = Path(__file__).parent.parent / 'templates' / 'bid_type_detection.yaml'
    if not detection_path.exists():
        return None

    if yaml is None:
        print('⚠️  未安装pyyaml，无法使用行业检测')
        return None

    try:
        with open(detection_path, 'r', encoding='utf-8') as f:
            rules = yaml.safe_load(f)
    except Exception as e:
        print(f'⚠️  行业检测规则加载失败: {e}')
        return None

    detection_rules = rules.get('detection_rules', {})
    if not detection_rules:
        return None

    text_lower = text.lower()
    scores = {}

    for type_key, type_config in detection_rules.items():
        if not isinstance(type_config, dict):
            continue
        score = 0
        keywords = type_config.get('keywords', {})

        # 强信号：每个2分
        for kw in keywords.get('strong', []):
            count = text_lower.count(kw.lower())
            if count > 0:
                score += count * 2

        # 中信号：每个1分
        for kw in keywords.get('medium', []):
            count = text_lower.count(kw.lower())
            if count > 0:
                score += count * 1

        # 排除条件
        exclude = type_config.get('exclude_if_dominant', [])
        exclude_score = 0
        for kw in exclude:
            exclude_score += text_lower.count(kw.lower())

        if exclude_score > score:
            score = 0  # 排除信号压倒性，不判定为此类型

        threshold = type_config.get('confidence_threshold', 3)
        if score >= threshold:
            scores[type_key] = score

    if not scores:
        return None

    # 取最高分
    best_type = max(scores, key=scores.get)
    template_map = {
        'engineering': 'engineering',
        'goods': 'government',  # 货物标用政府采购模板
        'service': 'government',  # 服务标用政府采购模板
    }

    # 显示检测结果
    template_name = template_map.get(best_type, 'government')
    type_name = {
        'engineering': '工程类',
        'goods': '货物类',
        'service': '服务类',
    }.get(best_type, best_type)
    print(f'🔍 行业自动检测: {type_name} (得分: {scores[best_type]}, 模板: {template_name})')

    return template_name


def run_deflavor_scan(text, mode='bid'):
    """对文本运行AI味雷达检测，返回检测报告

    集成 ai-flavor-radar 到流水线，检测AI味文本并输出报告。
    """
    # 添加 ai-flavor-radar 到路径
    radar_path = Path(os.environ.get("AI_FLAVOR_RADAR_PATH", "/tmp/ai-flavor-radar"))
    if not radar_path.exists():
        print('⚠️  AI味雷达未安装 (路径: /tmp/ai-flavor-radar)')
        print('💡  安装: git clone https://github.com/charlotty2026/ai-flavor-radar.git /tmp/ai-flavor-radar')
        return None

    sys.path.insert(0, str(radar_path))
    try:
        from ai_flavor_radar import FlavorRadar, format_text_report
    except ImportError as e:
        print(f'⚠️  AI味雷达导入失败: {e}')
        return None

    radar = FlavorRadar(mode=mode)
    result = radar.scan(text, file_path='<bid_engine>')

    print('\n' + '=' * 60)
    print('🧪 AI味雷达检测报告')
    print('=' * 60)
    print(f'📊 评分: {result.score}/100 ({result.grade})')
    print(f'🔍 命中: {result.hit_count} 处 (fatal={result.fatal_count}, high={result.high_count}, medium={result.medium_count}, low={result.low_count})')
    print()

    if result.hits:
        report = format_text_report(result, use_color=False)
        # 取前20行精简展示
        lines = report.split('\n')
        for line in lines[:30]:
            print(line)
        if len(lines) > 30:
            print(f'  ... 还有 {len(lines) - 30} 行 ...')
    else:
        print('✅ 未检测到AI味问题')

    print('=' * 60)
    return result


# ===== 命令行入口 =====
def main():
    parser = argparse.ArgumentParser(description='标书自动化引擎 v3.6')
    parser.add_argument('input', nargs='?', default=None, help='Markdown输入文件（--check-scoring/--check-priority模式不需要）')
    parser.add_argument('-o', '--output', default=None, help='输出docx路径')
    parser.add_argument('--scan', action='store_true', help='仅扫描全角半角')
    parser.add_argument('--check', action='store_true', help='生成后自检')
    parser.add_argument('--no-fix', action='store_true', help='跳过全角修复')
    parser.add_argument('--暗标', action='store_true', help='暗标模式(去公司标识)')
    parser.add_argument('--template', type=str, default=None,
                        choices=['government', 'enterprise', 'engineering'],
                        help='使用预设模板: government=政府采购, enterprise=企业投标, engineering=工程类')
    parser.add_argument('--detect-industry', action='store_true', default=False,
                        help='自动检测行业类型并加载对应模板（覆盖--template）')
    parser.add_argument('--deflavor', action='store_true', default=False,
                        help='生成后自动运行AI味雷达检测，输出去AI味报告')
    parser.add_argument('--config', type=str, default=None, help='指定config.yaml配置文件路径')
    parser.add_argument('--json', action='store_true', help='JSON输出')
    parser.add_argument('--no-toc', action='store_true', help='不插入目录域（默认自动生成目录）')
    parser.add_argument('--check-scoring', nargs=2, metavar=('MATRIX', 'CONTENT'), help='评分项覆盖矩阵检查')
    parser.add_argument('--check-priority', type=str, default=None, help='P0/P1/P2合规分级检查')
    parser.add_argument('--profile', type=str, nargs='?', const='auto', default=None,
                        help='自动注入企业资料库信息。接受路径参数，不加参数使用内置company_profile/')
    parser.add_argument('--mermaid-api', action='store_true', default=False,
                        help='使用mermaid.ink网络API渲染图表（不依赖本地mmdc命令）')
    args = parser.parse_args()

    if args.check_scoring:
        matrix_file, content_file = args.check_scoring
        if not os.path.exists(matrix_file):
            print(f'❌ 矩阵文件不存在: {matrix_file}'); sys.exit(1)
        if not os.path.exists(content_file):
            print(f'❌ 正文文件不存在: {content_file}'); sys.exit(1)
        results = check_scoring_coverage(matrix_file, content_file)
        print('\\n' + '=' * 60)
        print('📋 评分项覆盖矩阵检查报告')
        print('=' * 60)
        print(f'📊 评分项总数: {results["total"]}')
        print(f'✅ 已覆盖: {results["covered"]}')
        print(f'❌ 未覆盖: {len(results["uncovered"])}')
        print(f'⚠️  状态异常: {len(results["status_mismatch"])}')
        if results['uncovered']:
            print('\\n❌ 未覆盖的评分项:')
            for row in results['uncovered']:
                print(f'  {row["id"]} | {row["content"]} | 应在: {row["chapter"]}')
        if results['status_mismatch']:
            print('\\n⚠️  状态未标"已响应"的评分项:')
            for row in results['status_mismatch']:
                print(f'  {row["id"]} | {row["content"]} | 状态: {row["status"]}')
        status = '✅ PASS' if not results['uncovered'] and not results['status_mismatch'] else '❌ FAIL'
        print(f'\\n🏁 结论: {status}')
        print('=' * 60)
        sys.exit(0)

    if args.check_priority:
        if not os.path.exists(args.check_priority):
            print(f'❌ 文件不存在: {args.check_priority}'); sys.exit(1)
        results = check_priority_issues(args.check_priority, dark_mode=args.暗标)
        print('\\n' + '=' * 60)
        print('📋 P0/P1/P2 合规分级检查报告')
        print('=' * 60)
        print(f'🔴 P0-致命: {results["p0_count"]} 个')
        print(f'🟡 P1-重要: {results["p1_count"]} 个')
        for issue in results['issues']:
            icon = '🔴' if issue['level'] == 'P0' else '🟡'
            print(f'  {icon} 行{issue["line"]} | {issue["desc"]} | 匹配: "{issue["match"]}"')
        deliverable = '✅ 可交付' if results['deliverable'] else '❌ 不可交付（P0>0或P1>2）'
        print(f'\\n🏁 结论: {deliverable}')
        print('=' * 60)
        sys.exit(0)

    if not args.input or not os.path.exists(args.input):
        if not args.input:
            print('❌ 请指定Markdown文件')
        else:
            print(f'❌ 文件不存在: {args.input}')
        print('💡 提示: 请指定一个已有的Markdown文件。试试:')
        print('   bid engine samples/某外包项目标书.md    # 用内置示例')
        print('   bid engine 你的标书.md -o 输出.docx       # 用你自己的文件')
        print('   bid engine 你的标书.md --profile          # 自动注入企业资料')
        print('   bid engine 你的标书.md --mermaid-api      # 使用网络API渲染图表')
        print('   bid list                                # 查看所有命令')
        sys.exit(1)

    # --detect-industry: 自动检测行业类型并加载对应模板（覆盖--template）
    if args.detect_industry and args.input and os.path.exists(args.input):
        # 读取文件头2000字用于检测（足够判断行业类型）
        with open(args.input, 'r', encoding='utf-8') as f:
            head_text = f.read(8000)
        detected_template = detect_industry(head_text)
        if detected_template:
            args.template = detected_template
            print(f'📋 自动匹配模板: {detected_template}')
        else:
            print('⚠️  未能自动识别行业类型，使用默认配置')

    config = load_config(config_path=args.config, template_name=args.template)

    if args.check and args.input.endswith('.docx'):
        results = check_docx(args.input, config=config)
        if args.json: print(json.dumps(results, ensure_ascii=False, indent=2))
        else: print_check_report(results)
        sys.exit(0)

    with open(args.input, 'r', encoding='utf-8') as f: md_text = f.read()

    if args.scan:
        issues = scan_punctuation(md_text)
        if args.json:
            print(json.dumps({'issues': issues, 'count': len(issues)}, ensure_ascii=False))
        elif issues:
            print(f'\\n⚠️  发现 {len(issues)} 处全角半角问题:\\n')
            for issue in issues:
                print(f'  行{issue["line"]}: [{issue["type"]}] "{issue["char"]}"')
                print(f'        上下文: …{issue["context"]}…\\n')
        else:
            print('✅ 未发现全角半角问题')
        sys.exit(0)

    # --profile: 企业资料库自动注入
    if args.profile:
        if args.profile == 'auto':
            placeholders = load_company_profile()
        else:
            placeholders = load_company_profile(args.profile)
        if placeholders:
            inject_count = sum(1 for v in placeholders.values() if v)
            print(f'📋 企业资料库加载完成: {inject_count} 个字段可用')
            md_text = inject_profile(md_text, placeholders)
        else:
            print('⚠️  企业资料库未找到有效数据，请先填写 company_profile/ 目录下的模板文件')

    output = args.output or os.path.splitext(args.input)[0] + '_排版.docx'
    md_to_docx(md_text, output, auto_fix=not args.no_fix, dark_mode=args.暗标,
               config=config, no_toc=args.no_toc, use_mermaid_api=args.mermaid_api)
    print(f'✅ 已生成: {output}')

    # --deflavor: 生成后自动运行AI味雷达检测
    if args.deflavor:
        run_deflavor_scan(md_text, mode='bid')

    if args.check:
        results = check_docx(output, config=config)
        if args.json: print(json.dumps(results, ensure_ascii=False, indent=2))
        else: print_check_report(results)
    sys.exit(0)


if __name__ == '__main__': main()